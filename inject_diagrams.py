"""
Injects SRS diagrams into the correct sections of AURA_WEAR_FYP_Document.docx.

Diagram mapping (from SRS v4):
  image1.png  -> Figure 1: Use Case Diagram           -> after "3.1 Use Case Diagram"
  image2.png  -> Figure 2: Activity Diagram            -> after "1.10 System Design Approach" (Ch1)
  image3.png  -> Figure 3: Sequence Diagram (Summary)  -> after "3.4 System Sequence Diagrams" heading
  image4.png  -> Figure 4: System Sequence Diagram     -> after SSD section intro
  image5.png  -> Figure 5: Detailed Sequence Diagram   -> after Figure 4 in SSD section
  image6.png  -> Figure 7: System Architecture Diagram -> after "1.10 System Design Approach" body
  image7.png  -> Figure 8: ER Diagram                  -> after "3.5 Schema Diagram"
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DIAGRAMS = os.path.join(os.path.dirname(__file__), 'doc_parts', 'diagrams')
INPUT  = os.path.join(os.path.dirname(__file__), 'AURA_WEAR_FYP_Document.docx')
OUTPUT = os.path.join(os.path.dirname(__file__), 'AURA_WEAR_FYP_Document.docx')


def add_figure(doc, para_index, img_path, caption, width=Inches(5.5)):
    """Insert an image + caption after paragraph at para_index."""
    # We rebuild paragraphs list after each insertion because inserting
    # elements shifts indices. We use the XML element directly.
    target_para = doc.paragraphs[para_index]

    # Insert image paragraph after target
    from docx.oxml import OxmlElement
    from lxml import etree

    # Add caption paragraph
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(4)
    cap_p.paragraph_format.space_after = Pt(10)
    cap_run = cap_p.add_run(caption)
    cap_run.italic = True
    cap_run.font.name = 'Times New Roman'
    cap_run.font.size = Pt(11)

    # Add image paragraph
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(6)
    img_p.paragraph_format.space_after = Pt(4)
    img_run = img_p.add_run()
    img_run.add_picture(img_path, width=width)

    # Move both new paragraphs to right after target_para in the XML
    body = doc.element.body
    ref_elem = target_para._element

    # Move img_p before cap_p (they were appended at end)
    body.remove(img_p._element)
    body.remove(cap_p._element)
    ref_elem.addnext(cap_p._element)
    ref_elem.addnext(img_p._element)


def find_para(doc, search_text, start=0):
    """Find first paragraph index containing search_text at or after start."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if search_text.lower() in p.text.lower():
            return i
    return -1


def main():
    print(f"Loading: {INPUT}")
    doc = Document(INPUT)
    total = len(doc.paragraphs)
    print(f"Paragraphs before insertion: {total}")

    # ── 1. Use Case Diagram (image1.png) ───────────────────────────
    # Insert after the "[Use Case Diagram - See attached UML diagram]" placeholder
    idx = find_para(doc, 'Use Case Diagram - See attached UML diagram')
    if idx >= 0:
        add_figure(doc, idx,
                   os.path.join(DIAGRAMS, 'image1.png'),
                   'Figure 1: Use Case Diagram — AURA-WEAR AI-Integrated Custom Clothing Website',
                   width=Inches(5.8))
        print(f"  ✓ Figure 1 (Use Case Diagram) inserted at para {idx}")
    else:
        print("  ✗ Use Case Diagram placeholder not found")

    # ── 2. Activity Diagram (image2.png) ───────────────────────────
    # Insert after "1.10 System Design Approach" section
    idx = find_para(doc, '1.10 System Design Approach')
    if idx >= 0:
        # Find end of this section (next heading)
        end = find_para(doc, '1.11 Process Model', idx + 1)
        insert_at = end - 1 if end > idx else idx + 3
        add_figure(doc, insert_at,
                   os.path.join(DIAGRAMS, 'image2.png'),
                   'Figure 2: Activity Diagram — Customer Order Flow',
                   width=Inches(5.5))
        print(f"  ✓ Figure 2 (Activity Diagram) inserted at para {insert_at}")
    else:
        print("  ✗ System Design Approach not found")

    # ── 3. System Architecture Diagram (image6.png) ────────────────
    # Insert after "1.10 System Design Approach" body text (after Figure 2)
    idx = find_para(doc, 'Figure 2: Activity Diagram')
    if idx >= 0:
        add_figure(doc, idx,
                   os.path.join(DIAGRAMS, 'image6.png'),
                   'Figure 3: System Architecture Diagram — Layered 4-Tier Architecture',
                   width=Inches(5.5))
        print(f"  ✓ Figure 3 (Architecture Diagram) inserted at para {idx}")
    else:
        print("  ✗ Figure 2 caption not found for architecture placement")

    # ── 4. ER Diagram (image7.png) ─────────────────────────────────
    # Insert after "3.5 Schema Diagram" section heading
    idx = find_para(doc, '3.5 Schema Diagram')
    if idx >= 0:
        add_figure(doc, idx,
                   os.path.join(DIAGRAMS, 'image7.png'),
                   'Figure 4: Entity-Relationship (ER) Diagram — Database Schema',
                   width=Inches(5.8))
        print(f"  ✓ Figure 4 (ER Diagram) inserted at para {idx}")
    else:
        print("  ✗ Schema Diagram section not found")

    # ── 5. Summary Sequence Diagram (image3.png) ───────────────────
    # Insert after "3.4 System Sequence Diagrams" heading
    idx = find_para(doc, '3.4 System Sequence Diagrams')
    if idx >= 0:
        # Insert right after the heading + intro paragraph
        add_figure(doc, idx + 1,
                   os.path.join(DIAGRAMS, 'image3.png'),
                   'Figure 5: Sequence Diagram (Summary Level) — Customer to Admin Order Flow',
                   width=Inches(5.5))
        print(f"  ✓ Figure 5 (Sequence Diagram) inserted after para {idx+1}")
    else:
        print("  ✗ SSD section heading not found")

    # ── 6. SSD Detail (image4.png) ─────────────────────────────────
    # Insert after the summary sequence diagram caption
    idx = find_para(doc, 'Figure 5: Sequence Diagram (Summary Level)')
    if idx >= 0:
        add_figure(doc, idx,
                   os.path.join(DIAGRAMS, 'image4.png'),
                   'Figure 6: System Sequence Diagram (SSD) — Detailed Message Flow',
                   width=Inches(5.5))
        print(f"  ✓ Figure 6 (SSD Detail) inserted at para {idx}")
    else:
        print("  ✗ Summary Sequence Diagram caption not found")

    # ── 7. Detailed Sequence Diagram (image5.png) ──────────────────
    # Insert after Class Diagram section in Chapter 4
    idx = find_para(doc, '4.1 Class Diagram')
    if idx >= 0:
        end = find_para(doc, '4.2 Project Code', idx + 1)
        insert_at = end - 1 if end > idx else idx + 3
        add_figure(doc, insert_at,
                   os.path.join(DIAGRAMS, 'image5.png'),
                   'Figure 7: Class Diagram — AURA-WEAR System Classes and Relationships',
                   width=Inches(5.0))
        print(f"  ✓ Figure 7 (Class Diagram) inserted at para {insert_at}")
    else:
        print("  ✗ Class Diagram section not found")

    doc.save(OUTPUT)
    print(f"\n✓ Saved with diagrams: {OUTPUT}")
    print(f"  Paragraphs after insertion: {len(doc.paragraphs)}")
    import os as _os
    print(f"  File size: {_os.path.getsize(OUTPUT)/1024:.1f} KB")


if __name__ == '__main__':
    main()
