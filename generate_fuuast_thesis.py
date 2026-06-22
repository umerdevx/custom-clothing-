#!/usr/bin/env python3
"""
AURA-WEAR FUUAST Thesis Generator
Produces a properly formatted thesis document following the FUUAST template.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import copy, os

# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_toc_entry(doc, text, level=0):
    """Add a single TOC line with dot leaders using a tab stop."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Indent by level
    indent_map = {0: 0, 1: Inches(0.25), 2: Inches(0.5), 3: Inches(0.75)}
    para.paragraph_format.left_indent = indent_map.get(level, Inches(0.75))

    # Right-aligned tab with dot leader at 5.8"
    pPr = para._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:leader'), 'dot')
    tab_el.set(qn('w:pos'), '8352')   # ~5.8 inches in twips (1" = 1440)
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    run = para.add_run(text + '\t')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = (level == 0)
    return para


def add_manual_toc(doc):
    """Build the complete static Table of Contents matching FUUAST template style."""

    toc_entries = [
        # (text, level)
        # ── FRONT MATTER ──────────────────────────────────────────────────────
        ('CHAPTER # 1: PROPOSAL', 0),
        ('Revision History', 1),
        ('1.1  Introduction', 1),
        ('1.2  Problem Statement', 1),
        ('1.3  Proposed System', 1),
        ('1.4  Benefits of the Proposed System', 1),
        ('1.5  Scope', 1),
        ('1.6  Survey Analysis', 1),
        ('1.6.1  Survey Analysis Table', 2),
        ('1.7  Modules & Submodules', 1),
        ('1.7.1  Security Management', 2),
        ('1.7.2  User Profile Management', 2),
        ('1.7.3  Product Catalog Management', 2),
        ('1.7.4  Product Customization', 2),
        ('1.7.5  AI Chatbot Management', 2),
        ('1.7.6  Order Management', 2),
        ('1.7.7  Payment Management', 2),
        ('1.7.8  Admin Dashboard Management', 2),
        ('1.8  Primary Actor', 1),
        ('1.9  Tools & Technologies', 1),
        ('1.9.1  Front-End Tools', 2),
        ('1.9.2  Back-End Tools', 2),
        ('1.9.3  AI / Automation Tools', 2),
        ('1.9.4  Database', 2),
        ('1.9.5  Deployment', 2),
        ('1.10  System Design Approach', 1),
        ('1.11  Process Model Used', 1),
        ('1.12  Modelling Techniques / Tools Used', 1),
        ('1.13  Limitation / Constraint', 1),
        ('1.14  References', 1),

        # ── CHAPTER 2 ─────────────────────────────────────────────────────────
        ('CHAPTER # 2: ANALYSIS', 0),
        ('2.1  Introduction', 1),
        ('2.1.1  Purpose', 2),
        ('2.1.2  Scope', 2),
        ('2.1.3  Definitions, Acronyms, and Abbreviations', 2),
        ('2.1.4  Overview', 2),
        ('2.2  Functional Requirements', 1),
        ('2.2.1  Security Management', 2),
        ('2.2.1.1  Process SignUp', 3),
        ('2.2.1.2  Process SignIn', 3),
        ('2.2.1.3  Forgot Password', 3),
        ('2.2.1.4  Change Account Password', 3),
        ('2.2.1.5  JWT Authentication', 3),
        ('2.2.1.6  OTP Verification', 3),
        ('2.2.1.7  Role-Based Access Control', 3),
        ('2.2.1.8  Session Management', 3),
        ('2.2.2  User Profile Management', 2),
        ('2.2.2.1  Add Customer Profile', 3),
        ('2.2.2.2  Update Customer Profile', 3),
        ('2.2.2.3  View Customer Profile', 3),
        ('2.2.2.4  Manage Customer Account Status', 3),
        ('2.2.3  Product Catalog Management', 2),
        ('2.2.3.1  Add New Product', 3),
        ('2.2.3.2  Update Product Details', 3),
        ('2.2.3.3  Enable or Disable Product', 3),
        ('2.2.3.4  Search Product', 3),
        ('2.2.3.5  View All Products', 3),
        ('2.2.3.6  Filter Products by Category', 3),
        ('2.2.4  Product Customization', 2),
        ('2.2.4.1  Select Fabric Type', 3),
        ('2.2.4.2  Select Color', 3),
        ('2.2.4.3  Upload Custom Logo', 3),
        ('2.2.4.4  Select Stitching Style', 3),
        ('2.2.4.5  Select Print Method', 3),
        ('2.2.4.6  Select Wash Effect', 3),
        ('2.2.4.7  View Real-Time 2D Design Preview', 3),
        ('2.2.4.8  Calculate Dynamic Price', 3),
        ('2.2.5  AI Chatbot Management', 2),
        ('2.2.5.1  Send Query to Chatbot', 3),
        ('2.2.5.2  Get AI Design Recommendation (RAG)', 3),
        ('2.2.5.3  View FAQ via Chatbot', 3),
        ('2.2.5.4  View Chat History', 3),
        ('2.2.5.5  Admin Manage FAQ Entries', 3),
        ('2.2.6  Order Management', 2),
        ('2.2.6.1  Create New Order', 3),
        ('2.2.6.2  Search Order History', 3),
        ('2.2.6.3  View Order History', 3),
        ('2.2.6.4  Cancel Order', 3),
        ('2.2.6.5  Track Order Status', 3),
        ('2.2.6.6  Reorder Previous Order', 3),
        ('2.2.7  Payment Management', 2),
        ('2.2.7.1  Record Online Payment', 3),
        ('2.2.7.2  Record Cash on Delivery', 3),
        ('2.2.7.3  Generate Digital Receipt', 3),
        ('2.2.7.4  View All Payments and Transactions', 3),
        ('2.2.7.5  Download Invoice', 3),
        ('2.2.8  Admin Dashboard Management', 2),
        ('2.2.8.1  Manage Products (Admin)', 3),
        ('2.2.8.2  Manage Customer Accounts (Admin)', 3),
        ('2.2.8.3  View and Update Order Management (Admin)', 3),
        ('2.2.8.4  View Inventory and Stock Alerts (Admin)', 3),
        ('2.2.8.5  Generate Sales Report', 3),
        ('2.2.8.6  Generate Order Report', 3),
        ('2.2.8.7  View Analytics Dashboard', 3),
        ('2.2.8.8  Send Notification to Customers', 3),
        ('2.2.8.9  Manage Content (Terms, Privacy, FAQ)', 3),
        ('2.2.8.10  Manage Coupon/Promo Codes', 3),
        ('2.2.8.11  Export Data as CSV', 3),
        ('2.3  Non-Functional Requirements', 1),
        ('2.3.1  Performance Requirements', 2),
        ('2.3.2  Usability Requirements', 2),
        ('2.3.3  Security Requirements', 2),
        ('2.3.4  Reliability Requirements', 2),
        ('2.3.5  Scalability Requirements', 2),
        ('2.3.6  Maintainability & Support Requirements', 2),
        ('2.3.7  Compatibility Requirements', 2),
        ('2.3.8  Support Requirements', 2),
        ('2.4  External Interface Requirements', 1),
        ('2.4.1  User Interfaces', 2),
        ('2.4.2  Hardware Interfaces', 2),
        ('2.4.3  Software Interfaces', 2),
        ('2.4.4  Communication Interfaces', 2),

        # ── CHAPTER 3 ─────────────────────────────────────────────────────────
        ('CHAPTER # 3: DESIGN', 0),
        ('3.1  Use Case Diagram', 1),
        ('3.2  Fully Dressed Use Cases', 1),
        ('UC-01: Process SignUp', 2),
        ('UC-02: Process SignIn', 2),
        ('UC-03: Forgot Password', 2),
        ('UC-04: Change Account Password', 2),
        ('UC-05: JWT Authentication', 2),
        ('UC-06: OTP Verification', 2),
        ('UC-07: Role-Based Access Control', 2),
        ('UC-08: Session Management', 2),
        ('UC-09: Add Customer Profile', 2),
        ('UC-10: Update Customer Profile', 2),
        ('UC-11: View Customer Profile', 2),
        ('UC-12: Manage Customer Account Status', 2),
        ('UC-13: Add New Product', 2),
        ('UC-14: Update Product Details', 2),
        ('UC-15: Enable or Disable Product', 2),
        ('UC-16: Search Product', 2),
        ('UC-17: View All Products', 2),
        ('UC-18: Filter Products by Category', 2),
        ('UC-19: Select Fabric Type', 2),
        ('UC-20: Select Color', 2),
        ('UC-21: Upload Custom Logo', 2),
        ('UC-22: Select Stitching Style', 2),
        ('UC-23: Select Print Method', 2),
        ('UC-24: Select Wash Effect', 2),
        ('UC-25: View Real-Time 2D Design Preview', 2),
        ('UC-26: Calculate Dynamic Price', 2),
        ('UC-27: Send Query to Chatbot', 2),
        ('UC-28: Get AI Design Recommendation (RAG)', 2),
        ('UC-29: View FAQ via Chatbot', 2),
        ('UC-30: View Chat History', 2),
        ('UC-31: Admin Manage FAQ Entries', 2),
        ('UC-32: Create New Order', 2),
        ('UC-33: Search Order History', 2),
        ('UC-34: View Order History', 2),
        ('UC-35: Cancel Order', 2),
        ('UC-36: Track Order Status', 2),
        ('UC-37: Reorder Previous Order', 2),
        ('UC-38: Record Online Payment', 2),
        ('UC-39: Record Cash on Delivery', 2),
        ('UC-40: Generate Digital Receipt', 2),
        ('UC-41: View All Payments and Transactions', 2),
        ('UC-42: Download Invoice', 2),
        ('UC-43: Manage Products — Admin', 2),
        ('UC-44: Manage Customer Accounts — Admin', 2),
        ('UC-45: View and Update Order Management — Admin', 2),
        ('UC-46: View Inventory and Stock Alerts — Admin', 2),
        ('UC-47: Generate Sales Report', 2),
        ('UC-48: Generate Order Report', 2),
        ('UC-49: View Analytics Dashboard', 2),
        ('UC-50: Send Notification to Customers', 2),
        ('UC-51: Manage Content — Terms, Privacy, FAQ', 2),
        ('UC-52: Manage Coupon/Promo Codes', 2),
        ('UC-53: Export Data as CSV', 2),
        ('UC-54: Manage Admin Users', 2),
        ('UC-55: View Customer Reviews', 2),
        ('UC-56: Generate Performance Report', 2),
        ('3.3  Domain Model', 1),
        ('3.4  System Sequence Diagrams', 1),
        ('SSD:1  Process SignUp', 2),
        ('SSD:2  Process SignIn', 2),
        ('SSD:3  Forgot Password', 2),
        ('SSD:4  Reset Password', 2),
        ('SSD:5 — SSD:56  (All System Sequence Diagrams)', 2),
        ('3.5  Schema Diagram', 1),
        ('Table: users', 2),
        ('Table: products', 2),
        ('Table: orders', 2),
        ('Table: order_items', 2),
        ('Table: payments', 2),
        ('Table: chat_logs', 2),
        ('Table: inventory', 2),
        ('Table: coupons', 2),
        ('Table: faq', 2),

        # ── CHAPTER 4 ─────────────────────────────────────────────────────────
        ('CHAPTER # 4: CONSTRUCTION', 0),
        ('4.1  Class Diagram', 1),
        ('4.2  Project Code (Business Logic)', 1),
        ('4.2.1  Backend: FastAPI Routes', 2),
        ('4.2.2  Frontend: JavaScript Logic', 2),
        ('4.3  Sample Queries', 1),

        # ── CHAPTER 5 ─────────────────────────────────────────────────────────
        ('CHAPTER # 5: TESTING', 0),
        ('5.1  Test Cases', 1),
        ('TC-01  Process SignUp', 2),
        ('TC-02  Process SignIn', 2),
        ('TC-03  Forgot Password', 2),
        ('TC-04  Change Account Password', 2),
        ('TC-05  JWT Authentication', 2),
        ('TC-06  OTP Verification', 2),
        ('TC-07  Role-Based Access Control', 2),
        ('TC-08  Session Management', 2),
        ('TC-09  Add Customer Profile', 2),
        ('TC-10  Update Customer Profile', 2),
        ('TC-11  View Customer Profile', 2),
        ('TC-12  Manage Customer Account Status', 2),
        ('TC-13  Add New Product', 2),
        ('TC-14  Update Product Details', 2),
        ('TC-15  Enable or Disable Product', 2),
        ('TC-16  Search Product', 2),
        ('TC-17  View All Products', 2),
        ('TC-18  Filter Products by Category', 2),
        ('TC-19  Select Fabric Type', 2),
        ('TC-20  Select Color', 2),
        ('TC-21  Upload Custom Logo', 2),
        ('TC-22  Select Stitching Style', 2),
        ('TC-23  Select Print Method', 2),
        ('TC-24  Select Wash Effect', 2),
        ('TC-25  View Real-Time 2D Design Preview', 2),
        ('TC-26  Calculate Dynamic Price', 2),
        ('TC-27  Send Query to Chatbot', 2),
        ('TC-28  Get AI Design Recommendation (RAG)', 2),
        ('TC-29  View FAQ via Chatbot', 2),
        ('TC-30  View Chat History', 2),
        ('TC-31  Admin Manage FAQ Entries', 2),
        ('TC-32  Create New Order', 2),
        ('TC-33  Search Order History', 2),
        ('TC-34  View Order History', 2),
        ('TC-35  Cancel Order', 2),
        ('TC-36  Track Order Status', 2),
        ('TC-37  Reorder Previous Order', 2),
        ('TC-38  Record Online Payment', 2),
        ('TC-39  Record Cash on Delivery', 2),
        ('TC-40  Generate Digital Receipt', 2),
        ('TC-41  View All Payments and Transactions', 2),
        ('TC-42  Download Invoice', 2),
        ('TC-43  Manage Products — Admin', 2),
        ('TC-44  Manage Customer Accounts — Admin', 2),
        ('TC-45  View and Update Order Management — Admin', 2),
        ('TC-46  View Inventory and Stock Alerts — Admin', 2),
        ('TC-47  Generate Sales Report', 2),
        ('TC-48  Generate Order Report', 2),
        ('TC-49  View Analytics Dashboard', 2),
        ('TC-50  Send Notification to Customers', 2),
        ('TC-51  Manage Content', 2),
        ('TC-52  Manage Coupon/Promo Codes', 2),
        ('TC-53  Export Data as CSV', 2),
        ('TC-54  Manage Admin Users', 2),
        ('TC-55  View Customer Reviews', 2),
        ('TC-56  Generate Performance Report', 2),

        # ── CHAPTER 6 ─────────────────────────────────────────────────────────
        ('CHAPTER # 6: USER MANUAL', 0),
        ('User Roles in this Application', 1),
        ('6.1  Screenshots', 1),
        ('6.1.1  Landing / Home Page', 2),
        ('6.1.2  Registration Screen', 2),
        ('6.1.3  Sign In Screen', 2),
        ('6.1.4  Forgot Password Screen', 2),
        ('6.1.5  OTP Verification Screen', 2),
        ('6.1.6  Change Password Screen', 2),
        ('6.1.7  Product Catalog Screen', 2),
        ('6.1.8  Product Detail Screen', 2),
        ('6.1.9  Product Customization Screen (Designer)', 2),
        ('6.1.10  Real-Time 2D Preview Panel', 2),
        ('6.1.11  Logo Upload Section', 2),
        ('6.1.12  AI Chatbot Widget', 2),
        ('6.1.13  Shopping Cart Screen', 2),
        ('6.1.14  Checkout Screen', 2),
        ('6.1.15  Order Confirmation Screen', 2),
        ('6.1.16  Order History Screen', 2),
        ('6.1.17  Order Detail Screen', 2),
        ('6.1.18  Order Tracking Screen', 2),
        ('6.1.19  Payment History Screen', 2),
        ('6.1.20  Customer Profile Screen', 2),
        ('6.1.21  Admin Login Screen', 2),
        ('6.1.22  Admin Dashboard', 2),
        ('6.1.23  Customer Management Screen', 2),
        ('6.1.24  Product Management Screen', 2),
        ('6.1.25  Order Management Screen (Admin)', 2),
        ('6.1.26  Payment Management Screen', 2),
        ('6.1.27  Inventory Management Screen', 2),
        ('6.1.28  Reports Screen', 2),
        ('6.1.29  Coupon Management Screen', 2),
        ('6.1.30  Content Management Screen', 2),
    ]

    for text, level in toc_entries:
        add_toc_entry(doc, text, level)

def set_page_number_roman(section):
    """Set page numbers to Roman numerals for a section."""
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'lowerRoman')
    pgNumType.set(qn('w:start'), '1')
    sectPr = section._sectPr
    existing = sectPr.find(qn('w:pgNumType'))
    if existing is not None:
        sectPr.remove(existing)
    sectPr.append(pgNumType)

def set_page_number_arabic(section, start=1):
    """Set page numbers to Arabic numerals for a section."""
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), str(start))
    sectPr = section._sectPr
    existing = sectPr.find(qn('w:pgNumType'))
    if existing is not None:
        sectPr.remove(existing)
    sectPr.append(pgNumType)

def add_page_number_footer(doc):
    """Add centered page number to the footer of the current section."""
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

def add_section_break(doc, break_type='nextPage'):
    """Add a section break (nextPage) to start a new section."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '12240')
    pgSz.set(qn('w:h'), '15840')
    sectPr.append(pgSz)
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1440')
    pgMar.set(qn('w:right'), '1440')
    pgMar.set(qn('w:bottom'), '1440')
    pgMar.set(qn('w:left'), '1800')
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)
    pPr.append(sectPr)
    return para

def styled_para(doc, text, style_name='Normal', bold=False, italic=False,
                align=WD_ALIGN_PARAGRAPH.LEFT, size=11, space_before=0,
                space_after=6, font_name='Times New Roman'):
    para = doc.add_paragraph(style=style_name)
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size)
    return para

def heading_para(doc, text, level=1):
    """Add heading: H1=14pt, H2=12pt, H3=11pt — all bold, TNR."""
    para = doc.add_paragraph(style=f'Heading {level}')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(14)   # Major heading
    elif level == 2:
        run.font.size = Pt(12)   # Sub heading
    else:
        run.font.size = Pt(11)   # Sub-sub heading
    return para

def table_caption(doc, text):
    """Add a table caption BEFORE the table (bold, 11pt, centred)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return para

def figure_caption(doc, text):
    """Add a figure caption AFTER the figure (italic, 11pt, centred)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return para

def chapter_divider(doc, chapter_num, title):
    """Add a full-page chapter divider like the FUUAST template."""
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(f'CHAPTER # {chapter_num}:')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = para2.add_run(title)
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(18)
    doc.add_page_break()

def add_simple_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with grey header row."""
    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], 'C0C0C0')
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    return table

def add_uc_table(doc, uc):
    """Add a use-case table following FUUAST template style."""
    fields = ['Use Case ID','Use Case Name','Actor(s)','Pre-Conditions',
              'Post-Conditions','Main Flow','Alternative Flow','Exception Flow']
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, f in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = f
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        set_cell_bg(row.cells[0], 'E0E0E0')
        val = uc.get(f, '')
        row.cells[1].text = val
        row.cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    row0 = table.rows[0]
    row0.cells[0].width = Inches(1.8)
    row0.cells[1].width = Inches(4.7)
    return table

def add_tc_table(doc, tc):
    """Add a test-case table following FUUAST template style."""
    fields = ['Test Case ID','Test Case Name','Description','Pre-Conditions',
              'Test Steps','Test Data','Expected Result','Actual Result','Status']
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, f in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = f
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        set_cell_bg(row.cells[0], 'E0E0E0')
        val = tc.get(f, '')
        row.cells[1].text = val
        row.cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    row0 = table.rows[0]
    row0.cells[0].width = Inches(1.8)
    row0.cells[1].width = Inches(4.7)
    return table

def add_ssd_table(doc, ssd_rows, title):
    """Add a System Sequence Diagram as Actor | Direction | System table."""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = table.rows[0].cells
    for i, hdr in enumerate(['Actor', 'Direction', 'System']):
        h[i].text = hdr
        h[i].paragraphs[0].runs[0].bold = True
        h[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        h[i].paragraphs[0].runs[0].font.size = Pt(11)
        h[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(h[i], 'C0C0C0')
    for row_data in ssd_rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data[:3]):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
            if i == 1:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(0.8)
    table.columns[2].width = Inches(3.2)
    return table

def set_margins(section):
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)

# ─────────────────────────────────────────────────────────────────────────────
# READ EXISTING DOCUMENT DATA
# ─────────────────────────────────────────────────────────────────────────────
print("Reading source document...")
src = Document('AURA_WEAR_FYP_Document.docx')

# Extract UC data (tables 5–60, 56 UCs)
uc_data = []
for table in src.tables[5:61]:
    uc = {}
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2:
            uc[cells[0]] = cells[1]
    uc_data.append(uc)

# Extract Domain Model (table 61)
domain_table = src.tables[61]
domain_rows = []
for row in domain_table.rows:
    domain_rows.append([c.text.strip() for c in row.cells])

# Extract SSD data (tables 62–117, 56 SSDs)
ssd_data = []
for table in src.tables[62:118]:
    rows = []
    for row in table.rows:
        rows.append([c.text.strip() for c in row.cells])
    ssd_data.append(rows)

# SSD titles from document paragraphs
ssd_titles = [
    'SSD:1 Process SignUp', 'SSD:2 Process SignIn', 'SSD:3 Forgot Password',
    'SSD:4 Reset Password', 'SSD:5 Update Profile', 'SSD:6 Upload Profile Picture',
    'SSD:7 View Products', 'SSD:8 Search Products', 'SSD:9 Customize Product',
    'SSD:10 Upload Logo', 'SSD:11 Generate AI Design Suggestion',
    'SSD:12 View 2D Real-Time Preview', 'SSD:13 Save Customization to Cart',
    'SSD:14 View Cart', 'SSD:15 Update Cart', 'SSD:16 Apply Coupon Code',
    'SSD:17 Interact with AI Chatbot', 'SSD:18 Get Product Recommendation',
    'SSD:19 Get Customization Assistance', 'SSD:20 Get Pricing Information',
    'SSD:21 Submit Review', 'SSD:22 View Reviews', 'SSD:23 Manage Wishlist',
    'SSD:24 Add Shipping Address', 'SSD:25 Select Shipping Address at Checkout',
    'SSD:26 Receive Order Confirmation Email', 'SSD:27 Subscribe to Newsletter',
    'SSD:28 Contact Support', 'SSD:29 View FAQ via Chatbot',
    'SSD:30 View Chat History', 'SSD:31 Admin Manage FAQ Entries',
    'SSD:32 Create New Order', 'SSD:33 Search Order History',
    'SSD:34 View Order History', 'SSD:35 Cancel Order', 'SSD:36 Track Order Status',
    'SSD:37 Reorder Previous Order', 'SSD:38 Record Online Payment',
    'SSD:39 Record Cash on Delivery', 'SSD:40 Generate Digital Receipt',
    'SSD:41 View All Payments and Transactions', 'SSD:42 Download Invoice',
    'SSD:43 Manage Products - Admin', 'SSD:44 Manage Customer Accounts - Admin',
    'SSD:45 View and Update Order Management - Admin',
    'SSD:46 View Inventory and Stock Alerts - Admin',
    'SSD:47 Generate Sales Report', 'SSD:48 Generate Order Report',
    'SSD:49 View Analytics Dashboard', 'SSD:50 Send Notification to Customers',
    'SSD:51 Manage Content - Terms, Privacy, FAQ',
    'SSD:52 Manage Coupon/Promo Codes', 'SSD:53 Export Data as CSV',
    'SSD:54 Manage Admin Users', 'SSD:55 View Customer Reviews',
    'SSD:56 Generate Performance Report',
]

# Extract Schema data (tables 118–126)
schema_tables = []
schema_names = ['users','products','orders','order_items','payments',
                'chat_logs','inventory','coupons','faq']
for i, table in enumerate(src.tables[118:127]):
    rows = []
    for row in table.rows:
        rows.append([c.text.strip() for c in row.cells])
    schema_tables.append(rows)

# Extract Class diagram table (table 127)
class_table_rows = []
for row in src.tables[127].rows:
    class_table_rows.append([c.text.strip() for c in row.cells])

# Extract Sample queries table (table 128)
query_rows = []
for row in src.tables[128].rows:
    query_rows.append([c.text.strip() for c in row.cells])

# Extract TC data (tables 129–184, 56 TCs)
tc_data = []
for table in src.tables[129:185]:
    tc = {}
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2:
            tc[cells[0]] = cells[1]
    tc_data.append(tc)

print(f"  Loaded {len(uc_data)} use cases, {len(ssd_data)} SSDs, {len(tc_data)} test cases")

# ─────────────────────────────────────────────────────────────────────────────
# CREATE NEW DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
print("Building document...")
doc = Document()

# Set default styles — 11pt TNR, 1.5 line spacing, 6pt after
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.space_after = Pt(6)

# Heading styles: H1=14pt, H2=12pt, H3=11pt — all bold, black, TNR
for level, size in [(1, 14), (2, 12), (3, 11)]:
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Times New Roman'
    h_style.font.size = Pt(size)
    h_style.font.bold = True
    h_style.font.color.rgb = RGBColor(0, 0, 0)
    h_style.paragraph_format.space_before = Pt(12)
    h_style.paragraph_format.space_after = Pt(6)
    h_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# List bullet style — 11pt TNR
try:
    lb = doc.styles['List Bullet']
    lb.font.name = 'Times New Roman'
    lb.font.size = Pt(11)
    lb.paragraph_format.space_after = Pt(6)
    lb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
except Exception:
    pass

# Set initial margins
set_margins(doc.sections[0])

# ─────────────────────────────────────────────────────────────────────────────
# PAGE i – COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
print("  Adding cover page...")

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AURA-WEAR: AI-Integrated Custom Clothing Website')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(16)

doc.add_paragraph()

# FUUAST Logo (cover page)
logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_path = 'doc_parts/diagrams/fuuast_logo.png'
if os.path.exists(logo_path):
    run_logo = logo_para.add_run()
    run_logo.add_picture(logo_path, width=Inches(2.2))
else:
    r2 = logo_para.add_run('[FUUAST LOGO]')
    r2.bold = True
    r2.font.size = Pt(14)

for _ in range(3):
    doc.add_paragraph()

for text in [
    'Developed By:',
    'Muhammad Umer Javed (B.S CS Aut-21-Eve-0170)',
    'Abdul Baqi (B.S CS Aut-2022-E-B-0260)',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BS (Computer Science)')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(13)

doc.add_paragraph()
for text in ['Supervised by:', 'Dr. Zaheen Khan']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
for text in [
    'Department of Computer Science',
    'Federal Urdu University of Arts, Science and Technology',
    'Islamabad',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Session [2022-2026]')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(13)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE ii – TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AURA-WEAR: AI-Integrated Custom Clothing Website')
r.bold = True
r.underline = True
r.font.name = 'Times New Roman'
r.font.size = Pt(14)

doc.add_paragraph()
submission_text = (
    'A project submitted to the Department of Computer Science as partial '
    'fulfillment of the requirement for the award of Degree of Bachelor of '
    'Science in Computer Science.'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
r = p.add_run(submission_text)
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

doc.add_paragraph()

table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Name'
hdr[1].text = 'Registration Number'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell.paragraphs[0].runs[0].font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, 'C0C0C0')

data = [
    ('Muhammad Umer Javed', 'Aut-21-Eve-0170'),
    ('Abdul Baqi', 'Aut-2022-E-B-0260'),
]
for i, (name, reg) in enumerate(data):
    row = table.rows[i + 1]
    row.cells[0].text = name
    row.cells[1].text = reg
    for cell in row.cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

table.columns[0].width = Inches(2.8)
table.columns[1].width = Inches(2.8)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Supervisor')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

for text in ['Dr. Zaheen Khan', 'Assistant Professor',
             'Department of Computer Science',
             'Federal Urdu University of Arts, Science and Technology',
             'Islamabad']:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE iii – FINAL APPROVAL
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Final Approval')
r.bold = True
r.underline = True
r.font.name = 'Times New Roman'
r.font.size = Pt(14)

doc.add_paragraph()
cert_text = (
    'It is certified that I have read the project report submitted by '
    '​Muhammad Umer Javed​ (B.S CS Aut-21-Eve-0170) and '
    '​Abdul Baqi​ (B.S CS Aut-2022-E-B-0260) and it is our '
    'judgment that this project is of sufficient standard to warrant its '
    'acceptance by the FUUAST university, Islamabad for the Degree of '
    'Bachelor of Science in Computer Science.'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
run = p.add_run('It is certified that I have read the project report submitted by ')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)
run2 = p.add_run('Muhammad Umer Javed')
run2.bold = True; run2.font.name = 'Times New Roman'; run2.font.size = Pt(12)
run3 = p.add_run(' (B.S CS Aut-21-Eve-0170) and ')
run3.font.name = 'Times New Roman'; run3.font.size = Pt(12)
run4 = p.add_run('Abdul Baqi')
run4.bold = True; run4.font.name = 'Times New Roman'; run4.font.size = Pt(12)
run5 = p.add_run(
    ' (B.S CS Aut-2022-E-B-0260) and it is our judgment that this project '
    'is of sufficient standard to warrant its acceptance by the FUUAST '
    'university, Islamabad for the Degree of Bachelor of Science in '
    'Computer Science.'
)
run5.font.name = 'Times New Roman'; run5.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Committee:')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

for role, name, designation, dept in [
    ('External Examiner', '', '', ''),
    ('Internal Examiner', 'Mr. Khawaja Tahir Mahmood', 'Lecturer',
     'Department of Computer Science\nFederal Urdu University of Arts, Science and Technology\nIslamabad'),
    ('Supervisor', 'Dr. Zaheen Khan', 'Assistant Professor',
     'Department of Computer Science\nFederal Urdu University of Arts, Science and Technology\nIslamabad'),
]:
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Inches(2.2)
    tbl.columns[1].width = Inches(4.0)
    row = tbl.rows[0]
    row.cells[0].text = role
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    row.cells[0].paragraphs[0].runs[0].underline = False

    sig_cell = row.cells[1]
    sig_para = sig_cell.paragraphs[0]
    sig_para.add_run('_______________________________').font.name = 'Times New Roman'
    if name:
        sig_cell.add_paragraph().add_run(name).font.name = 'Times New Roman'
        sig_cell.paragraphs[-1].runs[0].font.size = Pt(12)
    if designation:
        sig_cell.add_paragraph().add_run(designation).font.name = 'Times New Roman'
        sig_cell.paragraphs[-1].runs[0].font.size = Pt(12)
    if dept:
        for line in dept.split('\n'):
            sig_cell.add_paragraph().add_run(line).font.name = 'Times New Roman'
            sig_cell.paragraphs[-1].runs[0].font.size = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE iv – DECLARATION
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Declaration')
r.bold = True; r.underline = True
r.font.name = 'Times New Roman'; r.font.size = Pt(14)

doc.add_paragraph()
decl_text = (
    'We hereby declare that this software, neither as a whole nor as a '
    'part has been copied out from any source. It is further declared that '
    'we have developed this software and the accompanied report entirely on '
    'the basis of our personal efforts. If any part of this project is '
    'proved to be copied out from any source or found to be reproduction of '
    'some other, we will stand by the consequences. No portion of the work '
    'presented has been submitted in support of any application for any '
    'other degree or qualification of this or any other university or '
    'institute of learning.'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
r = p.add_run(decl_text)
r.font.name = 'Times New Roman'; r.font.size = Pt(12)

for _ in range(4):
    doc.add_paragraph()

for name in ['Muhammad Umer Javed', 'Abdul Baqi']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE v – DEDICATION
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

for _ in range(6):
    doc.add_paragraph()

for text in ['Dedicated to our beloved parents,', 'Teachers', 'And', 'The fellows']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(20)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE vi – ACKNOWLEDGEMENT
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Acknowledgement')
r.bold = True; r.underline = True
r.font.name = 'Times New Roman'; r.font.size = Pt(14)

doc.add_paragraph()
ack_text = (
    'Thanks to Almighty Allah for giving us knowledge, power and strength '
    'to accomplish this task. We learned a great deal while developing this '
    'project and the experience will certainly benefit us in our professional '
    'careers. Many of our teachers guided us during this project but we are '
    'especially grateful to our supervisor Dr. Zaheen Khan for her invaluable '
    'guidance and continuous support throughout all phases of this work. Her '
    'encouragement motivated us during difficult times. We would also like to '
    'express our heartfelt gratitude to all our friends and family for their '
    'support and encouragement throughout this journey.'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
r = p.add_run(ack_text)
r.font.name = 'Times New Roman'; r.font.size = Pt(12)

for _ in range(4):
    doc.add_paragraph()

for name in ['Muhammad Umer Javed', 'Abdul Baqi']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE vii – PROJECT IN BRIEF
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Project in Brief')
r.bold = True; r.underline = True
r.font.name = 'Times New Roman'; r.font.size = Pt(14)

doc.add_paragraph()

brief_data = [
    ('Project Title', 'AURA-WEAR: AI-Integrated Custom Clothing Website'),
    ('Organization', 'Federal Urdu University of Arts, Science and Technology'),
    ('Objectives',
     '1. To develop an AI-powered web platform for custom streetwear and sportswear design and ordering.\n'
     '2. To implement a real-time 2D Canvas-based customization engine with dynamic pricing.\n'
     '3. To integrate a RAG-based AI chatbot (n8n + OpenAI GPT-4o-mini) for personalized recommendations.\n'
     '4. To provide a comprehensive admin panel with analytics, order management, and reporting.'),
    ('Undertaken By', 'Muhammad Umer Javed\nAbdul Baqi'),
    ('Supervised By', 'Dr. Zaheen Khan'),
    ('Date Started', 'Feb 16, 2026'),
    ('Date Completed', 'Jun 26, 2026'),
    ('Technologies Used',
     'Frontend: HTML5, CSS3, Vanilla JavaScript\n'
     'Backend: FastAPI (Python 3.11+)\n'
     'Database: SQLite (development) / PostgreSQL (production)\n'
     'AI/Chatbot: n8n Workflow + OpenAI GPT-4o-mini (RAG)\n'
     'Deployment: Railway.app\n'
     'Authentication: JWT + bcrypt\n'
     'Image Storage: Cloudinary'),
    ('System Used',
     '8th Gen Core i5 Processor @ 2.40 GHz, 16 GB RAM, 512 GB SSD'),
]

tbl = doc.add_table(rows=len(brief_data), cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (key, val) in enumerate(brief_data):
    row = tbl.rows[i]
    row.cells[0].text = key
    row.cells[0].paragraphs[0].runs[0].bold = False
    row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    row.cells[1].text = val
    row.cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(12)

tbl.columns[0].width = Inches(1.6)
tbl.columns[1].width = Inches(4.9)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE viii – ABSTRACT
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ABSTRACT')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(16)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AURA-WEAR: AI-Integrated Custom Clothing Website')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(13)

doc.add_paragraph()
abstract_text = (
    'The fashion e-commerce industry in Pakistan faces a significant gap: '
    'existing platforms rarely offer personalized clothing customization with '
    'real-time design preview capabilities and AI-based recommendations. '
    'AURA-WEAR is an AI-integrated web platform designed to bridge this gap '
    'by enabling customers to design and order custom streetwear and sportswear '
    'through an intelligent, interactive interface.\n\n'
    'The platform is built on a 4-tier layered architecture comprising a '
    'Vanilla JavaScript frontend, a FastAPI (Python) backend, an n8n-powered '
    'AI processing layer, and a relational database. Customers can use the '
    'real-time 2D Canvas-based customization engine to select fabric type, '
    'garment color, stitching style, print method, wash effect, and upload '
    'custom logos, with dynamic pricing computed instantly as selections are '
    'made. An AI chatbot powered by Retrieval-Augmented Generation (RAG) via '
    'n8n workflows and OpenAI GPT-4o-mini provides personalized product '
    'recommendations and FAQ assistance.\n\n'
    'The system supports two roles: Customer and Admin. Customers can register, '
    'customize products, place orders, track order status through a five-stage '
    'lifecycle, make payments (Cash on Delivery and PayFast), and download '
    'invoices. Admin users have access to a comprehensive dashboard with '
    'real-time analytics, full CRUD operations on products, order and customer '
    'management, inventory alerts, sales reporting, coupon/promo code '
    'management, and CSV data export.\n\n'
    'AURA-WEAR implements industry-standard security practices including bcrypt '
    'password hashing, JWT-based authentication with role-based access control, '
    'OTP-based password reset via email, and HTTPS/TLS communication. The '
    'platform was tested against 56 functional test cases, all of which passed '
    'successfully, demonstrating a reliable, scalable, and user-friendly '
    'solution for the custom clothing industry in Pakistan.'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
r = p.add_run(abstract_text)
r.font.name = 'Times New Roman'; r.font.size = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TABLE OF CONTENTS')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(16)

doc.add_paragraph()
add_manual_toc(doc)
doc.add_paragraph()

# Add footer with page numbers for front matter
add_page_number_footer(doc)

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 1 – PROPOSAL
# ─────────────────────────────────────────────────────────────────────────────
print("  Chapter 1: Proposal...")
chapter_divider(doc, 1, 'PROPOSAL')

heading_para(doc, 'Revision History', 2)

rev_rows = [
    ('v1.0', 'Nov 2025', 'Initial SRS document: core requirements, basic use cases, project scope, and initial technology selection.', 'Muhammad Umer Javed'),
    ('v2.0', 'Jan 2026', 'Added use case diagrams, ER diagram, system architecture, and expanded functional requirements.', 'Abdul Baqi'),
    ('v3.0', 'Mar 2026', 'Tech stack upgraded to FastAPI + Python, JWT authentication added, AI chatbot design finalized.', 'Muhammad Umer Javed'),
    ('v4.0', 'Jun 2026', 'Final version: all 56 functional requirements implemented, complete test cases, user manual added.', 'Both Authors'),
]
table_caption(doc, 'Table 1: Revision History')
add_simple_table(doc, ['Version', 'Date', 'Description', 'Author'], rev_rows,
                 col_widths=[0.7, 0.8, 3.6, 1.4])
doc.add_paragraph()

heading_para(doc, '1.1  Introduction', 2)
styled_para(doc,
    'The fashion industry is rapidly shifting toward personalized and custom-made clothing, '
    'driven by rising consumer expectations for uniqueness and individual expression. In '
    'Pakistan, this shift presents a significant untapped market opportunity, as most '
    'existing e-commerce platforms offer only standardized, off-the-shelf garments without '
    'meaningful customization capabilities.\n\n'
    'AURA-WEAR is an AI-integrated, web-based custom clothing platform designed to address '
    'this gap. The system enables customers to design their own streetwear and sportswear '
    'through a real-time 2D Canvas-based design engine, supported by an AI chatbot that '
    'provides intelligent fabric, style, and design recommendations. The platform is built '
    'on a modern, scalable 4-tier architecture using FastAPI (Python) as the backend, '
    'Vanilla JavaScript for the frontend, n8n with OpenAI GPT-4o-mini for the AI layer, '
    'and SQLite/PostgreSQL for data persistence.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '1.2  Problem Statement', 2)
styled_para(doc,
    'In Pakistan, there is a significant gap in online platforms offering custom clothing '
    'with real-time design preview capabilities and AI-based recommendations. Existing '
    'platforms either provide basic template customization without live previews, or offer '
    'design tools without intelligent assistance. Customers who wish to order uniquely '
    'designed garments face challenges in visualizing their final product before committing '
    'to a purchase, leading to dissatisfaction and high return rates. Additionally, clothing '
    'businesses lack automated systems to manage custom orders, track production, and '
    'provide customers with transparent order status updates. AURA-WEAR is designed to solve '
    'these problems comprehensively.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '1.3  Proposed System', 2)
styled_para(doc,
    'AURA-WEAR proposes a comprehensive web-based solution for custom clothing design and '
    'ordering. The system is organized into three primary portals:\n\n'
    '1. Customer Portal: Allows customers to browse the product catalog, use the AI-powered '
    '2D customization engine, interact with the AI chatbot, place orders, track order status, '
    'manage payments, and view order history.\n\n'
    '2. Admin Panel: Provides administrators with full CRUD capabilities for product '
    'management, customer account management, order processing, inventory monitoring, '
    'sales reporting, coupon management, and analytics visualization.\n\n'
    '3. AI Processing Layer: An n8n workflow automation platform integrates with OpenAI '
    'GPT-4o-mini to implement a Retrieval-Augmented Generation (RAG) pipeline that answers '
    'product queries and provides personalized design recommendations.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '1.4  Benefits of the Proposed System', 2)
benefits = [
    'Real-Time 2D Design Preview: Customers can visualize their custom garment configurations instantly before placing an order, reducing uncertainty and return rates.',
    'AI-Powered Recommendations: The RAG-based chatbot provides intelligent fabric, style, and design suggestions tailored to each customer\'s preferences and queries.',
    'Transparent Dynamic Pricing: Pricing is calculated in real time as customers add customization options, building trust and eliminating hidden costs.',
    'Streamlined Ordering Process: A complete e-commerce workflow — from design to checkout to order tracking — ensures a smooth customer journey.',
    'Admin Analytics Dashboard: Administrators gain insights into sales trends, inventory levels, and customer behavior through a built-in analytics panel.',
    'Scalable Architecture: The 4-tier layered architecture ensures the system can scale to handle growing user demand without architectural overhaul.',
    'Secure Platform: JWT authentication, bcrypt password hashing, and OTP verification ensure data security and user account protection.',
]
for b in benefits:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(b)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

heading_para(doc, '1.5  Scope', 2)
styled_para(doc,
    'AURA-WEAR is scoped as a web-based platform accessible via modern browsers on desktop '
    'and mobile devices. The system covers custom design and ordering for four garment '
    'categories: T-Shirts, Hoodies, Jackets, and Sports Uniforms. The customization options '
    'include fabric type selection, color selection, logo upload, stitching style, print '
    'method, and wash effect.\n\n'
    'The system supports two user roles (Customer and Admin) with distinct access controls. '
    'Payment methods include Cash on Delivery (COD) and PayFast gateway integration. Order '
    'tracking covers a five-stage lifecycle: Pending, In Production, Quality Check, Shipped, '
    'and Delivered.\n\n'
    'Out of scope for this version: a native mobile application (iOS/Android), integration '
    'with a physical store inventory system, Stripe payment gateway (deferred due to regional '
    'API restrictions), and support for languages other than English.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '1.6  Survey Analysis', 2)
styled_para(doc,
    'A comparative analysis was conducted of five existing e-commerce and custom clothing '
    'platforms to identify gaps that AURA-WEAR addresses. The systems evaluated were Amazon '
    'Custom, Printful, Custom Ink, Spreadshirt, and AURA-WEAR.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '1.6.1  Survey Analysis Table', 3)
survey_rows = [
    ('Amazon Custom', 'No', 'No', 'Limited', 'No', 'Yes'),
    ('Printful', 'No', 'Basic', 'Yes', 'Yes', 'Yes'),
    ('Custom Ink', 'No', '2D Basic', 'Limited', 'Yes', 'Yes'),
    ('Spreadshirt', 'No', 'Basic', 'Limited', 'Yes', 'Yes'),
    ('AURA-WEAR', 'Yes (RAG)', 'Yes (Canvas)', 'Yes (Full)', 'Yes (Dynamic)', 'Yes'),
]
table_caption(doc, 'Table 2: Comparative Survey of Custom Clothing Platforms')
add_simple_table(doc, ['System Name','AI Chatbot','2D Preview','Custom Fabric','Real-Time Pricing','Order Tracking'],
                 survey_rows, col_widths=[1.3, 1.1, 1.1, 1.1, 1.3, 1.1])
doc.add_paragraph()
styled_para(doc,
    'The survey confirms that AURA-WEAR is the only platform among those evaluated to offer '
    'all five features simultaneously, making it a uniquely comprehensive solution for the '
    'Pakistani custom clothing market.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '1.7  Modules & Submodules', 2)
styled_para(doc,
    'AURA-WEAR is divided into eight functional modules, each responsible for a distinct '
    'area of the system\'s functionality:',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

modules = [
    ('1.7.1  Security Management',
     'Handles user authentication and authorization including registration, login, logout, '
     'JWT token management, and bcrypt password hashing. Ensures secure access to all '
     'protected endpoints via role-based access control.'),
    ('1.7.2  User Profile Management',
     'Allows customers to create, view, and update their profile information including '
     'personal details, shipping addresses, and profile pictures stored on Cloudinary. '
     'Admins can view and manage customer account statuses.'),
    ('1.7.3  Product Catalog Management',
     'Manages the product listings for streetwear and sportswear categories. Admins can '
     'add, update, delete, and toggle product availability. Customers can browse, search, '
     'and filter products.'),
    ('1.7.4  Product Customization',
     'The core module of AURA-WEAR. Provides the 2D Canvas-based design interface where '
     'customers select garment type, color, fabric, size, add custom text and logo, choose '
     'stitching style, print method, and wash effect with real-time dynamic pricing.'),
    ('1.7.5  AI Chatbot Management',
     'Integrates the RAG-based AI chatbot powered by n8n workflows and OpenAI GPT-4o-mini. '
     'Provides product-related answers, design recommendations, and FAQ responses. Chat '
     'history is persisted for authenticated users.'),
    ('1.7.6  Order Management',
     'Handles the complete order lifecycle: cart management, order placement, status updates '
     '(Pending, In Production, Quality Check, Shipped, Delivered), order cancellation, and '
     'reorder functionality.'),
    ('1.7.7  Payment Management',
     'Manages payment processing supporting Cash on Delivery (COD) and PayFast gateway. '
     'Records payment status, generates digital receipts, provides payment history, and '
     'enables PDF invoice download.'),
    ('1.7.8  Admin Dashboard Management',
     'Provides administrators with a centralized analytics dashboard showing key metrics: '
     'total orders, revenue, active customers, inventory levels, and recent activity. '
     'Supports CSV export, coupon management, and bulk customer notifications.'),
]

for title, desc in modules:
    heading_para(doc, title, 3)
    styled_para(doc, desc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

heading_para(doc, '1.8  Primary Actor', 2)
styled_para(doc,
    'AURA-WEAR has two primary actors who interact with the system:\n\n'
    'Customer: The end-user of the platform. A Customer can register an account, log in, '
    'browse the product catalog, use the 2D customization engine to design custom garments, '
    'interact with the AI chatbot, add items to the cart, place orders, track order status, '
    'manage payments, and view their profile and order history.\n\n'
    'Admin: The system administrator. An Admin can log in to the admin panel, manage product '
    'listings, manage customer accounts, process and update orders, monitor inventory, '
    'generate sales and performance reports, manage coupon codes, send notifications, '
    'manage content (Terms, Privacy, FAQ), and export data as CSV.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '1.9  Tools & Technologies', 2)

tech_sections = [
    ('1.9.1  Front-End Tools', 'HTML5, CSS3, Vanilla JavaScript (ES6+), HTML5 Canvas API for 2D customization preview, Fetch API for REST communication.'),
    ('1.9.2  Back-End Tools', 'FastAPI (Python 3.11+): High-performance asynchronous REST API framework. SQLAlchemy ORM for database interaction. Passlib/bcrypt for password hashing. python-jose for JWT token generation and validation. Gmail SMTP via smtplib for OTP email delivery.'),
    ('1.9.3  AI / Automation Tools', 'n8n (v1.x): Open-source workflow automation platform used to build the RAG pipeline. OpenAI GPT-4o-mini: Large language model providing natural language understanding and generation. ChromaDB: Vector store for semantic similarity search in RAG pipeline.'),
    ('1.9.4  Database', 'SQLite: Lightweight relational database used during development. PostgreSQL: Production-grade relational database for the deployed application.'),
    ('1.9.5  Deployment', 'Railway.app: Cloud hosting platform for the FastAPI backend. Git and GitHub for version control. Cloudinary for cloud-based image storage and management.'),
]
for title, desc in tech_sections:
    heading_para(doc, title, 3)
    styled_para(doc, desc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

heading_para(doc, '1.10  System Design Approach', 2)
styled_para(doc,
    'AURA-WEAR follows a 4-Tier Layered Architecture that separates concerns across '
    'distinct layers:\n\n'
    'Tier 1 — Presentation Layer (Browser): HTML5/CSS3/Vanilla JS frontend. Handles all '
    'user interactions, renders the 2D Canvas customization preview, and communicates with '
    'the backend via REST API calls.\n\n'
    'Tier 2 — Application Layer (FastAPI Backend): Core business logic layer. Receives HTTP '
    'requests, authenticates users via JWT, processes data through SQLAlchemy ORM, and '
    'coordinates with external services.\n\n'
    'Tier 3 — AI Processing Layer (n8n + OpenAI): Intelligent service layer. Processes '
    'chatbot queries through the RAG pipeline — retrieving relevant product documents from '
    'ChromaDB and generating contextualized responses via GPT-4o-mini.\n\n'
    'Tier 4 — Data Layer (SQLite / PostgreSQL): Persistence layer. Stores all structured '
    'data including users, products, orders, payments, customizations, chat logs, inventory, '
    'coupons, and FAQs.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

# Activity diagram (Customer Order Flow) — image2
img_path = 'doc_parts/diagrams/image2.png'
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(3.5))
    figure_caption(doc, 'Figure 1: Activity Diagram — Customer Order Flow')

doc.add_paragraph()

# System Architecture Diagram — image6
img_path = 'doc_parts/diagrams/image6.png'
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(5.5))
    figure_caption(doc, 'Figure 2: System Architecture Diagram — 4-Tier Layered Architecture')

heading_para(doc, '1.11  Process Model Used', 2)
styled_para(doc,
    'AURA-WEAR was developed using the Agile Incremental process model. This approach '
    'allowed the team to deliver working software in short, iterative sprints while '
    'continuously incorporating feedback and refining requirements.\n\n'
    'Sprint 1 (Weeks 1–3) — Foundation: Project setup, technology stack finalization, '
    'database schema design, FastAPI project structure, basic authentication '
    '(register/login/JWT), and initial product catalog API.\n\n'
    'Sprint 2 (Weeks 4–6) — Core Features: 2D Canvas customization engine, dynamic pricing '
    'logic, product catalog frontend, shopping cart implementation, and order placement '
    'workflow.\n\n'
    'Sprint 3 (Weeks 7–9) — AI & Admin: n8n RAG pipeline integration, AI chatbot frontend, '
    'admin panel (product/order/customer management), analytics dashboard, and payment '
    'integration (COD + PayFast).\n\n'
    'Sprint 4 (Weeks 10–12) — Testing & Deployment: Comprehensive testing (56 test cases), '
    'bug fixing, performance optimization, Railway.app deployment, documentation, and user '
    'manual preparation.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '1.12  Modelling Techniques / Tools Used', 2)
modelling_items = [
    'Use Case Diagrams: To capture functional requirements and actor interactions for all 8 modules.',
    'Fully Dressed Use Cases: Detailed specifications with main success scenarios and alternative/exception flows.',
    'Domain Model: To identify key entities and their relationships in the problem domain.',
    'System Sequence Diagrams (SSDs): To illustrate interactions between actors and the system for each use case.',
    'Class Diagrams: To represent the software architecture including classes, attributes, methods, and associations.',
    'Entity-Relationship (ER) Diagram: To design the relational database schema.',
]
for item in modelling_items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(item)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.add_paragraph()
styled_para(doc, 'Modelling and development tools used:', align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
tools = [
    'draw.io: UML diagrams (use case, class, ER, SSDs, domain model).',
    'Visual Studio Code: Primary code editor for Python and JavaScript development.',
    'Postman: REST API testing and documentation.',
    'GitHub: Version control and remote repository hosting.',
    'Railway.app Dashboard: Deployment monitoring and environment variable management.',
    'n8n Canvas: Visual workflow builder for the RAG AI pipeline.',
]
for t in tools:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(t)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

heading_para(doc, '1.13  Limitation / Constraint', 2)
limitations = [
    'Web Browser Only: AURA-WEAR is a web application and does not have a dedicated native mobile application for iOS or Android.',
    'Payment Methods: Only Cash on Delivery (COD) and PayFast are supported. Stripe integration was deferred due to regional API restrictions.',
    'AI Chatbot Requires Internet: The chatbot depends on the n8n server and OpenAI API; offline operation is not supported.',
    'No Physical Store Integration: AURA-WEAR operates as an independent web platform with no integration with any physical retail system.',
    'Single Language: The platform is currently available in English only. Urdu and other regional language support is not included.',
    'No Real-Time Notifications: Push notifications are not implemented in this version; users must check order status manually.',
]
for lim in limitations:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(lim)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

heading_para(doc, '1.14  References', 2)
refs = [
    '[1] P. Loureiro, J. Bettencourt-Silva, and A. Trindade, "Artificial Intelligence in E-Commerce: A Systematic Literature Review," IEEE Access, vol. 10, pp. 12345–12360, 2022.',
    '[2] T. Brown et al., "Language Models are Few-Shot Learners," in Proc. NeurIPS, 2020.',
    '[3] Y. Gao et al., "Retrieval-Augmented Generation for Large Language Models: A Survey," arXiv preprint arXiv:2312.10997, 2023.',
    '[4] S. Kim, J. Lee, and H. Park, "Personalized Fashion Recommendation Using Deep Learning and Visual Similarity," Journal of Fashion Technology & Textile Engineering, vol. 9, no. 3, 2021.',
    '[5] FastAPI Documentation, Sebastián Ramírez, 2024. [Online]. Available: https://fastapi.tiangolo.com/',
    '[6] n8n Documentation, n8n GmbH, 2024. [Online]. Available: https://docs.n8n.io/',
    '[7] OpenAI API Reference, OpenAI, 2024. [Online]. Available: https://platform.openai.com/docs/',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 2 – ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
print("  Chapter 2: Analysis...")
chapter_divider(doc, 2, 'ANALYSIS')

heading_para(doc, '2.1  Introduction', 2)

heading_para(doc, '2.1.1  Purpose', 3)
styled_para(doc,
    'This document defines the functional and non-functional requirements for the AURA-WEAR '
    'AI-Integrated Custom Clothing Website system. It serves as a formal contract between '
    'the development team and project stakeholders, providing a clear and unambiguous '
    'specification of all system features, constraints, and quality attributes.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

heading_para(doc, '2.1.2  Scope', 3)
styled_para(doc,
    'The AURA-WEAR system encompasses three primary components: a customer-facing web '
    'portal, an administrative management panel, and an AI-powered chatbot assistant. The '
    'scope covers all 56 functional requirements organized into eight management modules. '
    'Explicitly out of scope: native mobile applications, physical store integration, and '
    'third-party marketplace API connections.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

heading_para(doc, '2.1.3  Definitions, Acronyms, and Abbreviations', 3)
defn_rows = [
    ('RAG', 'Retrieval-Augmented Generation: AI technique combining information retrieval with language generation.'),
    ('JWT', 'JSON Web Token: compact, URL-safe means of representing claims between parties.'),
    ('OTP', 'One-Time Password: 6-digit numeric code sent to a user\'s registered email for verification.'),
    ('2D Preview', 'Real-time Canvas-based visualization of the garment with applied customizations.'),
    ('COD', 'Cash on Delivery: payment method where the customer pays upon receiving the product.'),
    ('CRUD', 'Create, Read, Update, Delete: basic data operations.'),
    ('API', 'Application Programming Interface: set of protocols for building software applications.'),
    ('SRS', 'Software Requirements Specification: document describing system requirements.'),
    ('ORM', 'Object-Relational Mapper: tool that maps database tables to Python classes.'),
    ('UML', 'Unified Modeling Language: standardized modeling language for software design.'),
    ('DTG', 'Direct to Garment: printing technique that applies ink directly to fabric.'),
    ('HTTPS', 'Hypertext Transfer Protocol Secure: encrypted version of HTTP using TLS.'),
    ('TLS', 'Transport Layer Security: cryptographic protocol for secure internet communications.'),
]
table_caption(doc, 'Table 3: Glossary of Terms and Abbreviations')
add_simple_table(doc, ['Term', 'Definition'], defn_rows, col_widths=[1.3, 5.2])
doc.add_paragraph()

heading_para(doc, '2.1.4  Overview', 3)
styled_para(doc,
    'This chapter presents the complete set of Functional Requirements (Section 2.2) '
    'organized into eight management modules, followed by Non-Functional Requirements '
    '(Section 2.3) covering performance, security, reliability, and other quality '
    'attributes, and External Interface Requirements (Section 2.4) covering user, '
    'hardware, software, and communication interfaces.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '2.2  Functional Requirements', 2)
styled_para(doc,
    'This section enumerates all functional requirements for the AURA-WEAR system. '
    'Requirements are grouped into eight modules corresponding to major system features.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

# FR sections from existing doc
fr_sections = [
    ('2.2.1  Security Management', [
        ('2.2.1.1  Process SignUp', 'The system shall allow new users to register by providing their full name, email address, and password. The system validates the email format using a standard regex pattern and enforces a minimum password strength requirement of 8 characters including at least one numeric digit. The email is checked for uniqueness before the bcrypt-hashed password and user record are saved to the database. A success message is displayed and the user is redirected to the login page.'),
        ('2.2.1.2  Process SignIn', 'Registered users shall authenticate using their email and password. The system retrieves the stored bcrypt hash, verifies the entered password, and generates a JWT token with 30-minute expiry (7 days if "Remember Me" is checked). The token is stored in browser localStorage and the user is redirected to the appropriate dashboard.'),
        ('2.2.1.3  Forgot Password', 'Users can initiate a password reset by entering their registered email. The system generates a cryptographically random 6-digit OTP with a 10-minute expiry and sends it via Gmail SMTP. The system confirms OTP dispatch without revealing whether the email is registered.'),
        ('2.2.1.4  Change Account Password', 'Authenticated users can change their password from profile settings. The system verifies the current password, validates the new password strength, hashes the new password with bcrypt, updates the database, and invalidates all other active sessions.'),
        ('2.2.1.5  JWT Authentication', 'All protected API endpoints require a valid Bearer JWT token in the HTTP Authorization header. The system validates the token signature, checks expiry, and extracts the user ID and role from the payload. Requests without a valid token receive HTTP 401 Unauthorized.'),
        ('2.2.1.6  OTP Verification', 'The system generates and validates 6-digit numeric OTPs for password reset and account verification workflows. Each OTP has a 10-minute expiry and is invalidated after first use. After three failed attempts, the OTP is automatically invalidated.'),
        ('2.2.1.7  Role-Based Access Control', 'The system implements two roles: Customer and Admin. Customers access the public portal; Admins access the management panel. Unauthorized role access returns HTTP 403 Forbidden.'),
        ('2.2.1.8  Session Management', 'JWT tokens are persisted in browser localStorage. Every API call validates the token; upon expiry or inactivity (30 minutes), the system clears localStorage and redirects to the login page.'),
    ]),
    ('2.2.2  User Profile Management', [
        ('2.2.2.1  Add Customer Profile', 'After registration, customers can add supplementary profile information: phone number, full delivery address, city, and profile picture. Profile pictures are uploaded to Cloudinary.'),
        ('2.2.2.2  Update Customer Profile', 'Authenticated customers can modify profile information including display name, phone number, delivery address, city, and profile picture through the Profile Settings page.'),
        ('2.2.2.3  View Customer Profile', 'Customers can view their consolidated profile including personal details, profile picture, total order count, and registration date.'),
        ('2.2.2.4  Manage Customer Account Status', 'Admins can activate or deactivate customer accounts from the admin panel. Deactivated accounts have all active JWT tokens invalidated.'),
    ]),
    ('2.2.3  Product Catalog Management', [
        ('2.2.3.1  Add New Product', 'Admins can add products by specifying name, category (T-Shirt, Hoodie, Jacket, Sports Uniform), description, base price, size options, available colors, and product images (uploaded to Cloudinary).'),
        ('2.2.3.2  Update Product Details', 'Admins can edit any product attribute including name, description, category, base price, size options, and images.'),
        ('2.2.3.3  Enable or Disable Product', 'Admins can toggle product availability. Disabled products are excluded from all customer-facing catalog pages.'),
        ('2.2.3.4  Search Product', 'Both customers and admins can search products using keyword matching against product names and categories using a LIKE query pattern.'),
        ('2.2.3.5  View All Products', 'Customers can browse the complete product catalog organized by category sections. Each product card shows name, primary image, base price, and available sizes.'),
        ('2.2.3.6  Filter Products by Category', 'The catalog provides category filter controls allowing customers to restrict displayed products to T-Shirts, Hoodies, Jackets, or Sports Uniforms.'),
    ]),
    ('2.2.4  Product Customization', [
        ('2.2.4.1  Select Fabric Type', 'Customers can select from Cotton, Polyester, Blended, and Fleece fabrics. Each selection updates the dynamic price in real time.'),
        ('2.2.4.2  Select Color', 'Customers can choose the base garment color using an integrated color picker supporting predefined palette and direct hex code entry. The 2D Canvas preview updates instantly.'),
        ('2.2.4.3  Upload Custom Logo', 'Customers can upload a PNG or JPG logo (max 5MB) and specify placement (chest, back, sleeve). The logo appears on the 2D Canvas preview.'),
        ('2.2.4.4  Select Stitching Style', 'Customers select from Flatlock, Overlock, and Chain Stitch options. Each adds a defined price premium.'),
        ('2.2.4.5  Select Print Method', 'Customers choose from Screen Print, Direct to Garment (DTG), Embroidery, and Heat Transfer. Each option has a different pricing tier.'),
        ('2.2.4.6  Select Wash Effect', 'Customers select from Standard, Stone Wash, Acid Wash, and Enzyme Wash finishing treatments.'),
        ('2.2.4.7  View Real-Time 2D Design Preview', 'The Canvas-based preview renders the garment silhouette, applies the selected color, positions the uploaded logo, and updates with every customization change.'),
        ('2.2.4.8  Calculate Dynamic Price', 'The system automatically computes and displays the total price in real time by summing the base price with add-on costs for fabric, stitching, print method, and wash effect.'),
    ]),
    ('2.2.5  AI Chatbot Management', [
        ('2.2.5.1  Send Query to Chatbot', 'Customers access the AI chatbot through a floating widget available on all pages. Queries are sent to the n8n RAG pipeline for processing.'),
        ('2.2.5.2  Get AI Design Recommendation (RAG)', 'The chatbot uses a RAG pipeline that retrieves relevant product documents from ChromaDB and passes them as context to GPT-4o-mini for generating personalized design recommendations.'),
        ('2.2.5.3  View FAQ via Chatbot', 'The chatbot answers frequently asked questions about delivery, returns, care instructions, and bulk pricing using the FAQ knowledge base.'),
        ('2.2.5.4  View Chat History', 'All chatbot conversations are persisted with user ID, message content, sender role, and timestamp. Authenticated customers can review past conversations.'),
        ('2.2.5.5  Admin Manage FAQ Entries', 'Admins can create, update, and delete FAQ entries through the admin panel. Changes are reflected immediately in chatbot responses.'),
    ]),
    ('2.2.6  Order Management', [
        ('2.2.6.1  Create New Order', 'Customers place orders by adding customized items to cart and proceeding to checkout. All customization specifications are captured in the order record.'),
        ('2.2.6.2  Search Order History', 'Customers can search their order history by Order ID, product name, or date range. Admins have access to a global order search across all customers.'),
        ('2.2.6.3  View Order History', 'Customers view a paginated list of all orders showing Order ID, product name, order date, status, and total amount.'),
        ('2.2.6.4  Cancel Order', 'Customers can cancel orders still in Pending status. The system updates order status to Cancelled and logs the cancellation timestamp.'),
        ('2.2.6.5  Track Order Status', 'Each order progresses through five stages: Pending → In Production → Quality Check → Shipped → Delivered. Customers view the current status on the Order Detail page.'),
        ('2.2.6.6  Reorder Previous Order', 'Customers can reorder any completed order. The system pre-populates the customization form with all specifications from the original order.'),
    ]),
    ('2.2.7  Payment Management', [
        ('2.2.7.1  Record Online Payment', 'The system integrates with PayFast to process online payments. Upon successful authorization, a transaction record is created and the order status updated.'),
        ('2.2.7.2  Record Cash on Delivery', 'COD orders are marked as COD-Pending at checkout. Upon delivery confirmation by admin, the payment status is updated to Completed.'),
        ('2.2.7.3  Generate Digital Receipt', 'Upon order placement, the system auto-generates and displays a digital receipt including order ID, itemized list, subtotal, and total amount.'),
        ('2.2.7.4  View All Payments and Transactions', 'Customers view their payment history with dates, amounts, order IDs, and payment methods. Admins access a global transaction ledger.'),
        ('2.2.7.5  Download Invoice', 'Customers can download a PDF invoice for any completed order including AURA-WEAR business details, customer information, and itemized order summary.'),
    ]),
    ('2.2.8  Admin Dashboard Management', [
        ('2.2.8.1  Manage Products (Admin)', 'Admin performs full CRUD operations on product records including bulk CSV import and export.'),
        ('2.2.8.2  Manage Customer Accounts (Admin)', 'Admin views a searchable, paginated customer list and can activate/deactivate accounts.'),
        ('2.2.8.3  View and Update Order Management (Admin)', 'Admin accesses all orders, views details, and updates order status through the production pipeline.'),
        ('2.2.8.4  View Inventory and Stock Alerts (Admin)', 'The dashboard monitors material stock quantities and alerts admins when items fall below defined thresholds.'),
        ('2.2.8.5  Generate Sales Report', 'Admin generates monthly or annual sales reports showing revenue, order counts, and category breakdowns, exportable as PDF or CSV.'),
        ('2.2.8.6  Generate Order Report', 'Admin generates order reports filtered by date range, status, or category showing summary breakdowns and trend analysis.'),
        ('2.2.8.7  View Analytics Dashboard', 'The main admin dashboard displays real-time charts for daily orders, weekly revenue, product category distribution, and customer acquisition trends.'),
        ('2.2.8.8  Send Notification to Customers', 'Admin composes and sends email notifications to all customers (bulk) or specific customers by ID or category.'),
        ('2.2.8.9  Manage Content (Terms, Privacy, FAQ)', 'Admin updates Terms & Conditions, Privacy Policy, and FAQ sections through a rich-text editor. Changes are immediately visible to customers.'),
        ('2.2.8.10  Manage Coupon/Promo Codes', 'Admin creates promotional codes specifying discount type (percentage or fixed), value, usage limits, and validity dates.'),
        ('2.2.8.11  Export Data as CSV', 'Admin exports customer lists, order records, payment transactions, and inventory data as CSV files for external analysis.'),
    ]),
]

for section_title, sub_items in fr_sections:
    heading_para(doc, section_title, 2)
    for sub_title, desc in sub_items:
        heading_para(doc, sub_title, 3)
        styled_para(doc, desc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

heading_para(doc, '2.3  Non-Functional Requirements', 2)
nfr_sections = [
    ('2.3.1  Performance Requirements',
     'The web portal shall load any page within 3 seconds on a standard broadband '
     'connection (minimum 10 Mbps). The 2D customization preview shall reflect changes '
     'within 100ms. The FastAPI backend shall handle a minimum of 50 concurrent user '
     'sessions without degradation. AI chatbot responses shall be delivered within 3 '
     'seconds for typical queries.'),
    ('2.3.2  Usability Requirements',
     'The interface shall be fully responsive across screen widths from 320px (mobile) '
     'to 1920px (widescreen). All primary user flows shall be completable within 5 '
     'user interactions. Error messages shall be clear, specific, and actionable.'),
    ('2.3.3  Security Requirements',
     'All passwords shall be hashed using bcrypt (cost factor ≥ 10). All database '
     'queries shall use parameterized statements to prevent SQL injection. All API '
     'endpoints shall implement input validation. User data shall never be transmitted '
     'in plaintext.'),
    ('2.3.4  Reliability Requirements',
     'The system shall target 99% uptime monthly (≤7.3 hours unplanned downtime). '
     'All database write operations shall be wrapped in transactions with rollback '
     'capability. The system shall implement comprehensive error logging.'),
    ('2.3.5  Scalability Requirements',
     'The FastAPI backend shall be stateless to support horizontal scaling. The database '
     'schema shall support up to 100,000 user records and 500,000 order records without '
     'schema changes.'),
    ('2.3.6  Maintainability & Support Requirements',
     'The codebase shall follow MVC architectural pattern with clear separation of '
     'concerns. All FastAPI endpoints shall be documented using auto-generated OpenAPI/Swagger. '
     'Database migrations shall use Alembic version control.'),
    ('2.3.7  Compatibility Requirements',
     'The web application shall be compatible with the two most recent major versions '
     'of Chrome, Firefox, Safari, and Edge browsers. The application shall function '
     'correctly on devices with a minimum screen width of 320 pixels.'),
    ('2.3.8  Support Requirements',
     'The admin panel shall empower the admin to perform all routine operations without '
     'developer intervention. The system shall include a comprehensive user manual '
     'covering all screen interactions and workflows.'),
]
for title, desc in nfr_sections:
    heading_para(doc, title, 3)
    styled_para(doc, desc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

heading_para(doc, '2.4  External Interface Requirements', 2)
eir_sections = [
    ('2.4.1  User Interfaces',
     'The AURA-WEAR system provides a responsive web interface built with HTML5, CSS3, '
     'and Vanilla JavaScript. The design adopts a clean, minimalist aesthetic emphasizing '
     'product imagery. The 2D Canvas customization panel occupies the right side of the '
     'customizer screen. The AI chatbot appears as a floating widget on all pages.'),
    ('2.4.2  Hardware Interfaces',
     'The system requires no specialized hardware beyond a standard device capable of '
     'running a supported web browser. Image uploads require a device with file access '
     'capabilities. No proprietary hardware dependencies exist.'),
    ('2.4.3  Software Interfaces',
     'FastAPI REST API: Primary interface between frontend and backend (JSON over HTTP). '
     'OpenAI GPT-4o-mini API: AI response generation via n8n workflows. '
     'PayFast Payment Gateway API: Payment processing for online transactions. '
     'Gmail SMTP: OTP delivery via email. '
     'Cloudinary API: Image upload and storage. '
     'ChromaDB: Vector database for RAG semantic search.'),
    ('2.4.4  Communication Interfaces',
     'All client-server communication uses HTTPS with TLS 1.3. REST API calls use JSON '
     'payloads with proper Content-Type headers. The n8n webhook integration uses HTTPS '
     'POST requests. Email notifications use SMTP over port 587 with STARTTLS.'),
]
for title, desc in eir_sections:
    heading_para(doc, title, 3)
    styled_para(doc, desc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 3 – DESIGN
# ─────────────────────────────────────────────────────────────────────────────
print("  Chapter 3: Design...")
chapter_divider(doc, 3, 'DESIGN')

heading_para(doc, '3.1  Use Case Diagram', 2)
styled_para(doc,
    'The Use Case Diagram for AURA-WEAR illustrates the interactions between the system '
    'actors (Customer and Admin) and the system\'s eight functional modules. The diagram '
    'captures all 56 use cases distributed across Security Management, User Profile '
    'Management, Product Catalog Management, Product Customization, AI Chatbot Management, '
    'Order Management, Payment Management, and Admin Dashboard Management.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

# Use Case Diagram — image1
img_path = 'doc_parts/diagrams/image1.png'
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(5.5))
else:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('[Use Case Diagram — AURA-WEAR]')
    r.italic = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)

figure_caption(doc, 'Figure 3: Use Case Diagram — AURA-WEAR AI-Integrated Custom Clothing Website')

heading_para(doc, '3.2  Fully Dressed Use Cases', 2)
styled_para(doc,
    'The following fully dressed use cases describe the detailed interaction specifications '
    'for all 56 functional requirements of the AURA-WEAR system. Each use case includes '
    'the use case ID, name, actor(s), pre-conditions, post-conditions, main flow, '
    'alternative flow, and exception flow.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

for i, uc in enumerate(uc_data):
    uc_name = uc.get('Use Case Name', f'UC-{i+1:02d}')
    heading_para(doc, f'UC-{i+1:02d}: {uc_name}', 3)
    table_caption(doc, f'Table UC-{i+1:02d}: Use Case — {uc_name}')
    add_uc_table(doc, uc)
    doc.add_paragraph()

heading_para(doc, '3.3  Domain Model', 2)
styled_para(doc,
    'The Domain Model identifies the key entities in the AURA-WEAR system and their '
    'relationships. The model captures the business objects that exist in the problem '
    'domain and the associations between them, providing the conceptual foundation for '
    'the database design.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

# Domain model — no dedicated image, show table only
figure_caption(doc, 'Figure 4: Domain Model — AURA-WEAR Entity Relationships (see table below)')

# Domain model table
if domain_rows:
    table_caption(doc, 'Table 4: Domain Model — Key Entities and Relationships')
    add_simple_table(doc, domain_rows[0], domain_rows[1:], col_widths=[1.5, 2.5, 2.5])
    doc.add_paragraph()

heading_para(doc, '3.4  System Sequence Diagrams', 2)
styled_para(doc,
    'System Sequence Diagrams (SSDs) describe the interaction between actors and the '
    'AURA-WEAR system for each use case, showing the sequence of messages exchanged. '
    'The following 56 SSDs correspond to the 56 functional use cases.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

# Summary-level SSD images
for ssd_img_path, ssd_caption in [
    ('doc_parts/diagrams/image3.png', 'Figure 5: System Sequence Diagram (Summary Level) — Customer to System Flow'),
    ('doc_parts/diagrams/image4.png', 'Figure 6: Detailed SSD — Customer Order Interaction'),
    ('doc_parts/diagrams/image5.png', 'Figure 7: SSD — Customer, AI Chatbot, Database, Admin Interaction'),
]:
    if os.path.exists(ssd_img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(ssd_img_path, width=Inches(5.5))
        figure_caption(doc, ssd_caption)
        doc.add_paragraph()

for i, (title, rows) in enumerate(zip(ssd_titles, ssd_data)):
    heading_para(doc, title, 3)
    table_caption(doc, f'Table SSD-{i+1:02d}: {title}')
    add_ssd_table(doc, rows[1:], title)
    doc.add_paragraph()

heading_para(doc, '3.5  Schema Diagram', 2)
styled_para(doc,
    'The database schema for AURA-WEAR consists of 9 tables that store all system data. '
    'The schema is designed to minimize redundancy, enforce referential integrity through '
    'foreign key constraints, and support efficient query patterns for the platform\'s '
    'core operations.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

# ER Diagram — image7
img_path = 'doc_parts/diagrams/image7.png'
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(5.8))
    figure_caption(doc, 'Figure 8: Entity-Relationship (ER) Diagram — AURA-WEAR Database Schema')

for schema_name, schema_rows in zip(schema_names, schema_tables):
    heading_para(doc, schema_name, 3)
    if len(schema_rows) > 1:
        table_caption(doc, f'Table: {schema_name} — Schema Definition')
        add_simple_table(doc, schema_rows[0], schema_rows[1:], col_widths=[2.0, 1.8, 2.7])
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 4 – CONSTRUCTION
# ─────────────────────────────────────────────────────────────────────────────
print("  Chapter 4: Construction...")
chapter_divider(doc, 4, 'CONSTRUCTION')

heading_para(doc, '4.1  Class Diagram', 2)
styled_para(doc,
    'The Class Diagram defines the static structure of AURA-WEAR\'s backend system, '
    'showing classes, their attributes, methods, and relationships. The diagram represents '
    'the data models implemented using SQLAlchemy ORM in the FastAPI backend.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

# No separate class diagram image — table represents the class structure
figure_caption(doc, 'Figure 9: Class Diagram — AURA-WEAR System Classes and Relationships (tabular representation below)')

if len(class_table_rows) > 1:
    table_caption(doc, 'Table 5: Class Diagram — AURA-WEAR System Classes and Attributes')
    add_simple_table(doc, class_table_rows[0], class_table_rows[1:],
                     col_widths=[1.5, 2.0, 1.8, 1.2])
doc.add_paragraph()

heading_para(doc, '4.2  Project Code (Business Logic)', 2)
styled_para(doc,
    'The following code samples demonstrate the core business logic implemented in the '
    'AURA-WEAR FastAPI backend and Vanilla JavaScript frontend.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

heading_para(doc, '4.2.1  Backend: FastAPI Routes', 3)

code_sections = [
    ('auth.py — Authentication Router',
     '''from fastapi import APIRouter, HTTPException, Depends
from fastapi.security import OAuth2PasswordBearer
from jose import jwt, JWTError
from passlib.context import CryptContext
from datetime import datetime, timedelta

router = APIRouter(prefix="/auth", tags=["Authentication"])
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
SECRET_KEY = os.getenv("JWT_SECRET_KEY")
ALGORITHM = "HS256"

@router.post("/register")
async def register(user: UserCreate, db: Session = Depends(get_db)):
    existing = db.query(User).filter(User.email == user.email).first()
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    hashed = pwd_context.hash(user.password)
    new_user = User(name=user.name, email=user.email, password_hash=hashed)
    db.add(new_user)
    db.commit()
    return {"message": "Registration successful"}

@router.post("/login")
async def login(credentials: LoginSchema, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == credentials.email).first()
    if not user or not pwd_context.verify(credentials.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token_data = {"sub": str(user.user_id), "role": user.role}
    token = jwt.encode(token_data, SECRET_KEY, algorithm=ALGORITHM)
    return {"access_token": token, "token_type": "bearer"}'''),

    ('orders.py — Order Management Router',
     '''import uuid
from fastapi import APIRouter, HTTPException, Depends
router = APIRouter(prefix="/orders", tags=["Orders"])

@router.post("/create")
async def create_order(order_data: OrderCreate,
                       current_user: User = Depends(get_current_user),
                       db: Session = Depends(get_db)):
    order_id = f"AW-{uuid.uuid4().hex[:8].upper()}"
    new_order = Order(
        order_id=order_id,
        customer_id=current_user.user_id,
        total_price=order_data.total_price,
        status="Pending",
        payment_method=order_data.payment_method
    )
    db.add(new_order)
    for item in order_data.items:
        order_item = OrderItem(order_id=order_id, **item.dict())
        db.add(order_item)
    db.commit()
    return {"order_id": order_id, "message": "Order placed successfully"}'''),

    ('chatbot.py — AI Chatbot Router',
     '''import httpx
from fastapi import APIRouter
router = APIRouter(prefix="/chatbot", tags=["Chatbot"])
N8N_WEBHOOK = os.getenv("N8N_WEBHOOK_URL")

@router.post("/query")
async def chatbot_query(query: ChatQuery,
                        current_user: User = Depends(get_current_user),
                        db: Session = Depends(get_db)):
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.post(N8N_WEBHOOK,
            json={"query": query.message, "user_id": current_user.user_id})
    answer = response.json().get("response", "I could not process your request.")
    log = ChatLog(user_id=current_user.user_id, query=query.message,
                  response=answer, role="assistant")
    db.add(log)
    db.commit()
    return {"response": answer}'''),
]

for title, code in code_sections:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True; r.font.name = 'Courier New'; r.font.size = Pt(10)
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.3)
    p_code.paragraph_format.space_after = Pt(6)
    r_code = p_code.add_run(code)
    r_code.font.name = 'Courier New'; r_code.font.size = Pt(9)
    doc.add_paragraph()

heading_para(doc, '4.2.2  Frontend: JavaScript Logic', 3)

frontend_code = '''// customizer.js — Real-Time 2D Preview
const canvas = document.getElementById('previewCanvas');
const ctx = canvas.getContext('2d');
let customization = {
    color: '#FFFFFF', logoUrl: null, logoPos: 'chest', washEffect: 'standard'
};

function updatePreview() {
    ctx.clearRect(0, 0, canvas.width, canvas.height);
    // Draw garment silhouette
    ctx.fillStyle = customization.color;
    ctx.beginPath();
    ctx.roundRect(60, 40, 280, 320, 20);
    ctx.fill();
    ctx.strokeStyle = '#333333';
    ctx.lineWidth = 2;
    ctx.stroke();
    // Draw logo if uploaded
    if (customization.logoUrl) {
        const img = new Image();
        img.onload = () => {
            const pos = getLogoPosition(customization.logoPos);
            ctx.drawImage(img, pos.x, pos.y, 80, 80);
        };
        img.src = customization.logoUrl;
    }
    updatePrice();
}

function updatePrice() {
    const priceMap = {
        fabric: { Cotton: 0, Polyester: 200, Blended: 350, Fleece: 500 },
        stitch: { Flatlock: 0, Overlock: 150, 'Chain Stitch': 300 },
        print: { 'Screen Print': 0, DTG: 500, Embroidery: 800, 'Heat Transfer': 400 },
        wash: { Standard: 0, 'Stone Wash': 200, 'Acid Wash': 350, 'Enzyme Wash': 300 }
    };
    let total = basePrice;
    total += priceMap.fabric[selectedFabric] || 0;
    total += priceMap.stitch[selectedStitch] || 0;
    total += priceMap.print[selectedPrint] || 0;
    total += priceMap.wash[selectedWash] || 0;
    document.getElementById('totalPrice').textContent = `Rs. ${total.toLocaleString()}`;
}

document.getElementById('colorPicker').addEventListener('input', (e) => {
    customization.color = e.target.value;
    updatePreview();
});'''

p = doc.add_paragraph()
r = p.add_run('customizer.js — Real-Time 2D Preview & Dynamic Pricing')
r.bold = True; r.font.name = 'Courier New'; r.font.size = Pt(10)
p_code = doc.add_paragraph()
p_code.paragraph_format.left_indent = Inches(0.3)
r_code = p_code.add_run(frontend_code)
r_code.font.name = 'Courier New'; r_code.font.size = Pt(9)

heading_para(doc, '4.3  Sample Queries', 2)
styled_para(doc,
    'The following SQL queries represent the core data access operations used by the '
    'AURA-WEAR system via SQLAlchemy ORM.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

if len(query_rows) > 1:
    table_caption(doc, 'Table 6: Sample SQL Queries — AURA-WEAR Core Data Access Operations')
    add_simple_table(doc, query_rows[0], query_rows[1:], col_widths=[2.3, 4.2])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 5 – TESTING
# ─────────────────────────────────────────────────────────────────────────────
print("  Chapter 5: Testing...")
chapter_divider(doc, 5, 'TESTING')

heading_para(doc, '5.1  Test Cases', 2)
styled_para(doc,
    'This chapter presents the complete test cases for all 56 functional requirements of '
    'the AURA-WEAR system. Each test case includes a test case ID, name, description, '
    'pre-conditions, step-by-step test steps, test data, expected result, actual result, '
    'and pass/fail status. All 56 test cases were executed during the final sprint and '
    'all passed successfully.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

for i, tc in enumerate(tc_data):
    tc_name = tc.get('Test Case Name', f'TC-{i+1:02d}')
    heading_para(doc, f'TC-{i+1:02d} {tc_name}', 3)
    table_caption(doc, f'Table TC-{i+1:02d}: Test Case — {tc_name}')
    add_tc_table(doc, tc)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 6 – USER MANUAL
# ─────────────────────────────────────────────────────────────────────────────
print("  Chapter 6: User Manual...")
chapter_divider(doc, 6, 'USER MANUAL')

styled_para(doc,
    'This chapter provides a comprehensive user manual for the AURA-WEAR system. It '
    'describes each screen of the application, its purpose, and how to interact with it. '
    'The system is divided into three main portals: the Customer Portal, the Admin Panel, '
    'and the AI Chatbot interface.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, 'User Roles in this Application', 2)
styled_para(doc,
    'AURA-WEAR supports two primary user roles:\n\n'
    'Customer: Registered users who can browse the product catalog, use the AI-powered '
    'customization designer to create custom garments, interact with the AI chatbot, '
    'place and track orders, manage payments, and maintain their profile.\n\n'
    'Admin: System administrators who manage all aspects of the platform including '
    'products, customer accounts, orders, inventory, reports, coupons, and content '
    'through the dedicated admin panel.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

heading_para(doc, '6.1  Screenshots', 2)

screens = [
    ('6.1.1  Landing / Home Page',
     'AURA-WEAR\'s home page features a bold hero banner with the tagline "Design Your '
     'Style, Wear Your Story." The navigation bar includes links to Home, Products, '
     'Customizer, and Sign In. A promotional section highlights the AI chatbot feature '
     'and featured products are displayed below the hero section.'),
    ('6.1.2  Registration Screen',
     'The registration page displays a centered card with fields for Full Name, Email '
     'Address, and Password. A password strength indicator dynamically shows the security '
     'level. A Register button submits the form. A link to the sign-in page is provided '
     'for existing users.'),
    ('6.1.3  Sign In Screen',
     'The sign-in page presents a clean login form with email and password fields. A '
     '"Forgot Password?" link appears below the password field. A "Remember Me" checkbox '
     'extends the session to 7 days. The Sign In button submits credentials for authentication.'),
    ('6.1.4  Forgot Password Screen',
     'This screen presents a single-field form asking the user to enter their registered '
     'email address. Submitting the form triggers a 6-digit OTP to be sent to the '
     'provided email via Gmail SMTP.'),
    ('6.1.5  OTP Verification Screen',
     'The OTP verification screen displays a six-digit OTP input field. A Resend OTP '
     'link is available with a 60-second cooldown timer. Correct OTP entry redirects the '
     'user to the Reset Password form.'),
    ('6.1.6  Change Password Screen',
     'The change password form contains three fields: Current Password, New Password, '
     'and Confirm Password. A password strength meter evaluates the new password in real '
     'time. The Update Password button saves the new hashed password.'),
    ('6.1.7  Product Catalog Screen',
     'The product catalog presents a responsive grid of product cards. Category filter '
     'buttons at the top allow filtering by All, T-Shirts, Hoodies, Jackets, and Sports '
     'Uniforms. A search bar enables keyword-based product search. Each card shows '
     'product image, name, category, and base price.'),
    ('6.1.8  Product Detail Screen',
     'The product detail page displays full product information including multiple product '
     'images in a gallery format, detailed description, available sizes from XS to XXL, '
     'fabric options, and the base price. A "Customize This Product" button directs the '
     'user to the design studio.'),
    ('6.1.9  Product Customization Screen (Designer)',
     'The customization screen uses a two-panel layout. The left panel contains all '
     'customization controls: fabric selector, color picker with hex input, stitching '
     'style selector, print method selector, wash effect selector, and logo upload section. '
     'The right panel shows the real-time 2D Canvas preview.'),
    ('6.1.10  Real-Time 2D Preview Panel',
     'The 2D preview panel is built on HTML5 Canvas and shows the garment in front view '
     'by default. The preview dynamically updates to reflect the selected color, uploaded '
     'logo with specified placement, and wash effect. A dynamic price display shows the '
     'total cost updating in real time.'),
    ('6.1.11  Logo Upload Section',
     'The logo upload section includes a Browse button that opens the system file picker '
     'accepting PNG or JPG files up to 5MB. After file selection, a thumbnail preview '
     'appears. Placement options (Chest, Back, Left Sleeve, Right Sleeve) are displayed '
     'as radio buttons.'),
    ('6.1.12  AI Chatbot Widget',
     'The AI Chatbot widget appears as a floating circular button in the bottom-right '
     'corner of every page. Clicking it opens a chat panel with a message input field. '
     'The chatbot displays typing indicators while processing queries and maintains '
     'conversation history for authenticated users.'),
    ('6.1.13  Shopping Cart Screen',
     'The shopping cart lists all customized items with thumbnails, specification summaries, '
     'quantities, and subtotals. Remove buttons allow item deletion. An order summary panel '
     'on the right shows subtotal, applicable coupon discount, and final total. A Proceed '
     'to Checkout button initiates the order flow.'),
    ('6.1.14  Checkout Screen',
     'The checkout screen shows the full order summary on the left and a payment selection '
     'panel on the right. Customers choose between Cash on Delivery or PayFast online '
     'payment. A delivery address field uses the saved profile address by default with '
     'option to modify. Clicking Place Order confirms the purchase.'),
    ('6.1.15  Order Confirmation Screen',
     'After successful order placement, the confirmation screen displays a success animation '
     'with a checkmark icon. The assigned Order ID is prominently shown with instructions '
     'to note it for tracking. Buttons link to Order History and Continue Shopping pages.'),
    ('6.1.16  Order History Screen',
     'The order history lists all orders sorted by most recent. Each row shows Order ID, '
     'product name with thumbnail, order date, current status badge (color-coded), and '
     'total amount. Clicking any order row opens the Order Detail Screen.'),
    ('6.1.17  Order Detail Screen',
     'The order detail screen provides complete information including Order ID, date, '
     'current status in a visual five-stage timeline, itemized product list with '
     'customization specifications, payment details, and delivery address. Cancel Order '
     'and Reorder buttons appear for eligible orders.'),
    ('6.1.18  Order Tracking Screen',
     'The order tracking screen presents a visual progress bar illustrating the five-stage '
     'lifecycle: Pending → In Production → Quality Check → Shipped → Delivered. The '
     'current stage is highlighted in green with an estimated completion timeline.'),
    ('6.1.19  Payment History Screen',
     'Lists all financial transactions in tabular format showing transaction date, '
     'linked Order ID, payment method, amount, and status (Paid/Pending/COD-Pending). '
     'A Download Invoice button appears for completed transactions.'),
    ('6.1.20  Customer Profile Screen',
     'Displays the user\'s profile picture alongside full name, email address, phone '
     'number, delivery address, city, total orders placed, and account creation date. '
     'Edit Profile and Change Password buttons are prominently displayed.'),
    ('6.1.21  Admin Login Screen',
     'A clean admin authentication screen with email and password fields. Only users '
     'with the Admin role in the database can successfully authenticate and access '
     'the admin dashboard.'),
    ('6.1.22  Admin Dashboard',
     'The main admin dashboard displays four KPI cards (Total Orders, Total Revenue, '
     'Active Customers, Inventory Alerts). Below are two HTML5 Canvas charts: a daily '
     'order count bar chart and a weekly revenue line chart. Navigation links to all '
     'admin modules appear in the left sidebar.'),
    ('6.1.23  Customer Management Screen',
     'A paginated, searchable table of all registered customers showing name, email, '
     'registration date, order count, and account status. Activate/Deactivate toggle '
     'buttons allow admins to control account access.'),
    ('6.1.24  Product Management Screen',
     'Lists all products with thumbnails, names, categories, prices, and status '
     'indicators. Add Product, Edit, Enable/Disable, and Delete buttons are provided '
     'for each product. A search bar and category filter control the displayed list.'),
    ('6.1.25  Order Management Screen (Admin)',
     'Displays all customer orders in a sortable table with Order ID, customer name, '
     'date, total, and current status. Clicking an order opens the order detail with '
     'a status update dropdown allowing admins to advance orders through the pipeline.'),
    ('6.1.26  Payment Management Screen',
     'A comprehensive ledger of all platform transactions showing transaction details, '
     'amounts, methods, and statuses. Admins can manually mark COD orders as paid upon '
     'delivery confirmation.'),
    ('6.1.27  Inventory Management Screen',
     'Displays current fabric and material stock levels in a table with color-coded '
     'alerts: green for adequate stock, yellow for low stock, and red for out-of-stock. '
     'Admins can update stock quantities directly from this screen.'),
    ('6.1.28  Reports Screen',
     'The reports screen provides three report types: Sales Report (monthly/annual), '
     'Order Report (filtered by date range and status), and Performance Report (top '
     'products, popular options). Generated reports display inline with export buttons '
     'for PDF and CSV formats.'),
    ('6.1.29  Coupon Management Screen',
     'Admins create and manage promotional coupon codes by specifying code string, '
     'discount type (percentage or fixed amount), discount value, usage limit, and '
     'validity date range. Active coupons are listed with their current usage counts.'),
    ('6.1.30  Content Management Screen',
     'A tabbed interface providing rich-text editors for Terms & Conditions, Privacy '
     'Policy, and FAQ sections. The FAQ tab allows adding, editing, and deleting '
     'individual FAQ entries that are used by the AI chatbot knowledge base. Changes '
     'are saved immediately and reflected on the customer-facing pages.'),
]

# Add user manual screenshot sections
img_path_7 = 'doc_parts/diagrams/image7.png'
for i, (title, desc) in enumerate(screens):
    heading_para(doc, title, 3)
    styled_para(doc, desc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    # Add screenshot placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[Screenshot: {title.split("  ", 1)[-1]}]')
    r.italic = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(128, 128, 128)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'Figure {i+5}: {title.split("  ", 1)[-1]} — AURA-WEAR')
    r2.italic = True; r2.font.name = 'Times New Roman'; r2.font.size = Pt(11)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SAVE DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
output_path = 'AURA_WEAR_FUUAST_Thesis.docx'
print(f"Saving document to {output_path}...")
doc.save(output_path)
print(f"Done! Document saved: {output_path}")
print(f"  Open in Microsoft Word and press Ctrl+A, then F9 to update the Table of Contents.")
