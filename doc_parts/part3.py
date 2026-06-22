from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _h(doc, text, bold=True, size=13, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def _b(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def _uc_table(doc, uc_data):
    # Auto-heading from first two rows (UC ID + Name)
    uc_id = uc_data[0][1] if uc_data else ''
    uc_name = uc_data[1][1] if len(uc_data) > 1 else ''
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{uc_id}: {uc_name}")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # uc_data = list of (label, value) pairs
    table = doc.add_table(rows=len(uc_data), cols=2, style='Table Grid')
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    for i, (label, value) in enumerate(uc_data):
        row = table.rows[i]
        # Label cell
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.name = 'Times New Roman'
        lr.font.size = Pt(11)
        # Value cell
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        vr.font.name = 'Times New Roman'
        vr.font.size = Pt(11)
    doc.add_paragraph('')


def add_part3(doc):
    # Chapter title page
    for _ in range(4):
        doc.add_paragraph('')

    _h(doc, 'CHAPTER # 3:', bold=True, size=24, center=True)
    _h(doc, 'DESIGN', bold=True, size=24, center=True)
    doc.add_page_break()

    # 3.1 Use Case Diagram
    _h(doc, '3.1 Use Case Diagram', bold=True, size=13)
    _b(doc,
       'The Use Case Diagram for AURA-WEAR illustrates the interactions between system actors and '
       'the system\'s functional modules. The primary actors are the Customer and the Admin. The '
       'Customer interacts with modules including Security Management, User Profile, Product Catalog, '
       'Product Customization, AI Chatbot, Order Management, and Payment Management. The Admin '
       'interacts with all management panels including Product Management, Order Management, Customer '
       'Management, Inventory, Analytics, Reports, and Content Management. The diagram captures all '
       '56 functional use cases organized across 8 modules.')

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_note = p_note.add_run('[Use Case Diagram - See attached UML diagram]')
    run_note.italic = True
    run_note.font.name = 'Times New Roman'
    run_note.font.size = Pt(12)

    doc.add_page_break()

    # 3.2 Fully Dressed Use Cases
    _h(doc, '3.2 Fully Dressed Use Cases', bold=True, size=13)

    # UC-01
    _uc_table(doc, [
        ('Use Case ID', 'UC-01'),
        ('Use Case Name', 'Process SignUp'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'User not registered; internet connection available'),
        ('Post-Conditions', 'New account created; user redirected to login page'),
        ('Main Flow',
         '1. User navigates to the registration page.\n'
         '2. User enters full name, email address, and password.\n'
         '3. User clicks the Register button.\n'
         '4. System validates email format and password strength (min 8 chars, at least one number).\n'
         '5. System checks if email is already registered.\n'
         '6. System hashes the password using bcrypt and saves the user record to the database.\n'
         '7. System displays a success message and redirects to the login page.'),
        ('Alternative Flow', 'If email already exists, system displays "Email already registered" error without saving.'),
        ('Exception Flow', 'If database error occurs, system displays "Registration failed. Please try again." and logs the error.'),
    ])
    doc.add_paragraph('')

    # UC-02
    _uc_table(doc, [
        ('Use Case ID', 'UC-02'),
        ('Use Case Name', 'Process SignIn'),
        ('Actor(s)', 'Customer / Admin'),
        ('Pre-Conditions', 'User has a registered account; credentials are correct'),
        ('Post-Conditions', 'User is authenticated; JWT token issued; user redirected to dashboard'),
        ('Main Flow',
         '1. User navigates to the sign-in page.\n'
         '2. User enters registered email and password.\n'
         '3. User clicks the Sign In button.\n'
         '4. System validates input fields are not empty.\n'
         '5. System queries the database for the email.\n'
         '6. System verifies the bcrypt password hash.\n'
         '7. System generates a JWT token with 30-minute expiry.\n'
         '8. System stores the token in localStorage and redirects to the appropriate dashboard.'),
        ('Alternative Flow', 'If remember me is checked, token expiry is extended to 7 days.'),
        ('Exception Flow', 'If credentials are invalid, display "Invalid email or password" without revealing which field is wrong.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-03
    _uc_table(doc, [
        ('Use Case ID', 'UC-03'),
        ('Use Case Name', 'Forgot Password'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'User has a registered account; email address is valid'),
        ('Post-Conditions', 'OTP sent to user\'s email; password reset initiated'),
        ('Main Flow',
         '1. User clicks "Forgot Password" on the login page.\n'
         '2. System displays the Forgot Password form.\n'
         '3. User enters their registered email address.\n'
         '4. User clicks Submit.\n'
         '5. System verifies the email exists in the database.\n'
         '6. System generates a random 6-digit OTP with 10-minute expiry.\n'
         '7. System sends the OTP to the user\'s email via Gmail SMTP.\n'
         '8. System displays "OTP sent to your email address."'),
        ('Alternative Flow', 'If email not found, display "No account associated with this email."'),
        ('Exception Flow', 'If email service fails, display "Unable to send OTP. Please try again later."'),
    ])
    doc.add_paragraph('')

    # UC-04
    _uc_table(doc, [
        ('Use Case ID', 'UC-04'),
        ('Use Case Name', 'Change Account Password'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'User is authenticated; current password is known'),
        ('Post-Conditions', 'Password updated successfully in the database'),
        ('Main Flow',
         '1. Authenticated user navigates to Profile Settings.\n'
         '2. User clicks "Change Password."\n'
         '3. System displays the Change Password form.\n'
         '4. User enters current password, new password, and confirmation.\n'
         '5. System verifies the current password hash.\n'
         '6. System validates new password meets strength requirements.\n'
         '7. System hashes the new password and updates the database.\n'
         '8. System displays "Password changed successfully" and logs out all other sessions.'),
        ('Alternative Flow', 'If confirmation does not match, display "Passwords do not match."'),
        ('Exception Flow', 'If current password is incorrect, display "Current password is wrong."'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-05
    _uc_table(doc, [
        ('Use Case ID', 'UC-05'),
        ('Use Case Name', 'JWT Authentication'),
        ('Actor(s)', 'System (automatic)'),
        ('Pre-Conditions', 'User is making a request to a protected endpoint'),
        ('Post-Conditions', 'Request is allowed or rejected based on token validity'),
        ('Main Flow',
         '1. Client sends HTTP request with Bearer JWT in Authorization header.\n'
         '2. FastAPI middleware intercepts the request.\n'
         '3. System extracts and decodes the JWT token.\n'
         '4. System validates token signature using the secret key.\n'
         '5. System checks token expiry timestamp.\n'
         '6. If valid, system extracts user ID and role from token payload.\n'
         '7. System attaches user context to the request and proceeds to the endpoint handler.'),
        ('Alternative Flow', 'If token is missing, return HTTP 401 Unauthorized.'),
        ('Exception Flow', 'If token is expired, return HTTP 401 with message "Session expired. Please sign in again."'),
    ])
    doc.add_paragraph('')

    # UC-06
    _uc_table(doc, [
        ('Use Case ID', 'UC-06'),
        ('Use Case Name', 'OTP Verification'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'OTP has been sent to user\'s email; OTP has not expired'),
        ('Post-Conditions', 'OTP validated; user proceeds to reset password form'),
        ('Main Flow',
         '1. User receives OTP in email.\n'
         '2. User navigates to the OTP verification screen.\n'
         '3. User enters the 6-digit OTP.\n'
         '4. User clicks Verify.\n'
         '5. System retrieves stored OTP and expiry from database.\n'
         '6. System compares entered OTP with stored value.\n'
         '7. System checks OTP has not expired (within 10 minutes).\n'
         '8. If valid, system marks OTP as used and redirects user to Reset Password form.'),
        ('Alternative Flow', 'If OTP is incorrect, display "Invalid OTP. Please try again." (max 3 attempts).'),
        ('Exception Flow', 'If OTP expired, display "OTP has expired. Please request a new one."'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-07
    _uc_table(doc, [
        ('Use Case ID', 'UC-07'),
        ('Use Case Name', 'Role-Based Access Control'),
        ('Actor(s)', 'System (automatic)'),
        ('Pre-Conditions', 'User is authenticated with a JWT token'),
        ('Post-Conditions', 'Access granted or denied based on user role'),
        ('Main Flow',
         '1. System receives request from authenticated user.\n'
         '2. System extracts role from JWT payload (Customer or Admin).\n'
         '3. System checks the requested endpoint\'s required role.\n'
         '4. If user role matches required role, access is granted.\n'
         '5. Request proceeds to the handler function.\n'
         '6. Response is returned to the user.'),
        ('Alternative Flow', 'If Customer tries to access Admin endpoint, system returns HTTP 403 Forbidden.'),
        ('Exception Flow', 'If role field is missing from JWT, treat as unauthorized and return HTTP 401.'),
    ])
    doc.add_paragraph('')

    # UC-08
    _uc_table(doc, [
        ('Use Case ID', 'UC-08'),
        ('Use Case Name', 'Session Management'),
        ('Actor(s)', 'Customer / Admin'),
        ('Pre-Conditions', 'User is signed in; JWT token stored in localStorage'),
        ('Post-Conditions', 'Session maintained or expired; user redirected to login if expired'),
        ('Main Flow',
         '1. User performs any action on the website.\n'
         '2. Browser sends JWT token with every API request.\n'
         '3. System validates token on each request.\n'
         '4. If token is valid and not expired, session continues.\n'
         '5. System tracks last activity timestamp.\n'
         '6. If 30 minutes of inactivity, token is invalidated.\n'
         '7. Next request returns HTTP 401; frontend detects this.\n'
         '8. Frontend clears localStorage and redirects to login page with notification.'),
        ('Alternative Flow', 'If user clicks Logout manually, token is cleared from localStorage immediately.'),
        ('Exception Flow', 'If localStorage is unavailable, session defaults to in-memory with tab-close expiry.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-09
    _uc_table(doc, [
        ('Use Case ID', 'UC-09'),
        ('Use Case Name', 'Add Customer Profile'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is registered and signed in; profile not yet completed'),
        ('Post-Conditions', 'Profile information saved to database; profile marked as complete'),
        ('Main Flow',
         '1. Customer navigates to their Profile page after first login.\n'
         '2. System displays the Complete Profile form.\n'
         '3. Customer enters phone number, delivery address, city, and optionally uploads a profile picture.\n'
         '4. Customer clicks Save Profile.\n'
         '5. System validates phone number format.\n'
         '6. If profile picture uploaded, system stores it in Cloudinary.\n'
         '7. System saves all profile data linked to user ID.\n'
         '8. System displays "Profile saved successfully."'),
        ('Alternative Flow', 'Customer can skip profile completion and complete it later.'),
        ('Exception Flow', 'If image upload fails, profile saved without picture with a warning message.'),
    ])
    doc.add_paragraph('')

    # UC-10
    _uc_table(doc, [
        ('Use Case ID', 'UC-10'),
        ('Use Case Name', 'Update Customer Profile'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is signed in; profile exists'),
        ('Post-Conditions', 'Updated profile saved to database'),
        ('Main Flow',
         '1. Customer navigates to Profile Settings.\n'
         '2. System displays current profile information in editable form.\n'
         '3. Customer modifies desired fields (name, phone, address, city, picture).\n'
         '4. Customer clicks Update.\n'
         '5. System validates modified fields.\n'
         '6. If new profile picture selected, old image deleted from Cloudinary, new one uploaded.\n'
         '7. System updates the user record in database.\n'
         '8. System displays "Profile updated successfully."'),
        ('Alternative Flow', 'If no changes made, display "No changes detected."'),
        ('Exception Flow', 'If validation fails on any field, display field-specific error messages.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-11
    _uc_table(doc, [
        ('Use Case ID', 'UC-11'),
        ('Use Case Name', 'View Customer Profile'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is signed in'),
        ('Post-Conditions', 'Profile information displayed to the customer'),
        ('Main Flow',
         '1. Customer clicks on their profile icon or navigates to Profile page.\n'
         '2. System fetches user data from the database using JWT-extracted user ID.\n'
         '3. System retrieves profile picture URL from Cloudinary.\n'
         '4. System fetches order count from orders table.\n'
         '5. System displays profile page with personal details, profile picture, order count, and registration date.\n'
         '6. Edit Profile and Change Password buttons are displayed.'),
        ('Alternative Flow', 'If profile is incomplete, display a prompt to complete the profile.'),
        ('Exception Flow', 'If data fetch fails, display cached profile data with a reload option.'),
    ])
    doc.add_paragraph('')

    # UC-12
    _uc_table(doc, [
        ('Use Case ID', 'UC-12'),
        ('Use Case Name', 'Manage Customer Account Status'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; customer accounts exist in the system'),
        ('Post-Conditions', 'Account status updated; customer notified if deactivated'),
        ('Main Flow',
         '1. Admin navigates to Customer Management in admin panel.\n'
         '2. System displays list of all customers with status indicators.\n'
         '3. Admin searches for or selects a customer.\n'
         '4. Admin clicks Activate or Deactivate.\n'
         '5. System prompts for confirmation.\n'
         '6. Admin confirms the action.\n'
         '7. System updates account status in the database.\n'
         '8. If deactivated, system invalidates all active JWT tokens for that user.\n'
         '9. System logs the admin action.'),
        ('Alternative Flow', 'Admin can filter customers by active/inactive status.'),
        ('Exception Flow', 'Admin cannot deactivate another admin account.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-13
    _uc_table(doc, [
        ('Use Case ID', 'UC-13'),
        ('Use Case Name', 'Add New Product'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; product details are ready'),
        ('Post-Conditions', 'New product saved to database; visible in catalog'),
        ('Main Flow',
         '1. Admin navigates to Product Management > Add Product.\n'
         '2. Admin enters product name, category, description, and base price.\n'
         '3. Admin selects available sizes (XS to XXL).\n'
         '4. Admin uploads one or more product images.\n'
         '5. Admin sets fabric options and available customization options.\n'
         '6. Admin clicks Save Product.\n'
         '7. System validates all required fields.\n'
         '8. System uploads images to Cloudinary.\n'
         '9. System saves product record to database with Active status.'),
        ('Alternative Flow', 'Admin can save as Draft to publish later.'),
        ('Exception Flow', 'If image upload fails, prompt admin to retry without saving other fields.'),
    ])
    doc.add_paragraph('')

    # UC-14
    _uc_table(doc, [
        ('Use Case ID', 'UC-14'),
        ('Use Case Name', 'Update Product Details'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; product exists in catalog'),
        ('Post-Conditions', 'Product details updated in database; changes reflected in catalog'),
        ('Main Flow',
         '1. Admin navigates to Product Management and selects a product.\n'
         '2. System displays product edit form with current details.\n'
         '3. Admin modifies desired fields (name, price, description, sizes, images).\n'
         '4. Admin clicks Update Product.\n'
         '5. System validates modified fields.\n'
         '6. If new images added, they are uploaded to Cloudinary.\n'
         '7. System updates the product record in the database.\n'
         '8. System displays "Product updated successfully."'),
        ('Alternative Flow', 'Admin can cancel edit and return to product list without changes.'),
        ('Exception Flow', 'If price set to zero or negative, display "Price must be greater than 0."'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-15
    _uc_table(doc, [
        ('Use Case ID', 'UC-15'),
        ('Use Case Name', 'Enable or Disable Product'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; product exists'),
        ('Post-Conditions', 'Product status toggled; catalog updated accordingly'),
        ('Main Flow',
         '1. Admin navigates to Product Management.\n'
         '2. Admin locates the product by searching or scrolling.\n'
         '3. Admin clicks the Enable/Disable toggle next to the product.\n'
         '4. System prompts for confirmation.\n'
         '5. Admin confirms.\n'
         '6. System updates product status in database.\n'
         '7. Disabled products are hidden from the customer catalog.\n'
         '8. Admin panel shows updated status indicator.'),
        ('Alternative Flow', 'Admin can enable/disable multiple products using bulk action checkboxes.'),
        ('Exception Flow', 'If product has active pending orders, system warns but allows disable.'),
    ])
    doc.add_paragraph('')

    # UC-16
    _uc_table(doc, [
        ('Use Case ID', 'UC-16'),
        ('Use Case Name', 'Search Product'),
        ('Actor(s)', 'Customer / Admin'),
        ('Pre-Conditions', 'Product catalog has at least one active product'),
        ('Post-Conditions', 'Search results displayed matching the query'),
        ('Main Flow',
         '1. User enters a keyword in the search bar.\n'
         '2. User presses Enter or clicks the Search icon.\n'
         '3. System queries the database for products where name or description contains the keyword.\n'
         '4. System returns matching results.\n'
         '5. Results are displayed in a product grid.\n'
         '6. If no results, display "No products found for your search."'),
        ('Alternative Flow', 'Admin search includes disabled products in results (marked differently).'),
        ('Exception Flow', 'If search query is empty, display all active products.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-17
    _uc_table(doc, [
        ('Use Case ID', 'UC-17'),
        ('Use Case Name', 'View All Products'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is on the product catalog page; at least one active product exists'),
        ('Post-Conditions', 'Product catalog displayed with all active products'),
        ('Main Flow',
         '1. Customer navigates to the Products page.\n'
         '2. System queries database for all active products.\n'
         '3. System retrieves product images from Cloudinary URLs.\n'
         '4. System displays products in a responsive grid (3 columns on desktop, 1 on mobile).\n'
         '5. Each product card shows name, image, base price, and available categories.\n'
         '6. Customer can click any product to view its details.'),
        ('Alternative Flow', 'If no products available, display "No products currently available."'),
        ('Exception Flow', 'If image URLs fail to load, display a placeholder image.'),
    ])
    doc.add_paragraph('')

    # UC-18
    _uc_table(doc, [
        ('Use Case ID', 'UC-18'),
        ('Use Case Name', 'Filter Products by Category'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Product catalog page is loaded; multiple categories exist'),
        ('Post-Conditions', 'Filtered products displayed matching the selected category'),
        ('Main Flow',
         '1. Customer clicks on a category filter button (T-Shirts, Hoodies, Jackets, Sports Uniforms).\n'
         '2. System highlights the selected category filter.\n'
         '3. System filters the product grid to show only products in the selected category.\n'
         '4. System updates the product count display.\n'
         '5. Customer can clear filter to show all products.\n'
         '6. Multiple categories can be selected simultaneously.'),
        ('Alternative Flow', 'If selected category has no products, display "No products in this category."'),
        ('Exception Flow', 'Filter state is preserved if user navigates back to catalog.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-19
    _uc_table(doc, [
        ('Use Case ID', 'UC-19'),
        ('Use Case Name', 'Select Fabric Type'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer has selected a product for customization'),
        ('Post-Conditions', 'Fabric type recorded; price updated dynamically'),
        ('Main Flow',
         '1. Customer is on the product customization page.\n'
         '2. System displays available fabric options for the product (Cotton, Polyester, Blended, Fleece).\n'
         '3. Customer clicks on the desired fabric type.\n'
         '4. Selected fabric is highlighted.\n'
         '5. System calculates fabric price premium.\n'
         '6. System updates the dynamic price display.\n'
         '7. Fabric selection is stored in the customization state.'),
        ('Alternative Flow', 'If only one fabric is available, it is pre-selected and not changeable.'),
        ('Exception Flow', 'If fabric stock is unavailable, option is shown as greyed out.'),
    ])
    doc.add_paragraph('')

    # UC-20
    _uc_table(doc, [
        ('Use Case ID', 'UC-20'),
        ('Use Case Name', 'Select Color'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is on the customization page; fabric type selected'),
        ('Post-Conditions', 'Color choice recorded; 2D preview updated with new color'),
        ('Main Flow',
         '1. Customer views the color selection panel.\n'
         '2. System displays a color picker with preset color swatches.\n'
         '3. Customer selects a preset color or enters a hex code.\n'
         '4. System validates the hex code format.\n'
         '5. System updates the 2D Canvas preview to show the garment in the selected color.\n'
         '6. Color value is stored in the customization state.\n'
         '7. No additional price change for color selection.'),
        ('Alternative Flow', 'Customer can type a hex code directly into the input field.'),
        ('Exception Flow', 'If invalid hex code entered, revert to the last valid color.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-21
    _uc_table(doc, [
        ('Use Case ID', 'UC-21'),
        ('Use Case Name', 'Upload Custom Logo'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is on the customization page'),
        ('Post-Conditions', 'Logo uploaded; displayed on 2D preview at specified position'),
        ('Main Flow',
         '1. Customer clicks "Upload Logo" in the customization panel.\n'
         '2. System opens file picker (accepts PNG and JPG, max 5MB).\n'
         '3. Customer selects their logo file.\n'
         '4. System validates file type and size.\n'
         '5. System uploads the file and generates a preview URL.\n'
         '6. Customer selects logo placement (Left Chest, Center, Back).\n'
         '7. System overlays the logo on the 2D Canvas preview at the chosen position.\n'
         '8. Logo reference stored in customization state.'),
        ('Alternative Flow', 'Customer can remove the logo by clicking the Remove Logo button.'),
        ('Exception Flow', 'If file exceeds 5MB, display "File too large. Maximum size is 5MB."'),
    ])
    doc.add_paragraph('')

    # UC-22
    _uc_table(doc, [
        ('Use Case ID', 'UC-22'),
        ('Use Case Name', 'Select Stitching Style'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is on the customization page'),
        ('Post-Conditions', 'Stitching style recorded in customization state'),
        ('Main Flow',
         '1. Customer views the Stitching Style section on the customization panel.\n'
         '2. System displays available options: Flatlock, Overlock, Chain Stitch with brief descriptions.\n'
         '3. Customer clicks on the desired stitching style.\n'
         '4. Selected option is highlighted with a checkmark.\n'
         '5. System updates the dynamic price based on the stitching premium.\n'
         '6. Selection is stored in the customization state.'),
        ('Alternative Flow', 'If no stitching preference, default to Flatlock.'),
        ('Exception Flow', 'No exception: all stitching options are always available.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-23
    _uc_table(doc, [
        ('Use Case ID', 'UC-23'),
        ('Use Case Name', 'Select Print Method'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is on the customization page'),
        ('Post-Conditions', 'Print method recorded; price updated'),
        ('Main Flow',
         '1. Customer views the Print Method section.\n'
         '2. System displays options: Screen Print, DTG (Direct to Garment), Embroidery, Heat Transfer.\n'
         '3. Each option shows a price premium.\n'
         '4. Customer selects a print method.\n'
         '5. System records the selection.\n'
         '6. System updates the dynamic price calculation to include the print method cost.\n'
         '7. 2D preview may show a texture hint for embroidery selection.'),
        ('Alternative Flow', 'Customer can hover over each option to see a description tooltip.'),
        ('Exception Flow', 'No exception: all print methods are standard options.'),
    ])
    doc.add_paragraph('')

    # UC-24
    _uc_table(doc, [
        ('Use Case ID', 'UC-24'),
        ('Use Case Name', 'Select Wash Effect'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is on the customization page; base garment color selected'),
        ('Post-Conditions', 'Wash effect recorded; 2D preview updated to reflect effect'),
        ('Main Flow',
         '1. Customer views the Wash Effect section.\n'
         '2. System displays options: Standard, Stone Wash, Acid Wash, Enzyme Wash.\n'
         '3. Customer clicks on the desired effect.\n'
         '4. System applies a visual filter to the 2D Canvas preview to simulate the wash effect.\n'
         '5. Selection is stored in the customization state.\n'
         '6. Price update applied if premium wash selected.'),
        ('Alternative Flow', 'Default is Standard (no wash effect) if customer skips this section.'),
        ('Exception Flow', 'If Canvas rendering fails, selection is still saved without visual update.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-25
    _uc_table(doc, [
        ('Use Case ID', 'UC-25'),
        ('Use Case Name', 'View Real-Time 2D Design Preview'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is on the customization page'),
        ('Post-Conditions', 'Real-time 2D preview reflects all current customization selections'),
        ('Main Flow',
         '1. Customer opens the product customization page.\n'
         '2. System initializes the HTML5 Canvas with the product base template.\n'
         '3. System renders the current garment color on the canvas.\n'
         '4. As customer makes each customization selection (color, logo, wash), system updates the canvas in real time.\n'
         '5. Logo is rendered at the selected placement position.\n'
         '6. Wash effect filter is applied.\n'
         '7. Customer can zoom in on the preview.\n'
         '8. Customer can toggle between front and back views.'),
        ('Alternative Flow', 'If browser does not support Canvas, display a static placeholder image.'),
        ('Exception Flow', 'If preview rendering lags, debounce updates to 300ms after last user input.'),
    ])
    doc.add_paragraph('')

    # UC-26
    _uc_table(doc, [
        ('Use Case ID', 'UC-26'),
        ('Use Case Name', 'Calculate Dynamic Price'),
        ('Actor(s)', 'Customer (automatic)'),
        ('Pre-Conditions', 'Customer is on the customization page; at least one customization made'),
        ('Post-Conditions', 'Updated total price displayed in real time'),
        ('Main Flow',
         '1. Customer makes any customization selection.\n'
         '2. System triggers the price calculation function.\n'
         '3. System starts with the base product price.\n'
         '4. System adds fabric type premium (Cotton: +0%, Polyester: +5%, Blended: +10%, Fleece: +15%).\n'
         '5. System adds print method cost (Screen Print: +PKR 300, DTG: +PKR 500, Embroidery: +PKR 800, Heat Transfer: +PKR 400).\n'
         '6. System adds stitching premium if applicable.\n'
         '7. System applies any active coupon code discount.\n'
         '8. System displays updated total price in PKR.'),
        ('Alternative Flow', 'If coupon code applied, show discounted price with original crossed out.'),
        ('Exception Flow', 'If price calculation results in negative (unlikely), cap at base price.'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()

    # UC-27
    _uc_table(doc, [
        ('Use Case ID', 'UC-27'),
        ('Use Case Name', 'Send Query to Chatbot'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is on any page of the website; chatbot widget is loaded'),
        ('Post-Conditions', 'Query sent to n8n RAG pipeline; response displayed in chat window'),
        ('Main Flow',
         '1. Customer clicks the chatbot bubble icon (floating button).\n'
         '2. Chatbot widget opens with greeting message.\n'
         '3. Customer types their query in the input field.\n'
         '4. Customer presses Enter or clicks Send.\n'
         '5. System sends the query to the n8n webhook endpoint.\n'
         '6. n8n workflow processes the query through the RAG pipeline.\n'
         '7. Response is returned as JSON.\n'
         '8. System displays the AI response in the chat window with a typing animation.'),
        ('Alternative Flow', 'If customer is not logged in, chatbot works in anonymous mode without history saving.'),
        ('Exception Flow', 'If n8n endpoint is unreachable, display "AI assistant is temporarily unavailable."'),
    ])
    doc.add_paragraph('')

    # UC-28
    _uc_table(doc, [
        ('Use Case ID', 'UC-28'),
        ('Use Case Name', 'Get AI Design Recommendation (RAG)'),
        ('Actor(s)', 'Customer (via chatbot)'),
        ('Pre-Conditions', 'Customer has sent a fabric or design-related query; n8n is running'),
        ('Post-Conditions', 'Context-aware AI recommendation returned and displayed'),
        ('Main Flow',
         '1. Customer asks a design question (e.g., "What fabric is best for summer hoodies?").\n'
         '2. n8n receives query via webhook.\n'
         '3. n8n Node 1 queries SQLite database for relevant fabric/product records matching query keywords.\n'
         '4. n8n Node 2 formats the retrieved database context.\n'
         '5. n8n Node 3 constructs OpenAI GPT-4o-mini prompt with query + database context.\n'
         '6. n8n Node 4 calls OpenAI API and receives recommendation.\n'
         '7. n8n returns formatted response to the website.\n'
         '8. System displays the recommendation with relevant product suggestions.'),
        ('Alternative Flow', 'If no relevant database records found, GPT responds from its general knowledge.'),
        ('Exception Flow', 'If OpenAI API quota exceeded, display "AI recommendations temporarily unavailable. Please try again later."'),
    ])
    doc.add_paragraph('')

    doc.add_page_break()
