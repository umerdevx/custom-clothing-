from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _h(doc, text, bold=True, size=13, center=False, underline=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold; run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    if underline: run.underline = True
    return p


def _b(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    return p


def add_part2(doc):
    # ── Chapter title page ──────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CHAPTER # 2:")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ANALYSIS")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)

    doc.add_page_break()

    # ── 2.1 Introduction ────────────────────────────────────────────────────
    _h(doc, "2.1 Introduction", size=14)

    # 2.1.1 Purpose
    _h(doc, "2.1.1 Purpose", size=13)
    _b(doc,
       "This document defines the functional and non-functional requirements for the AURA-WEAR "
       "AI-Integrated Custom Clothing Website system. It serves as a formal contract between the "
       "development team and project stakeholders, establishing a shared understanding of what the "
       "system must do and the quality standards it must meet. The document guides all subsequent "
       "design, implementation, and testing phases of the project.")

    # 2.1.2 Scope
    _h(doc, "2.1.2 Scope", size=13)
    _b(doc,
       "The AURA-WEAR system encompasses three primary components: a customer-facing web portal, "
       "an administrative management panel, and an AI-powered chatbot assistant. The scope of this "
       "document covers the following functional areas that are in scope: user authentication and "
       "authorization, product catalog browsing and management, garment customization with real-time "
       "2D preview, order placement and tracking, payment processing, AI chatbot for design "
       "recommendations and FAQ, and the admin dashboard for managing all aspects of the platform.")
    _b(doc,
       "The following areas are explicitly out of scope for this release: a native mobile application "
       "for Android or iOS, integration with any physical retail or point-of-sale store system, "
       "and third-party marketplace integrations. These items may be considered for future iterations "
       "of the system.")

    # 2.1.3 Definitions, Acronyms, and Abbreviations
    _h(doc, "2.1.3 Definitions, Acronyms, and Abbreviations", size=13)
    _b(doc,
       "The following table defines the key terms, acronyms, and abbreviations used throughout "
       "this document.")

    terms = [
        ("RAG",   "Retrieval-Augmented Generation: an AI technique that combines information "
                  "retrieval from a knowledge base with a generative language model to produce "
                  "accurate, context-aware responses."),
        ("JWT",   "JSON Web Token: a compact, URL-safe means of representing claims to be "
                  "transferred between two parties, used here for stateless user authentication."),
        ("OTP",   "One-Time Password: a 6-digit numeric code sent to a user's registered email "
                  "address for identity verification during password reset."),
        ("2D Preview",
                  "Two-Dimensional Design Preview: a real-time Canvas-based visualization of the "
                  "customized garment displayed within the browser."),
        ("FR",    "Functional Requirement: a requirement that specifies a behavior or function "
                  "the system must be capable of performing."),
        ("NFR",   "Non-Functional Requirement: a requirement that defines quality attributes, "
                  "constraints, or performance criteria for the system."),
        ("API",   "Application Programming Interface: a set of definitions and protocols for "
                  "building and integrating application software."),
        ("CRUD",  "Create, Read, Update, Delete: the four basic operations of persistent storage "
                  "used in database interactions."),
        ("COD",   "Cash on Delivery: a payment method in which the customer pays for goods at "
                  "the time of delivery rather than in advance."),
        ("DTG",   "Direct to Garment: a printing method in which specialized inkjet technology "
                  "prints the design directly onto the fabric surface."),
        ("SRS",   "Software Requirements Specification: this document; a comprehensive description "
                  "of the intended purpose, environment, and requirements of the software system."),
        ("UI",    "User Interface: the visual and interactive elements through which a user "
                  "interacts with the software application."),
    ]

    tbl = doc.add_table(rows=1, cols=2, style='Table Grid')
    hdr = tbl.rows[0].cells
    for cell, label in zip(hdr, ["Term", "Definition"]):
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for term, definition in terms:
        row = tbl.add_row().cells
        r0 = row[0].paragraphs[0].add_run(term)
        r0.bold = True
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(12)
        r1 = row[1].paragraphs[0].add_run(definition)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)

    doc.add_paragraph()

    # 2.1.4 Overview
    _h(doc, "2.1.4 Overview", size=13)
    _b(doc,
       "The remainder of this document is organized as follows. Section 2.2 presents the complete "
       "set of Functional Requirements organized into eight management modules: Security, User "
       "Profile, Product Catalog, Product Customization, AI Chatbot, Order, Payment, and Admin "
       "Dashboard. Section 2.3 details the Non-Functional Requirements covering performance, "
       "usability, security, reliability, scalability, maintainability, compatibility, and support "
       "standards. Section 2.4 specifies External Interface Requirements including user, hardware, "
       "software, and communication interfaces.")

    # ── 2.2 Functional Requirements ─────────────────────────────────────────
    _h(doc, "2.2 Functional Requirements", size=14)
    _b(doc,
       "This section enumerates all functional requirements for the AURA-WEAR system. Requirements "
       "are grouped into modules corresponding to major system features. Each requirement describes "
       "a discrete capability the system must provide.")

    # ── 2.2.1 Security Management ───────────────────────────────────────────
    _h(doc, "2.2.1 Security Management", size=13)

    _h(doc, "2.2.1.1 Process SignUp", size=12, underline=True)
    _b(doc,
       "The system shall allow new users to register by providing their full name, email address, "
       "and password through the registration form. The system validates the email format using a "
       "standard regex pattern and enforces a minimum 8-character password containing at least one "
       "numeric character, hashing the accepted password using bcrypt before persisting it to the "
       "database. A confirmation message is displayed upon successful registration, and the user is "
       "redirected to the login page.")

    _h(doc, "2.2.1.2 Process SignIn", size=12, underline=True)
    _b(doc,
       "Registered users shall authenticate with the system using their email address and password "
       "through the sign-in form. The system retrieves the stored bcrypt hash for the provided email "
       "and verifies the submitted password; upon a successful match, a JWT access token with a "
       "30-minute expiry is issued and stored in the browser's localStorage. Failed authentication "
       "attempts display an appropriate error message without disclosing whether the email or "
       "password was incorrect.")

    _h(doc, "2.2.1.3 Forgot Password", size=12, underline=True)
    _b(doc,
       "Users who have forgotten their password can initiate a reset flow by entering their "
       "registered email address on the Forgot Password page. The system generates a cryptographically "
       "random 6-digit OTP, stores it with a timestamp, and dispatches it to the provided email "
       "address via Gmail SMTP integration. The OTP expires after 10 minutes from the time of "
       "generation, after which a new OTP must be requested.")

    _h(doc, "2.2.1.4 Change Account Password", size=12, underline=True)
    _b(doc,
       "Authenticated users can change their account password from the profile settings page. The "
       "system first requires the user to enter their current password for verification against the "
       "stored bcrypt hash, then enforces the same password-strength rules as registration (minimum "
       "8 characters, at least one digit) on the new password. Upon validation, the new password "
       "is hashed with bcrypt and the updated hash is persisted to the database.")

    _h(doc, "2.2.1.5 JWT Authentication", size=12, underline=True)
    _b(doc,
       "All protected API endpoints in the FastAPI backend require a valid Bearer JWT token in the "
       "HTTP Authorization header. On each request to a protected route, the system validates the "
       "token's cryptographic signature against the server-side secret key and checks that the token "
       "has not expired. Requests carrying an invalid, tampered, or expired token receive an HTTP "
       "401 Unauthorized response, and the client-side application redirects the user to the "
       "login page.")

    _h(doc, "2.2.1.6 OTP Verification", size=12, underline=True)
    _b(doc,
       "The system shall generate and validate 6-digit numeric OTPs for both password reset and "
       "account verification workflows. Each OTP is stored in the database alongside the associated "
       "user identifier and an expiry timestamp set to 10 minutes from creation. Once an OTP is "
       "successfully verified it is immediately marked as used and cannot be reused for any "
       "subsequent verification attempt.")

    _h(doc, "2.2.1.7 Role-Based Access Control", size=12, underline=True)
    _b(doc,
       "The system implements two distinct roles: Customer and Admin. Customers, created through "
       "the public registration page, are granted access to the product catalog, customization "
       "tool, order management, and chatbot features. Admin accounts are provisioned exclusively "
       "by existing administrators and have unrestricted access to all management panels including "
       "product management, user management, order management, reporting, and content management. "
       "Admin accounts cannot be created through the public registration endpoint.")

    _h(doc, "2.2.1.8 Session Management", size=12, underline=True)
    _b(doc,
       "The system manages authenticated user sessions using JWT tokens persisted in the browser's "
       "localStorage. Every page load and API call checks the token's expiry; when the token "
       "expires, the system automatically clears the stored token, redirects the user to the login "
       "page, and displays a session-expiry notification informing the user that their session has "
       "ended and prompting them to sign in again.")

    # ── 2.2.2 User Profile Management ───────────────────────────────────────
    _h(doc, "2.2.2 User Profile Management", size=13)

    _h(doc, "2.2.2.1 Add Customer Profile", size=12, underline=True)
    _b(doc,
       "After completing registration, a customer can enhance their account by adding supplementary "
       "profile information including their phone number, full delivery address, city, and a "
       "profile picture image. All profile information is validated on submission and stored in the "
       "database linked to the customer's user account record, enabling pre-population of delivery "
       "details during the checkout process.")

    _h(doc, "2.2.2.2 Update Customer Profile", size=12, underline=True)
    _b(doc,
       "Authenticated customers can modify their saved profile information: including display name, "
       "phone number, delivery address, city, and profile picture: through the Profile Settings "
       "page. Each field is individually validated before submission, and all accepted changes are "
       "persisted to the database immediately, with the updated information reflected across "
       "the platform without requiring a page reload.")

    _h(doc, "2.2.2.3 View Customer Profile", size=12, underline=True)
    _b(doc,
       "Customers can navigate to their Profile page to view a consolidated summary of their "
       "account information including their full name, email address, phone number, delivery "
       "address, profile picture, total order count, and account registration date. The profile "
       "page provides quick-access links to order history and account settings.")

    _h(doc, "2.2.2.4 Manage Customer Account Status", size=12, underline=True)
    _b(doc,
       "Admin users can view a complete list of all registered customer accounts from the admin "
       "panel, including each customer's registration date, order count, and current account status. "
       "Admin can toggle any customer account between Active and Deactivated status; deactivated "
       "accounts receive an appropriate error message when attempting to log in and are denied "
       "access to all authenticated features until reactivated by an administrator.")

    # ── 2.2.3 Product Catalog Management ────────────────────────────────────
    _h(doc, "2.2.3 Product Catalog Management", size=13)

    _h(doc, "2.2.3.1 Add New Product", size=12, underline=True)
    _b(doc,
       "Admin users can add new products to the AURA-WEAR catalog by completing a product creation "
       "form that captures the product name, category (T-Shirt, Hoodie, Jacket, or Sports Uniform), "
       "detailed description, base price, available sizes ranging from XS to XXL, and one or more "
       "product images uploaded to Cloudinary. Newly created products are assigned Active status "
       "by default and immediately appear in the customer-facing product catalog.")

    _h(doc, "2.2.3.2 Update Product Details", size=12, underline=True)
    _b(doc,
       "Admin can select any existing product from the admin panel and edit its attributes including "
       "name, description, category, base price, available size options, and product images. All "
       "submitted changes are validated and saved to the database immediately, with updates reflected "
       "in real time on the customer-facing product catalog without requiring manual cache "
       "invalidation.")

    _h(doc, "2.2.3.3 Enable or Disable Product", size=12, underline=True)
    _b(doc,
       "Admin can toggle a product's availability status between Enabled and Disabled through a "
       "status control in the product management panel. Disabled products are filtered out of all "
       "customer-facing catalog views and search results, making them invisible to customers, but "
       "their records are retained in the database to preserve referential integrity with existing "
       "order history records.")

    _h(doc, "2.2.3.4 Search Product", size=12, underline=True)
    _b(doc,
       "Both customers on the catalog page and admin users in the management panel can search for "
       "products by entering keywords into the search bar, which matches against product names and "
       "categories using a case-insensitive partial-match query. Search results are displayed in "
       "a dynamically filtered product grid that updates as the user types, highlighting "
       "matching products and hiding non-matching items.")

    _h(doc, "2.2.3.5 View All Products", size=12, underline=True)
    _b(doc,
       "Customers can browse the complete AURA-WEAR product catalog organized by category sections "
       "on the Products page. Each product card displays the product name, primary image, base "
       "price, and available fabric options. Customers can click any product card to navigate to "
       "the detailed product page where they can begin the customization process.")

    _h(doc, "2.2.3.6 Filter Products by Category", size=12, underline=True)
    _b(doc,
       "The customer product catalog page provides category filter controls allowing customers to "
       "restrict the displayed products to one or more specific categories: T-Shirts, Hoodies, "
       "Jackets, and Sports Uniforms. Multiple category filters can be applied simultaneously, "
       "and the product grid updates dynamically to display only products matching the selected "
       "categories. Clearing all filters restores the full catalog view.")

    # ── 2.2.4 Product Customization ──────────────────────────────────────────
    _h(doc, "2.2.4 Product Customization", size=13)

    _h(doc, "2.2.4.1 Select Fabric Type", size=12, underline=True)
    _b(doc,
       "When customizing a product, customers can select from the available fabric types: Cotton, "
       "Polyester, Blended, and Fleece. Each fabric option is presented with a brief description "
       "of its properties (e.g., breathability, warmth, durability). The selection of a fabric "
       "type dynamically adjusts the calculated total price displayed in the real-time price "
       "summary panel based on the fabric's associated price modifier.")

    _h(doc, "2.2.4.2 Select Color", size=12, underline=True)
    _b(doc,
       "Customers can choose the base color of their garment using an integrated color picker "
       "control that supports selection from a predefined color palette as well as direct hex "
       "code entry for precise color specification. Upon selecting a color, the real-time 2D "
       "preview panel updates immediately to render the garment in the chosen color, enabling "
       "customers to visualize the final appearance before placing their order.")

    _h(doc, "2.2.4.3 Upload Custom Logo", size=12, underline=True)
    _b(doc,
       "Customers can upload a custom logo or graphic image in PNG or JPG format, with a maximum "
       "file size of 5MB, to be applied to the garment. The customer specifies the placement "
       "position of the logo from the available options: left chest, center front, or back. "
       "Upon upload and position selection, the logo image is rendered on the 2D preview at "
       "the specified position, allowing the customer to confirm the placement before ordering.")

    _h(doc, "2.2.4.4 Select Stitching Style", size=12, underline=True)
    _b(doc,
       "Customers can select from three stitching styles for their custom garment: Flatlock, "
       "Overlock, and Chain Stitch. Each option is displayed in the customization panel with "
       "a concise description of its characteristics and typical applications, enabling customers "
       "to make an informed selection that suits their intended use for the garment.")

    _h(doc, "2.2.4.5 Select Print Method", size=12, underline=True)
    _b(doc,
       "Customers can choose their preferred print method for applying designs and logos to the "
       "garment from four available options: Screen Print, Direct to Garment (DTG), Embroidery, "
       "and Heat Transfer. Each print method has a different associated price adjustment that is "
       "applied to the dynamic price calculation and reflected immediately in the total price "
       "displayed to the customer.")

    _h(doc, "2.2.4.6 Select Wash Effect", size=12, underline=True)
    _b(doc,
       "Customers can select a finishing wash treatment to be applied to the garment from four "
       "options: Standard, Stone Wash, Acid Wash, and Enzyme Wash. Each wash effect option is "
       "described in the customization panel to help customers understand the visual and textural "
       "outcome, and the 2D preview updates its appearance to approximate the selected wash "
       "effect wherever technically feasible.")

    _h(doc, "2.2.4.7 View Real-Time 2D Design Preview", size=12, underline=True)
    _b(doc,
       "The product customization page displays a real-time, Canvas-based 2D preview of the "
       "garment on the right side of the customization panel. The preview renders the garment "
       "silhouette, applies the selected color fill, and overlays the uploaded logo at the "
       "specified placement position. The preview updates instantaneously: with a target "
       "refresh latency of 500 milliseconds or less: as the customer makes any customization "
       "selection, providing immediate visual feedback without requiring a page reload.")

    _h(doc, "2.2.4.8 Calculate Dynamic Price", size=12, underline=True)
    _b(doc,
       "The system automatically computes and displays the total order price in real time within "
       "the customization panel as the customer makes selections. The price is calculated by "
       "summing the product base price with the applicable price adjustments for the selected "
       "fabric type, print method, stitching style, and wash effect. The displayed price updates "
       "instantly with each customer selection, ensuring full pricing transparency before "
       "the customer proceeds to checkout.")

    # ── 2.2.5 AI Chatbot Management ──────────────────────────────────────────
    _h(doc, "2.2.5 AI Chatbot Management", size=13)

    _h(doc, "2.2.5.1 Send Query to Chatbot", size=12, underline=True)
    _b(doc,
       "Customers can access the AI chatbot assistant through a floating chat widget available "
       "on all pages of the AURA-WEAR website. The customer types a natural language query "
       "regarding products, fabric types, customization options, pricing, or order status into "
       "the chat input field and submits it. The system processes the query through the n8n "
       "workflow pipeline and returns a formatted response within the chat interface.")

    _h(doc, "2.2.5.2 Get AI Design Recommendation (RAG)", size=12, underline=True)
    _b(doc,
       "The AI chatbot leverages a Retrieval-Augmented Generation (RAG) pipeline implemented "
       "via the n8n workflow automation platform. When a customer submits a design-related query, "
       "the system first queries the SQLite database to retrieve relevant product catalog data, "
       "fabric specifications, and available customization options. This retrieved context is then "
       "passed along with the customer's query to the OpenAI GPT-4o-mini model, which generates "
       "an accurate, context-aware design recommendation grounded in the actual AURA-WEAR "
       "product data.")

    _h(doc, "2.2.5.3 View FAQ via Chatbot", size=12, underline=True)
    _b(doc,
       "The AI chatbot is capable of answering frequently asked questions on topics including "
       "estimated delivery times, the return and refund policy, garment care instructions, and "
       "bulk order pricing tiers. When a customer query matches an FAQ topic, the system retrieves "
       "the relevant FAQ entry from the database and the chatbot returns a formatted, readable "
       "response directly within the chat interface.")

    _h(doc, "2.2.5.4 View Chat History", size=12, underline=True)
    _b(doc,
       "All chatbot conversations for authenticated customers are persisted to the database with "
       "the customer's user identifier, the message content, the sender role (user or assistant), "
       "and a timestamp. Authenticated customers can access their complete chatbot conversation "
       "history from their profile page, where previous sessions are listed in chronological "
       "order for reference.")

    _h(doc, "2.2.5.5 Admin Manage FAQ Entries", size=12, underline=True)
    _b(doc,
       "Admin users can manage the chatbot's FAQ knowledge base through a dedicated FAQ "
       "management section in the admin panel. Admin can create new FAQ entries by specifying "
       "a question and its corresponding answer, edit existing entries, or delete outdated entries. "
       "All changes to the FAQ database take effect immediately, ensuring that the chatbot "
       "serves up-to-date information in its responses without requiring a system restart.")

    # ── 2.2.6 Order Management ───────────────────────────────────────────────
    _h(doc, "2.2.6 Order Management", size=13)

    _h(doc, "2.2.6.1 Create New Order", size=12, underline=True)
    _b(doc,
       "After completing the product customization process, a customer can add the configured "
       "item to their cart and proceed to place an order. The system captures and records all "
       "customization specifications (fabric, color, logo, stitching, print method, wash effect), "
       "calculates the final total price, generates a unique alphanumeric Order ID, associates "
       "the order with the customer's account, and assigns an initial status of Pending. A "
       "confirmation screen displays the order summary and Order ID upon successful placement.")

    _h(doc, "2.2.6.2 Search Order History", size=12, underline=True)
    _b(doc,
       "Customers can search within their personal order history by entering an Order ID, "
       "product name keyword, or selecting a date range on the Order History page. Admin users "
       "have access to a global order search that queries across all customer orders using the "
       "same search criteria. Matching orders are displayed in a filtered results list sorted "
       "by most recent date.")

    _h(doc, "2.2.6.3 View Order History", size=12, underline=True)
    _b(doc,
       "Customers can view a paginated list of all their past and current orders on the Order "
       "History page, with each entry displaying the Order ID, product name, order date, current "
       "status, and total amount. Selecting an individual order from the list navigates to the "
       "Order Detail page, which presents the complete order information including all "
       "customization specifications and the payment method used.")

    _h(doc, "2.2.6.4 Cancel Order", size=12, underline=True)
    _b(doc,
       "Customers can cancel orders that are still in Pending status by selecting the cancel "
       "option on the Order Detail page and confirming the cancellation. The system updates the "
       "order status to Cancelled, records the cancellation timestamp, and sends a cancellation "
       "confirmation to the customer's email. Orders in any status beyond Pending can only be "
       "cancelled by an Admin user through the admin order management panel.")

    _h(doc, "2.2.6.5 Track Order Status", size=12, underline=True)
    _b(doc,
       "Each order progresses through a defined lifecycle of statuses representing its production "
       "and delivery stages: Pending, In Production, Quality Check, Shipped, and Delivered. "
       "Customers can view the current status of any order from the Order Detail page, with "
       "a visual progress indicator showing the completed and upcoming stages. Admin users "
       "update the order status through the admin panel at each stage transition.")

    _h(doc, "2.2.6.6 Reorder Previous Order", size=12, underline=True)
    _b(doc,
       "Customers can initiate a reorder of any previously completed order by selecting the "
       "Reorder option on the Order Detail page. The system pre-populates the customization "
       "form with all specifications from the original order: including fabric type, color, "
       "logo, stitching style, print method, and wash effect: and presents the form to the "
       "customer for review and confirmation before creating a new order record.")

    # ── 2.2.7 Payment Management ─────────────────────────────────────────────
    _h(doc, "2.2.7 Payment Management", size=13)

    _h(doc, "2.2.7.1 Record Online Payment", size=12, underline=True)
    _b(doc,
       "The system integrates with PayFast and Stripe payment gateways to process online credit "
       "and debit card payments. Upon successful payment authorization from the payment gateway, "
       "the system records a transaction entry in the payments database table capturing the order "
       "ID, customer ID, payment amount, payment method, transaction reference number, payment "
       "timestamp, and final payment status. The associated order status is updated to reflect "
       "confirmed payment.")

    _h(doc, "2.2.7.2 Record Cash on Delivery", size=12, underline=True)
    _b(doc,
       "Customers who prefer to pay upon receipt can select Cash on Delivery (COD) as their "
       "payment method at the checkout stage. The system marks the order's payment status as "
       "COD-Pending, generates a unique COD reference number for use by the delivery agent, "
       "and includes the COD instructions in the order confirmation email sent to the customer. "
       "The payment status is updated to Paid by the admin upon confirmed delivery and collection.")

    _h(doc, "2.2.7.3 Generate Digital Receipt", size=12, underline=True)
    _b(doc,
       "Upon successful payment confirmation or order placement with COD, the system "
       "automatically generates a digital receipt and displays it to the customer on screen. "
       "The receipt includes the order ID, a complete customization summary, an itemized pricing "
       "breakdown showing base price and all adjustments, the payment method, the total amount "
       "paid, and a unique receipt reference number. The receipt is also stored and accessible "
       "from the customer's order history.")

    _h(doc, "2.2.7.4 View All Payments and Transactions", size=12, underline=True)
    _b(doc,
       "Authenticated customers can access a payment history page listing all their transactions "
       "with dates, amounts, associated order IDs, and payment methods. Admin users have access "
       "to a comprehensive global transactions view covering all customer payments, with filter "
       "controls for date range, payment method (online/COD), and payment status to facilitate "
       "financial reconciliation and reporting.")

    _h(doc, "2.2.7.5 Download Invoice", size=12, underline=True)
    _b(doc,
       "From the Order Detail page, customers can download a formally formatted invoice for any "
       "completed order as a PDF document. The generated invoice includes the AURA-WEAR business "
       "details, customer information, a complete list of ordered items with their customization "
       "specifications, the pricing breakdown, the payment method and status, and all relevant "
       "order and invoice reference numbers.")

    # ── 2.2.8 Admin Dashboard Management ────────────────────────────────────
    _h(doc, "2.2.8 Admin Dashboard Management", size=13)

    _h(doc, "2.2.8.1 Manage Products (Admin)", size=12, underline=True)
    _b(doc,
       "Admin users can perform full Create, Read, Update, and Delete (CRUD) operations on "
       "product records through the product management section of the admin panel. In addition "
       "to individual product editing, admin can perform bulk status toggles to enable or disable "
       "multiple products simultaneously and can update inventory quantity counts for each "
       "product-size combination to reflect current material stock levels.")

    _h(doc, "2.2.8.2 Manage Customer Accounts (Admin)", size=12, underline=True)
    _b(doc,
       "The admin panel presents a paginated, searchable customer list displaying each customer's "
       "name, email address, registration date, total number of orders placed, and current account "
       "status. Admin users can activate or deactivate any customer account directly from this "
       "list view and can click through to a detailed customer profile showing full contact "
       "information and order history.")

    _h(doc, "2.2.8.3 View and Update Order Management (Admin)", size=12, underline=True)
    _b(doc,
       "Admin can access a centralized order management panel displaying all customer orders "
       "sortable by order date, status, or total value. Admin can view the complete details of "
       "any individual order: including full customization specifications: and can advance the "
       "order's status through each lifecycle stage (Pending → In Production → Quality Check → "
       "Shipped → Delivered) via a status update control, triggering automated notifications "
       "to the customer at each stage.")

    _h(doc, "2.2.8.4 View Inventory and Stock Alerts (Admin)", size=12, underline=True)
    _b(doc,
       "The admin dashboard includes an inventory management view where admin can monitor current "
       "fabric and material stock quantities across all product categories. The system evaluates "
       "each material's quantity against a configurable low-stock threshold and surfaces low-stock "
       "alert indicators on the admin dashboard for materials that have fallen below the threshold, "
       "enabling timely procurement decisions.")

    _h(doc, "2.2.8.5 Generate Sales Report", size=12, underline=True)
    _b(doc,
       "Admin can generate comprehensive sales reports for a selected month or year through the "
       "reporting section of the admin panel. The generated report presents total revenue figures, "
       "total order counts, a ranking of the most popular products by order volume, and a list "
       "of top customers by order value. Completed reports can be exported as CSV files for "
       "offline analysis in spreadsheet applications.")

    _h(doc, "2.2.8.6 Generate Order Report", size=12, underline=True)
    _b(doc,
       "Admin can generate detailed order reports filtered by a custom date range, specific order "
       "status, or product category through the reporting panel. The generated report provides "
       "a summary breakdown of orders segmented by status: Pending, In Production, Quality Check, "
       "Shipped, and Delivered: and lists individual orders with their key details. Reports "
       "are exportable as CSV files for record-keeping and operational planning purposes.")

    _h(doc, "2.2.8.7 View Analytics Dashboard", size=12, underline=True)
    _b(doc,
       "The main admin dashboard page displays a suite of real-time analytical visualizations "
       "rendered using HTML5 Canvas-based charting. The charts present daily new order counts, "
       "weekly revenue trends, product popularity distribution across categories, and chatbot "
       "usage statistics including query volumes and common topics. Data underlying the charts "
       "is fetched from the FastAPI analytics endpoints and refreshed on each dashboard load.")

    _h(doc, "2.2.8.8 Send Notification to Customers", size=12, underline=True)
    _b(doc,
       "Admin can compose and dispatch email notifications to customers through the notification "
       "management panel. Admin can choose to send a notification to all registered customers "
       "(bulk broadcast) or target a specific customer by account. Notifications are sent via "
       "the Gmail SMTP integration and can be used to communicate new product launches, "
       "promotional offers, or important order-related updates.")

    _h(doc, "2.2.8.9 Manage Content (Terms, Privacy, FAQ)", size=12, underline=True)
    _b(doc,
       "Admin can update the platform's Terms & Conditions, Privacy Policy, and FAQ sections "
       "through a rich-text content editor integrated into the admin panel. Changes made through "
       "the editor are saved to the database and immediately reflected on the corresponding "
       "customer-facing pages, allowing the admin to keep legal and informational content current "
       "without requiring developer intervention or a system redeployment.")

    _h(doc, "2.2.8.10 Manage Coupon/Promo Codes", size=12, underline=True)
    _b(doc,
       "Admin can create promotional coupon codes through the discount management section of "
       "the admin panel by specifying the code string, discount type (percentage-based or "
       "fixed-amount), discount value, usage limit (total number of redemptions allowed), and "
       "an optional expiry date. Admin can toggle any coupon code between Active and Inactive "
       "status at any time, and the system enforces the usage limit by invalidating the code "
       "once it has been redeemed the maximum number of times.")

    _h(doc, "2.2.8.11 Export Data as CSV", size=12, underline=True)
    _b(doc,
       "Admin can export key operational datasets as CSV files directly from the admin panel "
       "for use in external analysis tools. The available exports include the complete customer "
       "list with registration and order summary data, the full order dataset with customization "
       "details and statuses, and aggregated sales report data. CSV files are generated "
       "server-side by the FastAPI backend and delivered to the admin's browser as a "
       "file download.")

    # ── 2.3 Non-Functional Requirements ─────────────────────────────────────
    _h(doc, "2.3 Non-Functional Requirements", size=14)
    _b(doc,
       "This section defines the quality attributes and constraints that the AURA-WEAR system "
       "must satisfy. Non-functional requirements govern how the system performs its functions "
       "rather than what functions it performs.")

    _h(doc, "2.3.1 Performance Requirements", size=13)
    _b(doc,
       "The AURA-WEAR web portal shall load any page within 3 seconds on a standard broadband "
       "internet connection (minimum 10 Mbps download speed). The real-time 2D customization "
       "preview panel shall reflect any customer selection change within 500 milliseconds. "
       "The AI chatbot shall return a response within 5 seconds for 95% of queries under "
       "normal operating load, accounting for the external OpenAI API call latency. All "
       "internal FastAPI REST API endpoints shall respond within 1 second for standard CRUD "
       "operations under expected user load. Performance shall be monitored against these "
       "thresholds during load testing prior to deployment.")

    _h(doc, "2.3.2 Usability Requirements", size=13)
    _b(doc,
       "The AURA-WEAR interface shall be fully responsive, supporting screen widths ranging "
       "from 320 pixels (small mobile devices) to 1920 pixels (widescreen desktop displays) "
       "without loss of functionality or visual integrity. The ordering workflow shall be "
       "completable in 5 clicks or fewer from the product catalog page to order confirmation. "
       "All form fields shall display inline validation feedback within 300 milliseconds of "
       "the user completing their input, providing immediate guidance without requiring form "
       "submission. The system shall comply with WCAG 2.1 Level AA accessibility guidelines, "
       "ensuring the platform is usable by individuals with visual and motor disabilities.")

    _h(doc, "2.3.3 Security Requirements", size=13)
    _b(doc,
       "All user passwords shall be hashed using the bcrypt algorithm with a cost factor of "
       "10 or greater before storage; plaintext passwords shall never be persisted in any form. "
       "All database queries shall use parameterized statements (SQLAlchemy ORM or explicit "
       "parameterization) to prevent SQL injection attacks. HTTPS with TLS 1.3 shall be enforced "
       "on all pages handling payment information, and all API endpoints shall be protected by "
       "TLS. JWT tokens shall be validated for signature integrity and expiry on every protected "
       "API request. All user-submitted text input shall be sanitized server-side to prevent "
       "cross-site scripting (XSS) and other injection attacks.")

    _h(doc, "2.3.4 Reliability Requirements", size=13)
    _b(doc,
       "The AURA-WEAR system shall target a minimum uptime of 99% measured on a monthly basis, "
       "equating to no more than approximately 7.3 hours of unplanned downtime per month. All "
       "database write operations affecting order, payment, or user account records shall be "
       "executed within ACID-compliant transactions to ensure data integrity in the event of "
       "partial failures. The system shall implement graceful degradation for the AI chatbot "
       "feature: when the OpenAI API is unavailable, the chatbot shall display a user-friendly "
       "fallback message and continue to serve FAQ responses from the local database. Automated "
       "daily backups of the SQLite database shall be performed and retained for a minimum of "
       "30 days.")

    _h(doc, "2.3.5 Scalability Requirements", size=13)
    _b(doc,
       "The FastAPI backend shall be designed as a stateless application, storing no session "
       "state in memory, to support horizontal scaling through the addition of multiple "
       "application server instances behind a load balancer without architectural changes. "
       "The overall system architecture shall be capable of supporting up to 10,000 concurrent "
       "active users without degradation below the stated performance thresholds. Database "
       "connection pooling shall be implemented to efficiently manage the pool of database "
       "connections shared across concurrent requests, preventing connection exhaustion under "
       "peak load conditions.")

    _h(doc, "2.3.6 Maintainability & Support Requirements", size=13)
    _b(doc,
       "The AURA-WEAR codebase shall follow a Model-View-Controller (MVC) architectural pattern "
       "with a clear separation of concerns between data models, business logic, and presentation "
       "layers. All FastAPI API endpoints and public functions shall include docstrings describing "
       "their purpose, parameters, and return values to facilitate future maintenance. The project "
       "shall use Git for version control with a feature-branch workflow, requiring all new "
       "features and bug fixes to be developed on dedicated branches and merged via pull requests "
       "with peer review. n8n workflow configurations shall be version-controlled by exporting "
       "workflow JSON files to the project repository.")

    _h(doc, "2.3.7 Compatibility Requirements", size=13)
    _b(doc,
       "The AURA-WEAR web application shall be compatible with the two most recent major "
       "versions of the following web browsers: Google Chrome, Mozilla Firefox, Apple Safari, "
       "and Microsoft Edge. The application shall function correctly on the Windows, macOS, "
       "and Linux desktop operating systems. The system requires no browser plugins, extensions, "
       "or native application installation; all functionality shall be delivered through "
       "standard web technologies (HTML5, CSS3, JavaScript) available natively in the "
       "supported browsers.")

    _h(doc, "2.3.8 Support Requirements", size=13)
    _b(doc,
       "The AURA-WEAR admin panel shall be designed to empower the admin user to perform all "
       "routine operational tasks: including product management, order processing, customer "
       "account management, report generation, and content updates: without requiring "
       "developer intervention or access to the server infrastructure. A comprehensive user "
       "manual covering all admin and customer-facing features shall be provided as part of "
       "the system deliverables. All system error messages displayed to end users shall be "
       "written in plain language, clearly describing what went wrong and providing actionable "
       "guidance for resolution, avoiding technical jargon or raw exception messages.")

    # ── 2.4 External Interface Requirements ──────────────────────────────────
    _h(doc, "2.4 External Interface Requirements", size=14)

    _h(doc, "2.4.1 User Interfaces", size=13)
    _b(doc,
       "The AURA-WEAR system shall provide a responsive web interface built with HTML5, CSS3, "
       "and vanilla JavaScript, adopting a clean, minimalist visual design aesthetic that "
       "emphasizes product imagery and the customization experience. The customization page "
       "shall feature a dedicated real-time 2D preview panel occupying the right half of the "
       "viewport on desktop screens. A floating chatbot widget with a toggle button shall be "
       "persistently accessible from every page without obstructing primary page content. "
       "The navigation shall collapse into a hamburger-style mobile menu on viewports narrower "
       "than 768 pixels to maintain usability on smartphone and tablet devices.")

    _h(doc, "2.4.2 Hardware Interfaces", size=13)
    _b(doc,
       "The AURA-WEAR system requires no specialized hardware on the client side beyond a "
       "standard desktop computer, laptop, or tablet device capable of running a supported "
       "web browser. The system has no dependencies on client-side peripherals such as printers "
       "or scanners. Camera or file system access on the client device is used optionally and "
       "only when the customer chooses to upload a custom logo image, which is initiated through "
       "the standard browser file input control and does not require any special device permissions "
       "beyond the browser's standard file access prompt.")

    _h(doc, "2.4.3 Software Interfaces", size=13)
    _b(doc,
       "The AURA-WEAR system interfaces with the following external software systems and APIs: "
       "the FastAPI REST API serves as the primary interface between the frontend and backend; "
       "the OpenAI GPT-4o-mini API is accessed via the n8n workflow automation platform for "
       "AI chatbot response generation; the PayFast payment gateway API and the Stripe payment "
       "API process online card payments; the Gmail SMTP API via the smtplib library handles "
       "all outbound email communication including OTPs, order confirmations, and notifications; "
       "and the Cloudinary image hosting API stores and serves all product and user-uploaded "
       "logo images.")

    _h(doc, "2.4.4 Communication Interfaces", size=13)
    _b(doc,
       "All communication between the client browser and the AURA-WEAR server shall be encrypted "
       "using HTTPS with TLS 1.3 to protect data in transit, including user credentials, payment "
       "information, and personal data. Real-time notifications: such as order status updates "
       "and chatbot response streaming: shall be delivered via WebSocket connections where "
       "supported, falling back to HTTP polling for environments where WebSocket connections "
       "are restricted. The backend exposes a RESTful JSON API adhering to standard HTTP "
       "methods (GET, POST, PUT, DELETE) and status codes. All outbound email notifications "
       "are transmitted via SMTP over port 587 with STARTTLS encryption.")

    doc.add_page_break()
