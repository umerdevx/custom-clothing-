from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _h(doc, text, bold=True, size=13, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def _b(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def _uc_table(doc, uc_data):
    table = doc.add_table(rows=len(uc_data), cols=2, style='Table Grid')
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    for i, (label, value) in enumerate(uc_data):
        row = table.rows[i]
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.name = 'Times New Roman'
        lr.font.size = Pt(11)
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        vr.font.name = 'Times New Roman'
        vr.font.size = Pt(11)
    doc.add_paragraph('')


def _ssd_table(doc, ssd_data):
    table = doc.add_table(rows=1 + len(ssd_data), cols=3, style='Table Grid')
    headers = ['Actor', 'Direction', 'System']
    for i, h in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
    for i, (direction, message) in enumerate(ssd_data):
        row = table.rows[i + 1]
        if direction == '→':
            p = row.cells[0].paragraphs[0]
            r = p.add_run(message)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            row.cells[1].paragraphs[0].add_run('→').font.name = 'Times New Roman'
        else:
            p = row.cells[2].paragraphs[0]
            r = p.add_run(message)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            row.cells[1].paragraphs[0].add_run('←').font.name = 'Times New Roman'
    doc.add_paragraph('')


def _code_block(doc, code_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.bold = False


def _schema_table(doc, table_name, columns):
    _h(doc, f'Table: {table_name}', bold=True, size=12)
    tbl = doc.add_table(rows=1 + len(columns), cols=3, style='Table Grid')
    headers = ['Column Name', 'Data Type', 'Constraints']
    for i, h in enumerate(headers):
        p = tbl.rows[0].cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
    for i, (col, dtype, constraint) in enumerate(columns):
        row = tbl.rows[i + 1]
        for j, val in enumerate([col, dtype, constraint]):
            p = row.cells[j].paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
    doc.add_paragraph('')


def add_part4(doc):
    # =========================================================
    # A) FULLY DRESSED USE CASES UC-29 to UC-56
    # =========================================================

    # UC-29
    _h(doc, 'UC-29: View FAQ via Chatbot', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-29'),
        ('Use Case Name', 'View FAQ via Chatbot'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Chatbot widget is open'),
        ('Post-Conditions', 'FAQ response displayed'),
        ('Main Flow', '1. Customer types FAQ question (e.g., "What is your return policy?").\n2. n8n receives query.\n3. n8n searches FAQ table in database for matching keywords.\n4. System finds matching FAQ entry.\n5. n8n returns the FAQ answer as formatted text.\n6. System displays FAQ answer in chat window.\n7. Related FAQ suggestions shown below the answer.'),
        ('Alternative Flow', 'If no FAQ match, query is passed to GPT for general response.'),
        ('Exception Flow', 'If FAQ table is empty, return generic support email response.'),
    ])

    # UC-30
    _h(doc, 'UC-30: View Chat History', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-30'),
        ('Use Case Name', 'View Chat History'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is signed in; past chat sessions exist'),
        ('Post-Conditions', 'Chat history displayed in chatbot interface'),
        ('Main Flow', '1. Authenticated customer opens the chatbot widget.\n2. System detects user is logged in via JWT.\n3. System queries chat_logs table for user\'s previous conversations.\n4. System loads last 20 conversation turns.\n5. Chat history is displayed in the chat window above the input field.\n6. User can scroll up to view older messages.\n7. Clear History button available.'),
        ('Alternative Flow', 'Anonymous users see empty history (no saving).'),
        ('Exception Flow', 'If history load fails, show empty state with option to retry.'),
    ])

    # UC-31
    _h(doc, 'UC-31: Admin Manage FAQ Entries', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-31'),
        ('Use Case Name', 'Admin Manage FAQ Entries'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in'),
        ('Post-Conditions', 'FAQ database updated'),
        ('Main Flow', '1. Admin navigates to Content Management > FAQ.\n2. System displays list of existing FAQ entries (question + answer pairs).\n3. Admin can Add, Edit, or Delete entries.\n4. Admin clicks Add FAQ.\n5. Admin enters question and answer text.\n6. Admin clicks Save.\n7. System validates fields are not empty.\n8. System saves new FAQ entry to database.\n9. New entry immediately available in chatbot knowledge base.'),
        ('Alternative Flow', 'Admin clicks Edit on existing FAQ and modifies text.'),
        ('Exception Flow', 'If FAQ entry has special characters, system sanitizes before saving.'),
    ])

    doc.add_page_break()

    # UC-32
    _h(doc, 'UC-32: Create New Order', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-32'),
        ('Use Case Name', 'Create New Order'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer signed in; customization complete; item in cart'),
        ('Post-Conditions', 'Order created with unique ID; status set to Pending'),
        ('Main Flow', '1. Customer reviews customization in cart.\n2. Customer clicks Proceed to Checkout.\n3. System displays order summary with customization details and total price.\n4. Customer confirms delivery address.\n5. Customer selects payment method (COD or online).\n6. Customer clicks Place Order.\n7. System generates unique alphanumeric Order ID.\n8. System saves order and order_items records to database.\n9. System sends order confirmation email.\n10. System displays order confirmation page with Order ID.'),
        ('Alternative Flow', 'Customer can edit cart before proceeding to checkout.'),
        ('Exception Flow', 'If payment fails, order is not created and user is notified.'),
    ])

    # UC-33
    _h(doc, 'UC-33: Search Order History', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-33'),
        ('Use Case Name', 'Search Order History'),
        ('Actor(s)', 'Customer / Admin'),
        ('Pre-Conditions', 'User is signed in; orders exist in database'),
        ('Post-Conditions', 'Filtered order list displayed'),
        ('Main Flow', '1. User navigates to Order History page.\n2. User enters search term (Order ID, product name, or date).\n3. System queries orders table with the search criteria.\n4. For customers, results are filtered to their own orders only.\n5. For admin, all orders are searchable.\n6. Matching orders displayed in list with status and date.\n7. User can click an order to view full details.'),
        ('Alternative Flow', 'User can filter by status (Pending, In Production, Shipped, Delivered).'),
        ('Exception Flow', 'If no orders match search, display "No orders found."'),
    ])

    # UC-34
    _h(doc, 'UC-34: View Order History', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-34'),
        ('Use Case Name', 'View Order History'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is signed in; at least one order exists'),
        ('Post-Conditions', 'Order history list displayed with all order details'),
        ('Main Flow', '1. Customer navigates to My Orders page.\n2. System queries orders table for orders linked to the customer ID.\n3. System orders results by date (newest first).\n4. System displays each order with: Order ID, date, product name, status badge, and total.\n5. Customer clicks an order to see full details.\n6. Full details show customization specs (fabric, color, logo, stitching, print, wash).\n7. Reorder button available on completed orders.'),
        ('Alternative Flow', 'If no orders, display "You haven\'t placed any orders yet."'),
        ('Exception Flow', 'If order data is inconsistent, show partial data with warning.'),
    ])

    doc.add_page_break()

    # UC-35
    _h(doc, 'UC-35: Cancel Order', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-35'),
        ('Use Case Name', 'Cancel Order'),
        ('Actor(s)', 'Customer / Admin'),
        ('Pre-Conditions', 'Order exists with Pending status'),
        ('Post-Conditions', 'Order status updated to Cancelled; customer notified'),
        ('Main Flow', '1. Customer navigates to Order History and selects a Pending order.\n2. Customer clicks Cancel Order.\n3. System displays confirmation dialog.\n4. Customer confirms cancellation.\n5. System updates order status to Cancelled in database.\n6. System records cancellation timestamp.\n7. System sends cancellation confirmation email.\n8. Order details page shows Cancelled status badge.'),
        ('Alternative Flow', 'Admin can cancel orders at any status from admin panel.'),
        ('Exception Flow', 'Orders with Shipped or Delivered status cannot be cancelled by customer.'),
    ])

    # UC-36
    _h(doc, 'UC-36: Track Order Status', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-36'),
        ('Use Case Name', 'Track Order Status'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Customer is signed in; order exists'),
        ('Post-Conditions', 'Current order status displayed with progress timeline'),
        ('Main Flow', '1. Customer navigates to Order History and selects an order.\n2. System fetches current order status from database.\n3. System displays a visual status timeline: Pending → In Production → Quality Check → Shipped → Delivered.\n4. Current stage is highlighted.\n5. Estimated dates shown for each stage.\n6. Customer receives push notification when status changes.\n7. Delivered orders show delivery date.'),
        ('Alternative Flow', 'If status is Cancelled, show cancellation date and reason.'),
        ('Exception Flow', 'If status data corrupted, show last known valid status.'),
    ])

    # UC-37
    _h(doc, 'UC-37: Reorder Previous Order', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-37'),
        ('Use Case Name', 'Reorder Previous Order'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Previous order exists with Delivered status'),
        ('Post-Conditions', 'New order created with same customization settings'),
        ('Main Flow', '1. Customer navigates to Order History.\n2. Customer finds a completed order.\n3. Customer clicks Reorder.\n4. System pre-fills the customization form with previous order\'s specifications.\n5. System displays the 2D preview with previous customizations applied.\n6. Customer reviews and can modify any settings.\n7. Customer clicks Place Order.\n8. System creates a new order with the same or modified specifications.'),
        ('Alternative Flow', 'Customer can change quantity when reordering.'),
        ('Exception Flow', 'If a selected fabric or product is no longer available, prompt customer to make alternative selection.'),
    ])

    doc.add_page_break()

    # UC-38
    _h(doc, 'UC-38: Record Online Payment', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-38'),
        ('Use Case Name', 'Record Online Payment'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Order placed; payment method is online (PayFast/Stripe)'),
        ('Post-Conditions', 'Payment recorded; order confirmed; receipt generated'),
        ('Main Flow', '1. Customer selects online payment at checkout.\n2. System displays payment method options (PayFast, Stripe).\n3. Customer selects preferred gateway.\n4. System redirects to payment gateway\'s secure page.\n5. Customer enters card details on gateway page.\n6. Gateway processes payment and sends callback to AURA-WEAR API.\n7. System receives payment confirmation (transaction ID, amount, status).\n8. System records payment in payments table.\n9. System updates order status.\n10. System generates digital receipt.'),
        ('Alternative Flow', 'Customer can switch to COD if online payment fails.'),
        ('Exception Flow', 'If payment gateway timeout, display "Payment processing. Please do not refresh." and check status.'),
    ])

    # UC-39
    _h(doc, 'UC-39: Record Cash on Delivery', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-39'),
        ('Use Case Name', 'Record Cash on Delivery'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Order placed; COD selected as payment method'),
        ('Post-Conditions', 'Order saved with COD status; reference number generated'),
        ('Main Flow', '1. Customer selects Cash on Delivery at checkout.\n2. System validates delivery address is complete.\n3. Customer confirms order.\n4. System creates order with payment_method=COD and payment_status=Pending.\n5. System generates a unique COD reference number.\n6. System sends order confirmation email with COD reference.\n7. Delivery agent uses COD reference to collect payment.\n8. Admin marks payment as Received upon delivery confirmation.'),
        ('Alternative Flow', 'Admin can update COD order payment status from admin panel.'),
        ('Exception Flow', 'If delivery address is in restricted area, display "COD not available for your location."'),
    ])

    # UC-40
    _h(doc, 'UC-40: Generate Digital Receipt', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-40'),
        ('Use Case Name', 'Generate Digital Receipt'),
        ('Actor(s)', 'System (automatic)'),
        ('Pre-Conditions', 'Order confirmed (online payment or COD)'),
        ('Post-Conditions', 'Digital receipt created; accessible by customer'),
        ('Main Flow', '1. Order placement is confirmed.\n2. System gathers order details (ID, date, items, customization, pricing).\n3. System calculates item totals and grand total.\n4. System generates receipt with unique receipt number.\n5. Receipt includes payment method, delivery address, and estimated delivery date.\n6. System stores receipt record in database.\n7. System sends receipt to customer email.\n8. Receipt accessible from Order Details page.'),
        ('Alternative Flow', 'Customer can regenerate receipt from order detail page.'),
        ('Exception Flow', 'If email sending fails, receipt is still saved and accessible from the website.'),
    ])

    doc.add_page_break()

    # UC-41
    _h(doc, 'UC-41: View All Payments and Transactions', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-41'),
        ('Use Case Name', 'View All Payments and Transactions'),
        ('Actor(s)', 'Customer / Admin'),
        ('Pre-Conditions', 'User is signed in; transactions exist'),
        ('Post-Conditions', 'Payment history displayed'),
        ('Main Flow', '1. Customer navigates to Payment History.\n2. System queries payments table filtered by customer ID.\n3. System displays list of all transactions with: receipt number, date, amount, payment method, status.\n4. Admin can view all customers\' transactions without filter.\n5. Admin can filter by date range, payment method, or status.\n6. Clicking a transaction shows full order details.'),
        ('Alternative Flow', 'Admin can export transaction data as CSV.'),
        ('Exception Flow', 'If no transactions found, display "No payment history available."'),
    ])

    # UC-42
    _h(doc, 'UC-42: Download Invoice', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-42'),
        ('Use Case Name', 'Download Invoice'),
        ('Actor(s)', 'Customer'),
        ('Pre-Conditions', 'Order has confirmed payment; invoice exists'),
        ('Post-Conditions', 'Invoice PDF downloaded to customer\'s device'),
        ('Main Flow', '1. Customer opens Order Details page.\n2. Customer clicks Download Invoice.\n3. System retrieves order and payment data.\n4. System generates formatted PDF invoice with: company logo, order details, customization summary, pricing breakdown, tax (if any), total, and payment confirmation.\n5. System returns the PDF file to the browser.\n6. Browser prompts download dialog.\n7. Invoice saved as "AURA-WEAR-Invoice-[OrderID].pdf".'),
        ('Alternative Flow', 'Customer can also receive invoice via email by clicking Email Invoice.'),
        ('Exception Flow', 'If PDF generation fails, display "Invoice generation failed. Please contact support."'),
    ])

    # UC-43
    _h(doc, 'UC-43: Manage Products - Admin', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-43'),
        ('Use Case Name', 'Manage Products - Admin'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in'),
        ('Post-Conditions', 'Product catalog updated as per admin action'),
        ('Main Flow', '1. Admin navigates to Product Management.\n2. System displays paginated product list with search and filter.\n3. Admin performs desired action: Add, Edit, Enable/Disable, or Delete.\n4. For Add: Admin fills product form and submits.\n5. For Edit: Admin modifies fields and saves.\n6. For Enable/Disable: Admin toggles status.\n7. System validates and saves changes.\n8. Success notification displayed.'),
        ('Alternative Flow', 'Admin can use bulk actions to enable/disable multiple products.'),
        ('Exception Flow', 'Cannot delete product with existing order history: soft delete only.'),
    ])

    doc.add_page_break()

    # UC-44
    _h(doc, 'UC-44: Manage Customer Accounts - Admin', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-44'),
        ('Use Case Name', 'Manage Customer Accounts - Admin'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; customers are registered'),
        ('Post-Conditions', 'Customer account status updated or information viewed'),
        ('Main Flow', '1. Admin navigates to Customer Management.\n2. System displays paginated customer list with search.\n3. Admin can view customer profile, order history, and account status.\n4. Admin selects a customer and clicks Activate or Deactivate.\n5. System confirms action.\n6. System updates account status.\n7. If deactivated, customer\'s active sessions are terminated.\n8. Admin action is logged.'),
        ('Alternative Flow', 'Admin can search customers by name, email, or registration date.'),
        ('Exception Flow', 'Admin cannot deactivate their own account.'),
    ])

    # UC-45
    _h(doc, 'UC-45: View and Update Order Management - Admin', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-45'),
        ('Use Case Name', 'View and Update Order Management - Admin'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; orders exist'),
        ('Post-Conditions', 'Order status updated; customer notified of status change'),
        ('Main Flow', '1. Admin navigates to Order Management.\n2. System displays all orders sorted by newest first.\n3. Admin can filter by status, date, or customer.\n4. Admin selects an order to view details including customization specs.\n5. Admin clicks Update Status.\n6. Admin selects new status from dropdown.\n7. System updates order status in database.\n8. System sends status update notification email to customer.'),
        ('Alternative Flow', 'Admin can add notes to an order for internal tracking.'),
        ('Exception Flow', 'Orders cannot be moved backwards in status (e.g., from Shipped to Pending).'),
    ])

    # UC-46
    _h(doc, 'UC-46: View Inventory and Stock Alerts - Admin', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-46'),
        ('Use Case Name', 'View Inventory and Stock Alerts - Admin'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; inventory records exist'),
        ('Post-Conditions', 'Inventory status viewed; low-stock items identified'),
        ('Main Flow', '1. Admin navigates to Inventory Management.\n2. System displays inventory table: material name, quantity, unit, last updated.\n3. Items below threshold (configurable, default 10 units) are highlighted in red.\n4. Admin updates quantity when stock arrives.\n5. Admin sets alert thresholds per material.\n6. Dashboard shows low-stock count badge.\n7. Admin can export inventory report.'),
        ('Alternative Flow', 'Admin can set custom thresholds for different materials.'),
        ('Exception Flow', 'If inventory table is empty, show prompt to add initial inventory.'),
    ])

    doc.add_page_break()

    # UC-47
    _h(doc, 'UC-47: Generate Sales Report', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-47'),
        ('Use Case Name', 'Generate Sales Report'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; order and payment data exists'),
        ('Post-Conditions', 'Sales report generated and displayed; CSV export available'),
        ('Main Flow', '1. Admin navigates to Reports > Sales Report.\n2. Admin selects date range (monthly, custom).\n3. System queries payments and orders tables for the period.\n4. System calculates: total revenue, number of orders, average order value, top products.\n5. System displays report as table and chart.\n6. Admin can export report as CSV.\n7. Report data includes comparison with previous period.'),
        ('Alternative Flow', 'Admin can select annual view to compare monthly performance.'),
        ('Exception Flow', 'If no data for selected period, display "No sales data available for this period."'),
    ])

    # UC-48
    _h(doc, 'UC-48: Generate Order Report', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-48'),
        ('Use Case Name', 'Generate Order Report'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; orders exist'),
        ('Post-Conditions', 'Order report generated with status breakdown'),
        ('Main Flow', '1. Admin navigates to Reports > Order Report.\n2. Admin selects date range and optional status filter.\n3. System queries orders table for the period.\n4. System categorizes orders by status: Pending, In Production, Quality Check, Shipped, Delivered, Cancelled.\n5. System displays counts and percentages for each status.\n6. Chart shows order trend over the selected period.\n7. Admin exports as CSV.'),
        ('Alternative Flow', 'Admin can filter by specific product to see its order history.'),
        ('Exception Flow', 'If database query times out, display partial data with reload option.'),
    ])

    # UC-49
    _h(doc, 'UC-49: View Analytics Dashboard', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-49'),
        ('Use Case Name', 'View Analytics Dashboard'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in'),
        ('Post-Conditions', 'Analytics dashboard displayed with real-time charts'),
        ('Main Flow', '1. Admin navigates to the Admin Dashboard.\n2. System displays summary widgets: total orders today, weekly revenue, total customers, pending orders.\n3. System renders a Canvas chart of daily orders for the past 30 days.\n4. System renders a revenue trend chart.\n5. System shows most popular products by order count.\n6. System shows chatbot usage statistics (queries per day).\n7. Dashboard auto-refreshes every 5 minutes.'),
        ('Alternative Flow', 'Admin can toggle chart period (7 days, 30 days, 90 days).'),
        ('Exception Flow', 'If chart data fails to load, display last cached data with timestamp.'),
    ])

    doc.add_page_break()

    # UC-50
    _h(doc, 'UC-50: Send Notification to Customers', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-50'),
        ('Use Case Name', 'Send Notification to Customers'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; customer email list available'),
        ('Post-Conditions', 'Notification emails sent to selected customers'),
        ('Main Flow', '1. Admin navigates to Notifications.\n2. Admin selects notification type: All Customers, Specific Customer, or Active Customers.\n3. Admin enters notification subject and message body.\n4. Admin clicks Preview to see the email template.\n5. Admin clicks Send.\n6. System queues emails via FastAPI-Mail.\n7. System sends emails via Gmail SMTP.\n8. System records notification in notification_logs table.'),
        ('Alternative Flow', 'Admin can schedule notification for a future date/time.'),
        ('Exception Flow', 'If email service unavailable, queue notifications and retry after 1 hour.'),
    ])

    # UC-51
    _h(doc, 'UC-51: Manage Content - Terms, Privacy, FAQ', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-51'),
        ('Use Case Name', 'Manage Content - Terms, Privacy, FAQ'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in'),
        ('Post-Conditions', 'Content updated and immediately visible to customers'),
        ('Main Flow', '1. Admin navigates to Content Management.\n2. System displays tabs: Terms & Conditions, Privacy Policy, FAQ.\n3. Admin clicks on the tab to edit.\n4. System displays current content in a text editor.\n5. Admin modifies content.\n6. Admin clicks Save Changes.\n7. System updates content in the database with timestamp.\n8. Content is immediately visible on the customer-facing pages.'),
        ('Alternative Flow', 'Admin can preview changes before saving.'),
        ('Exception Flow', 'If content exceeds max length, display "Content too long. Maximum 10,000 characters."'),
    ])

    # UC-52
    _h(doc, 'UC-52: Manage Coupon/Promo Codes', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-52'),
        ('Use Case Name', 'Manage Coupon/Promo Codes'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in'),
        ('Post-Conditions', 'Coupon code created, activated, or deactivated'),
        ('Main Flow', '1. Admin navigates to Promotions > Coupon Codes.\n2. System displays list of existing coupon codes.\n3. Admin clicks Add New Coupon.\n4. Admin enters: code, discount type (%, fixed), discount value, expiry date, usage limit.\n5. Admin clicks Save.\n6. System validates the coupon code is unique.\n7. System saves coupon to database with Active status.\n8. Customers can apply the code at checkout.'),
        ('Alternative Flow', 'Admin can deactivate a coupon without deleting it.'),
        ('Exception Flow', 'If coupon code already exists, display "This code is already in use."'),
    ])

    doc.add_page_break()

    # UC-53
    _h(doc, 'UC-53: Export Data as CSV', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-53'),
        ('Use Case Name', 'Export Data as CSV'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; data exists to export'),
        ('Post-Conditions', 'CSV file downloaded to admin\'s device'),
        ('Main Flow', '1. Admin navigates to Reports or Customer Management.\n2. Admin clicks Export as CSV.\n3. System queries relevant database table.\n4. System formats data as CSV with column headers.\n5. System generates the file server-side.\n6. Browser initiates file download.\n7. File saved as "[Report-Type]-[Date].csv".'),
        ('Alternative Flow', 'Admin can apply filters before export to export subset of data.'),
        ('Exception Flow', 'If data set is too large (>50,000 rows), system emails CSV link instead of direct download.'),
    ])

    # UC-54
    _h(doc, 'UC-54: Manage Admin Users', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-54'),
        ('Use Case Name', 'Manage Admin Users'),
        ('Actor(s)', 'Admin (Super Admin)'),
        ('Pre-Conditions', 'Super Admin is signed in'),
        ('Post-Conditions', 'Admin user account created or modified'),
        ('Main Flow', '1. Super Admin navigates to Admin Management.\n2. System displays list of admin users.\n3. Super Admin clicks Add Admin.\n4. Super Admin enters new admin\'s name and email.\n5. System generates a temporary password and sends it to the new admin\'s email.\n6. System creates admin account with Admin role.\n7. New admin must change temporary password on first login.'),
        ('Alternative Flow', 'Super Admin can deactivate an admin account.'),
        ('Exception Flow', 'Cannot delete the last remaining super admin account.'),
    ])

    # UC-55
    _h(doc, 'UC-55: View Customer Reviews', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-55'),
        ('Use Case Name', 'View Customer Reviews'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; reviews exist'),
        ('Post-Conditions', 'Customer reviews displayed; admin can flag or respond'),
        ('Main Flow', '1. Admin navigates to Reviews & Ratings section.\n2. System displays all customer reviews with: customer name, product, rating (1-5 stars), review text, date.\n3. Admin filters by product or rating.\n4. Admin can flag inappropriate reviews.\n5. Flagged reviews are hidden from customer-facing pages.\n6. Admin can respond to reviews on behalf of the company.\n7. System records admin response and displays it below the review.'),
        ('Alternative Flow', 'Admin can export review data for quality analysis.'),
        ('Exception Flow', 'If review text contains prohibited content, auto-flag using keyword filter.'),
    ])

    doc.add_page_break()

    # UC-56
    _h(doc, 'UC-56: Generate Provider Performance Report', bold=True, size=12)
    _uc_table(doc, [
        ('Use Case ID', 'UC-56'),
        ('Use Case Name', 'Generate Provider Performance Report'),
        ('Actor(s)', 'Admin'),
        ('Pre-Conditions', 'Admin is signed in; order data exists'),
        ('Post-Conditions', 'Performance report generated showing product and sales metrics'),
        ('Main Flow', '1. Admin navigates to Reports > Performance Report.\n2. Admin selects date range.\n3. System analyzes order data to calculate performance metrics.\n4. System generates report including: top-selling products, revenue by product category, customization preference trends (most popular fabric, color, print method), and order completion time average.\n5. Charts visualize trends.\n6. Report exported as CSV or PDF.\n7. Report saved to reports_log for future reference.'),
        ('Alternative Flow', 'Admin can schedule monthly automated performance reports via email.'),
        ('Exception Flow', 'If insufficient data for meaningful analysis, display minimum 10-order threshold warning.'),
    ])

    doc.add_page_break()

    # =========================================================
    # B) SECTION 3.3 DOMAIN MODEL
    # =========================================================
    _h(doc, '3.3 Domain Model', bold=True, size=14)
    _b(doc, 'The Domain Model identifies the key entities in the AURA-WEAR system and their relationships. The model captures the business objects and their associations without specifying implementation details.')

    domain_table = doc.add_table(rows=1 + 9, cols=3, style='Table Grid')
    domain_headers = ['Entity', 'Attributes', 'Relationships']
    for i, h in enumerate(domain_headers):
        p = domain_table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

    domain_rows = [
        ('User', 'user_id, name, email, password_hash, role, phone, address, status, created_at', 'Places Orders, Has Profile, Sends Chat Messages'),
        ('Product', 'product_id, name, category, description, base_price, status, created_at', 'Has Customization Options, Appears in Order Items'),
        ('Order', 'order_id, customer_id, total_price, status, payment_method, created_at', 'Belongs to User, Contains Order Items, Has Payment'),
        ('Order_Item', 'item_id, order_id, product_id, fabric, color, logo_url, stitch, print_method, wash, price', 'Part of Order, References Product'),
        ('Payment', 'payment_id, order_id, amount, method, status, transaction_id, created_at', 'Belongs to Order'),
        ('Chat_Log', 'log_id, user_id, query, response, timestamp', 'Belongs to User'),
        ('Inventory', 'inventory_id, material_name, quantity, unit, threshold, updated_at', 'Managed by Admin'),
        ('Coupon', 'coupon_id, code, discount_type, value, expiry, usage_limit, used_count, status', 'Applied to Order'),
        ('FAQ', 'faq_id, question, answer, created_at', 'Used by Chatbot'),
    ]
    for i, (entity, attrs, rels) in enumerate(domain_rows):
        row = domain_table.rows[i + 1]
        for j, val in enumerate([entity, attrs, rels]):
            p = row.cells[j].paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

    doc.add_paragraph('')
    doc.add_page_break()

    # =========================================================
    # C) SECTION 3.4 SYSTEM SEQUENCE DIAGRAMS
    # =========================================================
    _h(doc, '3.4 System Sequence Diagrams', bold=True, size=14)
    _b(doc, 'System Sequence Diagrams describe the interaction between the Actor and the System for each use case, showing the sequence of messages exchanged.')

    # SSD 1-28 (previously defined use cases)
    _h(doc, 'SSD:1 Process SignUp', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Enter registration form data'),
        ('←', 'Validate input fields'),
        ('→', 'Submit registration'),
        ('←', 'Hash password and save'),
        ('←', 'Return success confirmation'),
    ])

    _h(doc, 'SSD:2 Process SignIn', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Enter email and password'),
        ('←', 'Validate credentials against database'),
        ('→', 'Submit login'),
        ('←', 'Generate JWT token'),
        ('←', 'Return token and redirect'),
    ])

    _h(doc, 'SSD:3 Forgot Password', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Enter registered email'),
        ('←', 'Verify email exists in database'),
        ('←', 'Generate 6-digit OTP'),
        ('←', 'Send OTP via Gmail SMTP'),
        ('←', 'Display OTP sent confirmation'),
    ])

    _h(doc, 'SSD:4 Reset Password', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Enter OTP and new password'),
        ('←', 'Validate OTP against stored value'),
        ('←', 'Check OTP expiry (10 minutes)'),
        ('→', 'Submit new password'),
        ('←', 'Hash new password and update database'),
        ('←', 'Return password reset confirmation'),
    ])

    _h(doc, 'SSD:5 Update Profile', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to profile page'),
        ('←', 'Retrieve current profile data from database'),
        ('→', 'Edit profile fields and submit'),
        ('←', 'Validate updated fields'),
        ('←', 'Save changes to database'),
        ('←', 'Return profile update success message'),
    ])

    _h(doc, 'SSD:6 Upload Profile Picture', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Select profile image file'),
        ('←', 'Validate file type and size'),
        ('→', 'Upload image'),
        ('←', 'Save image to server and generate URL'),
        ('←', 'Update profile_pic_url in database'),
        ('←', 'Display updated profile picture'),
    ])

    _h(doc, 'SSD:7 View Products', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to products page'),
        ('←', 'Query active products from database'),
        ('←', 'Return product list with images and base prices'),
        ('→', 'Select a product to view details'),
        ('←', 'Retrieve product details and available options'),
        ('←', 'Display product detail page'),
    ])

    _h(doc, 'SSD:8 Search Products', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Enter search keyword'),
        ('←', 'Query products with LIKE filter on name and description'),
        ('←', 'Return matching active products'),
        ('→', 'Apply category or price filter'),
        ('←', 'Re-query with additional filters'),
        ('←', 'Display filtered results'),
    ])

    _h(doc, 'SSD:9 Customize Product', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Select product and open customizer'),
        ('←', 'Load customization options (fabric, color, print, wash, stitch)'),
        ('→', 'Select fabric type and color'),
        ('←', 'Update 2D preview canvas in real time'),
        ('→', 'Upload logo and select placement'),
        ('←', 'Render logo on preview and update price'),
    ])

    _h(doc, 'SSD:10 Upload Logo', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Select logo image from device'),
        ('←', 'Validate file type (PNG, JPG, SVG) and size'),
        ('→', 'Upload logo file'),
        ('←', 'Save logo to server and return URL'),
        ('←', 'Render logo on 2D canvas at selected placement'),
    ])

    doc.add_page_break()

    _h(doc, 'SSD:11 Generate AI Design Suggestion', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Enter design description prompt in chat'),
        ('←', 'Send prompt to n8n webhook'),
        ('←', 'n8n forwards request to Gemini/GPT API'),
        ('←', 'AI generates design recommendation'),
        ('←', 'System applies suggestion to customizer preview'),
        ('←', 'Display AI-suggested design with explanation'),
    ])

    _h(doc, 'SSD:12 View 2D Real-Time Preview', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Modify any customization option'),
        ('←', 'Capture change event via JavaScript listener'),
        ('←', 'Redraw HTML5 canvas with updated properties'),
        ('←', 'Apply wash effect and logo overlay'),
        ('←', 'Recalculate and display updated price'),
    ])

    _h(doc, 'SSD:13 Save Customization to Cart', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Click Add to Cart'),
        ('←', 'Validate customization completeness'),
        ('←', 'Serialize customization data to session/localStorage'),
        ('←', 'Update cart item count in navigation'),
        ('←', 'Display cart added confirmation message'),
    ])

    _h(doc, 'SSD:14 View Cart', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to cart page'),
        ('←', 'Retrieve cart items from session/localStorage'),
        ('←', 'Display cart items with customization thumbnails'),
        ('←', 'Show quantity, unit price, and total'),
        ('→', 'Update quantity or remove item'),
        ('←', 'Recalculate cart total and refresh display'),
    ])

    _h(doc, 'SSD:15 Update Cart', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Change quantity of cart item'),
        ('←', 'Validate new quantity is positive integer'),
        ('←', 'Update quantity in cart data'),
        ('←', 'Recalculate item and grand total'),
        ('←', 'Refresh cart display with updated values'),
    ])

    _h(doc, 'SSD:16 Apply Coupon Code', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Enter coupon code at checkout'),
        ('←', 'Query coupons table for matching active code'),
        ('←', 'Verify expiry date and usage limit'),
        ('←', 'Calculate discount amount'),
        ('←', 'Display updated total with discount applied'),
    ])

    _h(doc, 'SSD:17 Interact with AI Chatbot', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Open chatbot widget and type message'),
        ('←', 'Send message to n8n webhook endpoint'),
        ('←', 'n8n processes via AI workflow'),
        ('←', 'AI generates contextual response'),
        ('←', 'Display response in chat window with markdown formatting'),
    ])

    _h(doc, 'SSD:18 Get Product Recommendation', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Request product recommendation via chatbot'),
        ('←', 'n8n queries product database for active products'),
        ('←', 'AI analyzes customer preferences from chat context'),
        ('←', 'Generate personalized product recommendation'),
        ('←', 'Display recommendation with product links in chat'),
    ])

    _h(doc, 'SSD:19 Get Customization Assistance', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Ask customization question in chatbot'),
        ('←', 'n8n identifies query as customization-related'),
        ('←', 'AI retrieves relevant customization knowledge'),
        ('←', 'Generate step-by-step customization guidance'),
        ('←', 'Display guidance with option to open customizer'),
    ])

    _h(doc, 'SSD:20 Get Pricing Information', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Ask pricing question in chatbot'),
        ('←', 'n8n identifies pricing query'),
        ('←', 'System queries product base prices from database'),
        ('←', 'AI calculates price range based on customization options'),
        ('←', 'Display pricing breakdown in chat'),
    ])

    doc.add_page_break()

    _h(doc, 'SSD:21 Submit Review', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to product page and click Write Review'),
        ('←', 'Verify customer has purchased the product'),
        ('→', 'Enter rating and review text and submit'),
        ('←', 'Validate review content and rating range'),
        ('←', 'Save review to database'),
        ('←', 'Display review under product with confirmation'),
    ])

    _h(doc, 'SSD:22 View Reviews', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to product detail page'),
        ('←', 'Query reviews table for product reviews'),
        ('←', 'Sort reviews by newest and highest rating'),
        ('←', 'Display star ratings and review text'),
        ('←', 'Show average rating and total review count'),
    ])

    _h(doc, 'SSD:23 Manage Wishlist', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Click wishlist icon on product'),
        ('←', 'Verify customer is authenticated'),
        ('←', 'Toggle product in wishlist (add or remove)'),
        ('←', 'Update wishlist icon state on page'),
        ('←', 'Display wishlist update confirmation'),
    ])

    _h(doc, 'SSD:24 Add Shipping Address', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to profile > addresses and click Add'),
        ('←', 'Display address entry form'),
        ('→', 'Enter address details and save'),
        ('←', 'Validate address fields are complete'),
        ('←', 'Save address to user profile in database'),
        ('←', 'Display saved address in address list'),
    ])

    _h(doc, 'SSD:25 Select Shipping Address at Checkout', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Proceed to checkout'),
        ('←', 'Retrieve saved addresses from user profile'),
        ('→', 'Select delivery address from list'),
        ('←', 'Set selected address as delivery address for order'),
        ('←', 'Display selected address in order summary'),
    ])

    _h(doc, 'SSD:26 Receive Order Confirmation Email', bold=True, size=11)
    _ssd_table(doc, [
        ('←', 'Order created successfully in database'),
        ('←', 'System compiles order confirmation data'),
        ('←', 'System calls FastAPI-Mail with order details'),
        ('←', 'Email sent via Gmail SMTP to customer'),
        ('←', 'Email delivery confirmation logged'),
    ])

    _h(doc, 'SSD:27 Subscribe to Newsletter', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Enter email in newsletter subscription form'),
        ('←', 'Validate email format'),
        ('←', 'Check for duplicate subscription'),
        ('←', 'Save email to subscriptions table'),
        ('←', 'Send welcome email via Gmail SMTP'),
        ('←', 'Display subscription confirmation'),
    ])

    _h(doc, 'SSD:28 Contact Support', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Open contact form and enter name, email, and message'),
        ('←', 'Validate all required fields'),
        ('→', 'Submit contact form'),
        ('←', 'Save support request to database'),
        ('←', 'Send auto-reply email to customer'),
        ('←', 'Display submission confirmation message'),
    ])

    doc.add_page_break()

    _h(doc, 'SSD:29 View FAQ via Chatbot', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Type FAQ question in chatbot widget'),
        ('←', 'Send query to n8n webhook'),
        ('←', 'n8n searches FAQ table for keyword match'),
        ('←', 'Return matched FAQ answer'),
        ('←', 'Display FAQ answer with related suggestions'),
    ])

    _h(doc, 'SSD:30 View Chat History', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Open chatbot widget while logged in'),
        ('←', 'Detect JWT token and extract user ID'),
        ('←', 'Query chat_logs for last 20 conversations'),
        ('←', 'Return conversation history'),
        ('←', 'Display history above input field in chat window'),
    ])

    _h(doc, 'SSD:31 Admin Manage FAQ Entries', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Content Management > FAQ'),
        ('←', 'Retrieve existing FAQ entries from database'),
        ('→', 'Click Add FAQ and enter question and answer'),
        ('←', 'Validate fields are not empty'),
        ('←', 'Save new FAQ entry to database'),
        ('←', 'Confirm entry added and update FAQ list'),
    ])

    _h(doc, 'SSD:32 Create New Order', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Click Proceed to Checkout from cart'),
        ('←', 'Display order summary and address confirmation'),
        ('→', 'Select payment method and click Place Order'),
        ('←', 'Generate unique Order ID'),
        ('←', 'Save order and order_items to database'),
        ('←', 'Display order confirmation with Order ID'),
    ])

    _h(doc, 'SSD:33 Search Order History', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Enter search term in Order History search box'),
        ('←', 'Query orders table with search criteria'),
        ('←', 'Filter results by customer ID for non-admin'),
        ('←', 'Return matching orders list'),
        ('←', 'Display filtered order list with status and date'),
    ])

    _h(doc, 'SSD:34 View Order History', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to My Orders page'),
        ('←', 'Query orders table filtered by customer ID'),
        ('←', 'Sort results by date descending'),
        ('←', 'Return order list with status badges'),
        ('→', 'Click order to view details'),
        ('←', 'Display full order details including customization specs'),
    ])

    _h(doc, 'SSD:35 Cancel Order', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Select Pending order and click Cancel Order'),
        ('←', 'Display cancellation confirmation dialog'),
        ('→', 'Confirm cancellation'),
        ('←', 'Update order status to Cancelled in database'),
        ('←', 'Send cancellation confirmation email'),
        ('←', 'Show Cancelled status badge on order page'),
    ])

    _h(doc, 'SSD:36 Track Order Status', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Select order and click Track Order'),
        ('←', 'Fetch current order status from database'),
        ('←', 'Generate visual status timeline'),
        ('←', 'Highlight current stage in timeline'),
        ('←', 'Display estimated dates for each stage'),
    ])

    doc.add_page_break()

    _h(doc, 'SSD:37 Reorder Previous Order', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Click Reorder on completed order'),
        ('←', 'Retrieve order specifications from database'),
        ('←', 'Pre-fill customization form with previous specs'),
        ('←', 'Render 2D preview with previous customizations'),
        ('→', 'Confirm or modify and click Place Order'),
        ('←', 'Create new order with specifications'),
    ])

    _h(doc, 'SSD:38 Record Online Payment', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Select online payment and click Pay Now'),
        ('←', 'Redirect to payment gateway secure page'),
        ('→', 'Enter card details and confirm payment'),
        ('←', 'Gateway sends payment callback to API'),
        ('←', 'Record payment in payments table'),
        ('←', 'Display payment confirmation and generate receipt'),
    ])

    _h(doc, 'SSD:39 Record Cash on Delivery', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Select Cash on Delivery at checkout'),
        ('←', 'Validate delivery address is complete'),
        ('→', 'Confirm order'),
        ('←', 'Create order with COD payment status Pending'),
        ('←', 'Generate COD reference number'),
        ('←', 'Send order confirmation with COD reference via email'),
    ])

    _h(doc, 'SSD:40 Generate Digital Receipt', bold=True, size=11)
    _ssd_table(doc, [
        ('←', 'Order confirmed in system'),
        ('←', 'Collect order details and pricing data'),
        ('←', 'Generate receipt with unique receipt number'),
        ('←', 'Store receipt in database'),
        ('←', 'Send receipt to customer email'),
        ('←', 'Make receipt accessible on Order Details page'),
    ])

    _h(doc, 'SSD:41 View All Payments and Transactions', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Payment History'),
        ('←', 'Query payments table filtered by customer ID'),
        ('←', 'Return transaction list'),
        ('←', 'Display transactions with amount, method, and status'),
        ('→', 'Click transaction for full order details'),
        ('←', 'Display linked order details'),
    ])

    _h(doc, 'SSD:42 Download Invoice', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Click Download Invoice on Order Details page'),
        ('←', 'Retrieve order and payment data from database'),
        ('←', 'Generate formatted PDF with customization summary'),
        ('←', 'Return PDF file to browser'),
        ('←', 'Browser initiates download as AURA-WEAR-Invoice-[ID].pdf'),
    ])

    _h(doc, 'SSD:43 Manage Products - Admin', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Product Management'),
        ('←', 'Display paginated product list'),
        ('→', 'Click Add Product and fill form'),
        ('←', 'Validate product fields'),
        ('←', 'Save product to database'),
        ('←', 'Display success notification and updated product list'),
    ])

    _h(doc, 'SSD:44 Manage Customer Accounts - Admin', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Customer Management'),
        ('←', 'Display paginated customer list'),
        ('→', 'Select customer and click Deactivate'),
        ('←', 'Confirm deactivation action'),
        ('←', 'Update account status and terminate sessions'),
        ('←', 'Log admin action and display confirmation'),
    ])

    doc.add_page_break()

    _h(doc, 'SSD:45 View and Update Order Management - Admin', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Order Management'),
        ('←', 'Display all orders sorted by newest first'),
        ('→', 'Select order and click Update Status'),
        ('←', 'Display status dropdown options'),
        ('→', 'Select new status and confirm'),
        ('←', 'Update status in database and notify customer via email'),
    ])

    _h(doc, 'SSD:46 View Inventory and Stock Alerts - Admin', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Inventory Management'),
        ('←', 'Query inventory table and compare against thresholds'),
        ('←', 'Display inventory list with low-stock items highlighted'),
        ('→', 'Update quantity for a material'),
        ('←', 'Save updated quantity to database'),
        ('←', 'Refresh inventory display and update dashboard badge'),
    ])

    _h(doc, 'SSD:47 Generate Sales Report', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Reports > Sales Report and select date range'),
        ('←', 'Query payments and orders for selected period'),
        ('←', 'Calculate total revenue and average order value'),
        ('←', 'Generate report table and chart'),
        ('→', 'Click Export as CSV'),
        ('←', 'Generate and download CSV file'),
    ])

    _h(doc, 'SSD:48 Generate Order Report', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Reports > Order Report and select filters'),
        ('←', 'Query orders table for selected period and status'),
        ('←', 'Categorize orders by status and count'),
        ('←', 'Display counts and percentages per status'),
        ('←', 'Render order trend chart'),
        ('→', 'Click Export as CSV'),
    ])

    _h(doc, 'SSD:49 View Analytics Dashboard', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Admin Dashboard'),
        ('←', 'Query aggregate stats from orders, payments, and users tables'),
        ('←', 'Return dashboard summary widgets data'),
        ('←', 'Render daily order chart for past 30 days'),
        ('←', 'Display chatbot usage statistics'),
        ('←', 'Schedule auto-refresh every 5 minutes'),
    ])

    _h(doc, 'SSD:50 Send Notification to Customers', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Notifications and compose message'),
        ('←', 'Display email preview'),
        ('→', 'Click Send'),
        ('←', 'Queue emails via FastAPI-Mail'),
        ('←', 'Send emails via Gmail SMTP'),
        ('←', 'Record notification in logs and display delivery status'),
    ])

    doc.add_page_break()

    _h(doc, 'SSD:51 Manage Content - Terms, Privacy, FAQ', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Content Management and select tab'),
        ('←', 'Retrieve current content from database'),
        ('→', 'Modify content and click Save Changes'),
        ('←', 'Validate content length and format'),
        ('←', 'Update content in database with timestamp'),
        ('←', 'Confirm update and make visible on customer pages'),
    ])

    _h(doc, 'SSD:52 Manage Coupon/Promo Codes', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Promotions > Coupon Codes and click Add'),
        ('←', 'Display coupon creation form'),
        ('→', 'Enter coupon details and click Save'),
        ('←', 'Validate code uniqueness'),
        ('←', 'Save coupon to database with Active status'),
        ('←', 'Display new coupon in coupon list'),
    ])

    _h(doc, 'SSD:53 Export Data as CSV', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Click Export as CSV on Reports or Customer Management page'),
        ('←', 'Query relevant database table'),
        ('←', 'Format data as CSV with column headers'),
        ('←', 'Generate file server-side'),
        ('←', 'Browser initiates file download'),
    ])

    _h(doc, 'SSD:54 Manage Admin Users', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Admin Management and click Add Admin'),
        ('←', 'Display admin creation form'),
        ('→', 'Enter name and email and confirm'),
        ('←', 'Generate temporary password'),
        ('←', 'Create admin account and send credentials via email'),
        ('←', 'Display new admin in admin list'),
    ])

    _h(doc, 'SSD:55 View Customer Reviews', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Reviews and Ratings section'),
        ('←', 'Query reviews table for all reviews'),
        ('←', 'Display reviews with star ratings and text'),
        ('→', 'Flag inappropriate review'),
        ('←', 'Update review flag status and hide from customer pages'),
        ('←', 'Confirm flag action'),
    ])

    _h(doc, 'SSD:56 Generate Performance Report', bold=True, size=11)
    _ssd_table(doc, [
        ('→', 'Navigate to Reports > Performance Report and select date range'),
        ('←', 'Analyze order data for performance metrics'),
        ('←', 'Calculate top-selling products and customization trends'),
        ('←', 'Generate charts and report table'),
        ('→', 'Click Export as CSV or PDF'),
        ('←', 'Generate and download report file'),
    ])

    doc.add_page_break()

    # =========================================================
    # D) SECTION 3.5 SCHEMA DIAGRAM
    # =========================================================
    _h(doc, '3.5 Schema Diagram', bold=True, size=14)
    _b(doc, 'The database schema for AURA-WEAR consists of 9 tables that store all system data. Each table is described below with its columns, data types, and constraints.')

    _schema_table(doc, 'users', [
        ('user_id', 'INT', 'PRIMARY KEY, AUTO_INCREMENT'),
        ('name', 'VARCHAR(100)', 'NOT NULL'),
        ('email', 'VARCHAR(150)', 'UNIQUE, NOT NULL'),
        ('password_hash', 'VARCHAR(255)', 'NOT NULL'),
        ('role', 'ENUM(customer, admin)', 'DEFAULT customer'),
        ('phone', 'VARCHAR(20)', 'NULLABLE'),
        ('address', 'TEXT', 'NULLABLE'),
        ('city', 'VARCHAR(50)', 'NULLABLE'),
        ('profile_pic_url', 'TEXT', 'NULLABLE'),
        ('status', 'ENUM(active, inactive)', 'DEFAULT active'),
        ('created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
    ])

    _schema_table(doc, 'products', [
        ('product_id', 'INT', 'PRIMARY KEY, AUTO_INCREMENT'),
        ('name', 'VARCHAR(150)', 'NOT NULL'),
        ('category', 'VARCHAR(50)', 'NOT NULL'),
        ('description', 'TEXT', 'NULLABLE'),
        ('base_price', 'DECIMAL(10,2)', 'NOT NULL'),
        ('available_sizes', 'TEXT', 'NULLABLE'),
        ('image_urls', 'TEXT', 'NULLABLE'),
        ('status', 'ENUM(active, inactive)', 'DEFAULT active'),
        ('created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
    ])

    _schema_table(doc, 'orders', [
        ('order_id', 'VARCHAR(20)', 'PRIMARY KEY'),
        ('customer_id', 'INT', 'FOREIGN KEY -> users(user_id)'),
        ('total_price', 'DECIMAL(10,2)', 'NOT NULL'),
        ('payment_method', 'VARCHAR(20)', 'NOT NULL'),
        ('payment_status', 'VARCHAR(20)', 'DEFAULT Pending'),
        ('status', 'VARCHAR(30)', 'DEFAULT Pending'),
        ('delivery_address', 'TEXT', 'NOT NULL'),
        ('created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
        ('updated_at', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP'),
    ])

    _schema_table(doc, 'order_items', [
        ('item_id', 'INT', 'PRIMARY KEY, AUTO_INCREMENT'),
        ('order_id', 'VARCHAR(20)', 'FOREIGN KEY -> orders(order_id)'),
        ('product_id', 'INT', 'FOREIGN KEY -> products(product_id)'),
        ('fabric_type', 'VARCHAR(30)', 'NULLABLE'),
        ('color_hex', 'VARCHAR(10)', 'NULLABLE'),
        ('logo_url', 'TEXT', 'NULLABLE'),
        ('logo_placement', 'VARCHAR(20)', 'NULLABLE'),
        ('stitch_style', 'VARCHAR(30)', 'NULLABLE'),
        ('print_method', 'VARCHAR(30)', 'NULLABLE'),
        ('wash_effect', 'VARCHAR(30)', 'NULLABLE'),
        ('quantity', 'INT', 'DEFAULT 1'),
        ('unit_price', 'DECIMAL(10,2)', 'NOT NULL'),
    ])

    _schema_table(doc, 'payments', [
        ('payment_id', 'INT', 'PRIMARY KEY, AUTO_INCREMENT'),
        ('order_id', 'VARCHAR(20)', 'FOREIGN KEY -> orders(order_id)'),
        ('amount', 'DECIMAL(10,2)', 'NOT NULL'),
        ('method', 'VARCHAR(20)', 'NOT NULL'),
        ('status', 'VARCHAR(20)', 'DEFAULT Pending'),
        ('transaction_id', 'VARCHAR(100)', 'NULLABLE'),
        ('created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
    ])

    _schema_table(doc, 'chat_logs', [
        ('log_id', 'INT', 'PRIMARY KEY, AUTO_INCREMENT'),
        ('user_id', 'INT', 'FOREIGN KEY -> users(user_id), NULLABLE'),
        ('query', 'TEXT', 'NOT NULL'),
        ('response', 'TEXT', 'NOT NULL'),
        ('timestamp', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
    ])

    _schema_table(doc, 'inventory', [
        ('inventory_id', 'INT', 'PRIMARY KEY, AUTO_INCREMENT'),
        ('material_name', 'VARCHAR(100)', 'NOT NULL'),
        ('quantity', 'DECIMAL(10,2)', 'DEFAULT 0'),
        ('unit', 'VARCHAR(20)', 'NOT NULL'),
        ('alert_threshold', 'DECIMAL(10,2)', 'DEFAULT 10'),
        ('last_updated', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
    ])

    _schema_table(doc, 'coupons', [
        ('coupon_id', 'INT', 'PRIMARY KEY, AUTO_INCREMENT'),
        ('code', 'VARCHAR(20)', 'UNIQUE, NOT NULL'),
        ('discount_type', 'ENUM(percent, fixed)', 'NOT NULL'),
        ('value', 'DECIMAL(10,2)', 'NOT NULL'),
        ('expiry_date', 'DATE', 'NOT NULL'),
        ('usage_limit', 'INT', 'DEFAULT 1'),
        ('used_count', 'INT', 'DEFAULT 0'),
        ('status', 'ENUM(active, inactive)', 'DEFAULT active'),
    ])

    _schema_table(doc, 'faq', [
        ('faq_id', 'INT', 'PRIMARY KEY, AUTO_INCREMENT'),
        ('question', 'TEXT', 'NOT NULL'),
        ('answer', 'TEXT', 'NOT NULL'),
        ('created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
    ])

    doc.add_page_break()

    # =========================================================
    # E) CHAPTER 4: CONSTRUCTION
    # =========================================================
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    _h(doc, 'CHAPTER # 4:', bold=True, size=24, center=True)
    _h(doc, 'CONSTRUCTION', bold=True, size=24, center=True)

    doc.add_page_break()

    # 4.1 Class Diagram
    _h(doc, '4.1 Class Diagram', bold=True, size=14)
    _b(doc, 'The Class Diagram defines the static structure of AURA-WEAR\'s backend system, showing classes, their attributes, methods, and relationships.')

    class_table = doc.add_table(rows=1 + 9, cols=4, style='Table Grid')
    class_headers = ['Class', 'Attributes', 'Methods', 'Relationships']
    for i, h in enumerate(class_headers):
        p = class_table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)

    class_rows = [
        ('UserModel', 'user_id, name, email, password_hash, role, phone, address, status', 'create_user(), get_user(), update_user(), verify_password(), generate_jwt()', 'Has many Orders, Has many ChatLogs'),
        ('ProductModel', 'product_id, name, category, description, base_price, status', 'add_product(), get_product(), update_product(), toggle_status()', 'Has many OrderItems'),
        ('OrderModel', 'order_id, customer_id, total_price, status, payment_method', 'create_order(), get_order(), update_status(), cancel_order()', 'Belongs to User, Has many OrderItems, Has one Payment'),
        ('OrderItemModel', 'item_id, order_id, product_id, fabric, color, logo, stitch, print, wash, price', 'add_item(), calculate_price()', 'Belongs to Order, References Product'),
        ('PaymentModel', 'payment_id, order_id, amount, method, status, transaction_id', 'record_payment(), generate_receipt(), download_invoice()', 'Belongs to Order'),
        ('ChatbotService', 'n8n_webhook_url, openai_key', 'send_query(), get_response(), save_log(), get_history()', 'Uses ChatLog, Calls n8n API'),
        ('AdminService', ': ', 'get_dashboard_stats(), generate_report(), export_csv(), send_notification()', 'Manages all Models'),
        ('CouponModel', 'code, discount_type, value, expiry, limit, used_count', 'validate_coupon(), apply_discount(), deactivate()', 'Applied to Order'),
        ('InventoryModel', 'material, quantity, threshold', 'update_quantity(), check_alerts()', 'Managed by Admin'),
    ]
    for i, (cls, attrs, methods, rels) in enumerate(class_rows):
        row = class_table.rows[i + 1]
        for j, val in enumerate([cls, attrs, methods, rels]):
            p = row.cells[j].paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

    doc.add_paragraph('')
    doc.add_page_break()

    # 4.2 Project Code (Business Logic)
    _h(doc, '4.2 Project Code (Business Logic)', bold=True, size=14)

    # 4.2.1 Backend (FastAPI)
    _h(doc, '4.2.1 Backend: FastAPI Routes', bold=True, size=13)
    _b(doc, 'The following code samples demonstrate the core business logic implemented in the AURA-WEAR FastAPI backend.')

    _h(doc, 'auth.py: Authentication Router', bold=True, size=12)
    _code_block(doc, """from fastapi import APIRouter, HTTPException, Depends
from fastapi.security import OAuth2PasswordBearer
from jose import jwt, JWTError
from passlib.context import CryptContext
from datetime import datetime, timedelta
import os

router = APIRouter(prefix="/auth", tags=["Authentication"])
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
SECRET_KEY = os.getenv("SECRET_KEY", "aura-wear-secret")
ALGORITHM = "HS256"

@router.post("/register")
async def register(user_data: UserCreate, db: Session = Depends(get_db)):
    if db.query(User).filter(User.email == user_data.email).first():
        raise HTTPException(status_code=400, detail="Email already registered")
    hashed_password = pwd_context.hash(user_data.password)
    new_user = User(name=user_data.name, email=user_data.email,
                    password_hash=hashed_password, role="customer")
    db.add(new_user)
    db.commit()
    return {"message": "Registration successful"}

@router.post("/login")
async def login(credentials: LoginSchema, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == credentials.email).first()
    if not user or not pwd_context.verify(credentials.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = jwt.encode({"sub": str(user.user_id), "role": user.role,
                        "exp": datetime.utcnow() + timedelta(minutes=30)},
                       SECRET_KEY, algorithm=ALGORITHM)
    return {"access_token": token, "token_type": "bearer"}""")

    _h(doc, 'orders.py: Order Management Router', bold=True, size=12)
    _code_block(doc, """import uuid
from fastapi import APIRouter, HTTPException, Depends
router = APIRouter(prefix="/orders", tags=["Orders"])

@router.post("/create")
async def create_order(order_data: OrderCreate,
                       current_user: User = Depends(get_current_user),
                       db: Session = Depends(get_db)):
    order_id = "AW-" + str(uuid.uuid4())[:8].upper()
    new_order = Order(
        order_id=order_id,
        customer_id=current_user.user_id,
        total_price=order_data.total_price,
        payment_method=order_data.payment_method,
        delivery_address=order_data.delivery_address,
        status="Pending"
    )
    db.add(new_order)
    for item in order_data.items:
        order_item = OrderItem(order_id=order_id, **item.dict())
        db.add(order_item)
    db.commit()
    return {"order_id": order_id, "status": "Pending"}

@router.get("/history")
async def get_order_history(current_user: User = Depends(get_current_user),
                            db: Session = Depends(get_db)):
    orders = db.query(Order).filter(Order.customer_id == current_user.user_id
                                    ).order_by(Order.created_at.desc()).all()
    return orders""")

    _h(doc, 'chatbot.py: AI Chatbot Router', bold=True, size=12)
    _code_block(doc, """import httpx
from fastapi import APIRouter
router = APIRouter(prefix="/chatbot", tags=["Chatbot"])
N8N_WEBHOOK = os.getenv("N8N_WEBHOOK_URL")

@router.post("/query")
async def chatbot_query(query: ChatQuery,
                        current_user: Optional[User] = Depends(get_optional_user),
                        db: Session = Depends(get_db)):
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.post(N8N_WEBHOOK, json={"query": query.message,
                                                         "user_id": current_user.user_id if current_user else None})
    result = response.json()
    if current_user:
        log = ChatLog(user_id=current_user.user_id,
                      query=query.message, response=result["response"])
        db.add(log)
        db.commit()
    return {"response": result["response"]}""")

    doc.add_page_break()

    # 4.2.2 Frontend (JavaScript)
    _h(doc, '4.2.2 Frontend: JavaScript Logic', bold=True, size=13)
    _b(doc, 'The following code sample demonstrates the real-time 2D preview and pricing logic implemented in the AURA-WEAR frontend customizer.')

    _h(doc, 'customizer.js: Real-Time 2D Preview', bold=True, size=12)
    _code_block(doc, """const canvas = document.getElementById('previewCanvas');
const ctx = canvas.getContext('2d');
let customization = { color: '#FFFFFF', logoUrl: null, logoPos: 'chest', washEffect: 'standard' };

function updatePreview() {
    ctx.clearRect(0, 0, canvas.width, canvas.height);
    drawGarment(customization.color);
    if (customization.logoUrl) drawLogo(customization.logoUrl, customization.logoPos);
    applyWashEffect(customization.washEffect);
    updatePrice();
}

function drawGarment(color) {
    ctx.fillStyle = color;
    ctx.beginPath();
    ctx.roundRect(50, 30, 300, 380, 20);
    ctx.fill();
}

function updatePrice() {
    const base = parseFloat(document.getElementById('basePrice').value);
    const fabricPremium = { cotton: 0, polyester: 0.05, blended: 0.10, fleece: 0.15 };
    const printCost = { screen: 300, dtg: 500, embroidery: 800, heat: 400 };
    const fabric = document.getElementById('fabricSelect').value;
    const print = document.getElementById('printSelect').value;
    const total = base * (1 + (fabricPremium[fabric] || 0)) + (printCost[print] || 0);
    document.getElementById('totalPrice').textContent = 'PKR ' + total.toFixed(0);
}""")

    doc.add_page_break()

    # 4.3 Sample Queries
    _h(doc, '4.3 Sample Queries', bold=True, size=14)
    _b(doc, 'The following SQL queries represent the core data access operations used by the AURA-WEAR system.')

    query_table = doc.add_table(rows=1 + 9, cols=2, style='Table Grid')
    qh0 = query_table.rows[0].cells[0].paragraphs[0]
    qr0 = qh0.add_run('Query Name')
    qr0.bold = True
    qr0.font.name = 'Times New Roman'
    qr0.font.size = Pt(11)
    qh1 = query_table.rows[0].cells[1].paragraphs[0]
    qr1 = qh1.add_run('SQL Query')
    qr1.bold = True
    qr1.font.name = 'Times New Roman'
    qr1.font.size = Pt(11)

    query_rows = [
        ('Get All Active Products', "SELECT * FROM products WHERE status = 'active' ORDER BY created_at DESC;"),
        ('Get Customer Orders', "SELECT o.*, oi.* FROM orders o JOIN order_items oi ON o.order_id = oi.order_id WHERE o.customer_id = ? ORDER BY o.created_at DESC;"),
        ('Monthly Revenue', "SELECT SUM(amount) as revenue, COUNT(*) as orders FROM payments WHERE status='completed' AND strftime('%Y-%m', created_at) = strftime('%Y-%m', 'now');"),
        ('Search Products', "SELECT * FROM products WHERE (name LIKE ? OR description LIKE ?) AND status='active';"),
        ('Low Stock Inventory', 'SELECT * FROM inventory WHERE quantity <= alert_threshold;'),
        ('Validate Coupon', "SELECT * FROM coupons WHERE code=? AND status='active' AND expiry_date >= date('now') AND used_count < usage_limit;"),
        ('Get Chat History', 'SELECT query, response, timestamp FROM chat_logs WHERE user_id=? ORDER BY timestamp DESC LIMIT 20;'),
        ('Update Order Status', "UPDATE orders SET status=?, updated_at=datetime('now') WHERE order_id=?;"),
        ('Get Admin Dashboard Stats', "SELECT (SELECT COUNT(*) FROM orders WHERE date(created_at)=date('now')) as today_orders, (SELECT SUM(amount) FROM payments WHERE strftime('%W', created_at)=strftime('%W','now') AND status='completed') as weekly_revenue, (SELECT COUNT(*) FROM users WHERE role='customer') as total_customers;"),
    ]

    for i, (qname, qsql) in enumerate(query_rows):
        row = query_table.rows[i + 1]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(qname)
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(10)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(qsql)
        r1.font.name = 'Courier New'
        r1.font.size = Pt(9)

    doc.add_paragraph('')
    doc.add_page_break()
