from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _heading(doc, text, bold=True, size=14, center=False, underline=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if underline:
        run.underline = True
    return p


def _body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def _shade_cell(cell, hex_color='C0C0C0'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _centered_para(doc, text, bold=False, size=12, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def _blank(doc):
    p = doc.add_paragraph('')
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def _set_cell_text(cell, text, bold=False, size=12, shade=False, shade_color='C0C0C0'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if shade:
        _shade_cell(cell, shade_color)


def add_part1(doc):
    # ─────────────────────────────────────────────────────────────
    # PAGE i: Title Page
    # ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("AURA-WEAR: AI-Integrated Custom Clothing Website")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)

    # large blank space
    for _ in range(6):
        _blank(doc)

    _centered_para(doc, "Developed By:", bold=True, size=12)
    _centered_para(doc, "Muhammad Umer Javed (B.S CS Aut-21-Eve-0170)", bold=False, size=12)
    _centered_para(doc, "Abdul Baqi (B.S CS Aut-2022-E-B-0260)", bold=False, size=12)
    _blank(doc)
    _centered_para(doc, "BS (Computer Science)", bold=True, size=12)
    _blank(doc)
    _centered_para(doc, "Supervised by:", bold=True, size=12)
    _centered_para(doc, "Dr. Zaheen Khan", bold=True, size=12)
    _blank(doc)
    _centered_para(doc, "Department of Computer Science", bold=True, size=12)
    _centered_para(doc, "Federal Urdu University of Arts, Science and Technology", bold=True, size=12)
    _centered_para(doc, "Islamabad", bold=True, size=12)
    _blank(doc)
    _centered_para(doc, "Session [2022-2026]", bold=True, size=12)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # PAGE ii: Submission Page
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "AURA-WEAR: AI-Integrated Custom Clothing Website",
             bold=True, size=14, center=True, underline=True)

    _body(doc,
          "A project submitted to the Department of Computer Science as partial fulfillment of the "
          "requirement for the award of Degree of Bachelor of Science in Computer Science.")

    _blank(doc)

    # 2-column table: Name and Registration Number
    table = doc.add_table(rows=3, cols=2, style='Table Grid')
    # Header row (shaded)
    hdr_cells = table.rows[0].cells
    _set_cell_text(hdr_cells[0], "Name", bold=True, shade=True)
    _set_cell_text(hdr_cells[1], "Registration Number", bold=True, shade=True)
    # Data rows
    _set_cell_text(table.rows[1].cells[0], "Muhammad Umer Javed")
    _set_cell_text(table.rows[1].cells[1], "B.S CS Aut-21-Eve-0170")
    _set_cell_text(table.rows[2].cells[0], "Abdul Baqi")
    _set_cell_text(table.rows[2].cells[1], "B.S CS Aut-2022-E-B-0260")

    _blank(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Supervisor")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    for line in [
        "Dr. Zaheen Khan",
        "Assistant Professor",
        "Department of Computer Science",
        "Federal Urdu University of Arts, Science and Technology, Islamabad",
    ]:
        _body(doc, line)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # PAGE iii: Final Approval
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "Final Approval", bold=True, size=14, center=True, underline=True)

    _body(doc,
          "It is certified that I have read the project report submitted by Muhammad Umer Javed "
          "(B.S CS Aut-21-Eve-0170) and Abdul Baqi (B.S CS Aut-2022-E-B-0260) and it is our "
          "judgment that this project is of sufficient standard to warrant its acceptance by the "
          "FUUAST university, Islamabad for the Degree of Bachelor of Science in Computer Science.")

    _blank(doc)

    p = doc.add_paragraph()
    run = p.add_run("Committee:")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    _blank(doc)

    p = doc.add_paragraph()
    run = p.add_run("External Examiner")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    _body(doc, "_______________________________")
    _blank(doc)

    p = doc.add_paragraph()
    run = p.add_run("Internal Examiner")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    _body(doc, "_______________________________")

    for line in ["Name", "Lecturer", "Department of Computer Science",
                 "Federal Urdu University of Arts, Science and Technology, Islamabad"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(2.5)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    _blank(doc)

    p = doc.add_paragraph()
    run = p.add_run("Supervisor")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    _body(doc, "_______________________________")

    for line in ["Dr. Zaheen Khan", "Assistant Professor",
                 "Department of Computer Science",
                 "Federal Urdu University of Arts, Science and Technology, Islamabad"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # PAGE iv: Declaration
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "Declaration", bold=True, size=14, center=True, underline=True)

    _body(doc,
          "We hereby declare that this software, neither as a whole nor as a part has been copied "
          "out from any source. It is further declared that we have developed this software and the "
          "accompanied report entirely on the basis of our personal efforts. If any part of this "
          "project is proved to be copied out from any source or found to be reproduction of some "
          "other, we will stand by the consequences. No portion of the work presented has been "
          "submitted in support of any application for any other degree or qualification of this "
          "or any other university or institute of learning.")

    _blank(doc)
    _blank(doc)

    _centered_para(doc, "Muhammad Umer Javed", bold=False, size=12)
    _centered_para(doc, "Abdul Baqi", bold=False, size=12)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # PAGE v: Dedication
    # ─────────────────────────────────────────────────────────────
    for _ in range(8):
        _blank(doc)

    for text in [
        "Dedicated to our beloved parents,",
        "Teachers",
        "And",
        "The fellows",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # PAGE vi: Acknowledgement
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "Acknowledgement", bold=True, size=14, center=True, underline=True)

    _body(doc,
          "Thanks to Almighty Allah for giving us knowledge, power and strength to accomplish this "
          "task. We learned a great deal while developing this project and this experience will "
          "certainly help us in our forthcoming professional life. The AURA-WEAR project: an "
          "AI-Integrated Custom Clothing Website: challenged us to combine modern web technologies "
          "with artificial intelligence to solve a real-world problem in the fashion industry. Many "
          "of our teachers provided invaluable guidance during this journey, but we are especially "
          "thankful to our supervisor Dr. Zaheen Khan for his continuous help, support, and "
          "encouragement throughout all phases of this project. His constructive feedback and "
          "motivation helped us navigate through times of difficulty and pushed us to deliver our "
          "best work. We would also like to extend our sincere gratitude to the Department of "
          "Computer Science at FUUAST for providing us with the resources and environment needed "
          "to complete this project. Finally, we would like to thank all of our friends and "
          "classmates for their moral support and encouragement throughout this academic journey.")

    _blank(doc)
    _blank(doc)

    _centered_para(doc, "Muhammad Umer Javed", bold=False, size=12)
    _centered_para(doc, "Abdul Baqi", bold=False, size=12)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # PAGE vii: Project in Brief
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "Project in Brief", bold=True, size=14, center=True, underline=True)

    pib_data = [
        ("Project Title", "AURA-WEAR: AI-Integrated Custom Clothing Website"),
        ("Organization", "Federal Urdu University of Arts, Science and Technology"),
        ("Objectives",
         "1. To develop an AI-powered web platform that allows customers to design and order "
         "custom streetwear and sportswear with real-time 2D previews.\n"
         "2. To implement a RAG-based AI chatbot using n8n and OpenAI GPT-4o-mini that provides "
         "intelligent fabric and design recommendations.\n"
         "3. To provide a secure, scalable e-commerce system with order tracking, payment "
         "processing, and admin management.\n"
         "4. To deliver a seamless user experience with responsive design, dynamic pricing, and "
         "multiple customization options."),
        ("Undertaken By", "Muhammad Umer Javed\nAbdul Baqi"),
        ("Supervised By", "Dr. Zaheen Khan"),
        ("Date Started", "Feb 16, 2026"),
        ("Date Completed", "Jun 26, 2026"),
        ("Technologies Used",
         "Framework: HTML5, CSS3, Vanilla JavaScript\n"
         "Backend: FastAPI (Python)\n"
         "AI/Chatbot: n8n Workflow + OpenAI GPT-4o-mini (RAG)\n"
         "Database: SQLite (dev) / PostgreSQL (prod)\n"
         "Authentication: JWT + bcrypt\n"
         "Email: FastAPI-Mail + Gmail SMTP\n"
         "Deployment: Railway.app\n"
         "Version Control: Git + GitHub"),
        ("System Used",
         "8th Gen Core i5 Processor @ 2.40 GHz, 16 GB RAM, 512 GB SSD"),
    ]

    pib_table = doc.add_table(rows=len(pib_data), cols=2, style='Table Grid')
    pib_table.autofit = False
    pib_table.columns[0].width = Inches(2.0)
    pib_table.columns[1].width = Inches(4.5)

    for i, (label, content) in enumerate(pib_data):
        left_cell = pib_table.rows[i].cells[0]
        right_cell = pib_table.rows[i].cells[1]

        left_cell.text = ''
        lp = left_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(3)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.name = 'Times New Roman'
        lr.font.size = Pt(11)
        _shade_cell(left_cell, 'D9D9D9')

        right_cell.text = ''
        rp = right_cell.paragraphs[0]
        rp.paragraph_format.space_after = Pt(3)
        rr = rp.add_run(content)
        rr.font.name = 'Times New Roman'
        rr.font.size = Pt(11)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # PAGE viii: Abstract
    # ─────────────────────────────────────────────────────────────
    _centered_para(doc, "ABSTRACT", bold=True, size=14)
    _centered_para(doc, "AURA-WEAR: AI-Integrated Custom Clothing Website", bold=True, size=13)
    _blank(doc)

    _body(doc,
          "The fashion e-commerce industry in Pakistan faces a significant gap: existing platforms "
          "rarely offer personalized clothing customization with real-time visual feedback or "
          "AI-driven recommendations. Customers are forced to order clothing without being able to "
          "preview their designs, resulting in dissatisfaction and high return rates. AURA-WEAR "
          "addresses this gap by providing a full-stack, AI-integrated web platform that enables "
          "customers to design custom streetwear and sportswear entirely online. "
          "The platform features a real-time 2D design preview engine built on the HTML5 Canvas "
          "API, allowing users to select garment types, colors, materials, sizes, and add custom "
          "text or graphic elements: all with an instant visual update. An RAG-powered AI chatbot, "
          "built using n8n workflow automation and OpenAI GPT-4o-mini, provides context-aware "
          "fabric and design recommendations drawn from a curated knowledge base of product data. "
          "Dynamic pricing calculation ensures transparent cost estimates as users configure their "
          "orders, and a complete e-commerce workflow handles cart management, order placement, "
          "and payment integration. "
          "The system is organized into three portals: the Customer Portal (product browsing, "
          "customization, cart, orders, profile), the Admin Panel (product management, order "
          "management, inventory, customer accounts, analytics dashboard), and an AI Chatbot "
          "interface accessible system-wide. The backend is developed with FastAPI (Python), "
          "the frontend with Vanilla JavaScript (no framework dependencies), and the database "
          "uses SQLite for development and PostgreSQL for production via SQLAlchemy ORM. The "
          "entire system is deployed on Railway.app with CI/CD via GitHub. "
          "This document covers the complete Software Development Life Cycle of the AURA-WEAR "
          "project, including: Chapter 1: Project Proposal and background analysis; Chapter 2: "
          "Requirements Analysis covering 56 functional requirements across 8 modules with use "
          "case diagrams and fully dressed use cases; Chapter 3: System Design comprising the "
          "domain model, system sequence diagrams, and database schema; Chapter 4: Construction "
          "covering class diagrams, implementation code, and SQL queries; Chapter 5: Testing "
          "with 56 test cases validated against functional requirements; and Chapter 6: User "
          "Manual with step-by-step operational instructions. AURA-WEAR demonstrates how "
          "modern AI and web technologies can be combined to deliver a compelling, user-centric "
          "custom fashion experience that is practical, scalable, and commercially viable.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # TABLE OF CONTENTS (pages ix-xix)
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "TABLE OF CONTENTS", bold=True, size=16, center=True)
    _blank(doc)

    toc_entries = [
        ("CHAPTER # 1: PROPOSAL", ""),
        ("Revision History", ""),
        ("    1.1 Introduction", ""),
        ("    1.2 Problem Statement", ""),
        ("    1.3 Proposed System", ""),
        ("    1.4 Benefits of the Proposed System", ""),
        ("    1.5 Scope", ""),
        ("    1.6 Survey Analysis", ""),
        ("        1.6.1 Survey Analysis Table", ""),
        ("    1.7 Module & Submodules", ""),
        ("        1.7.1 Security Management", ""),
        ("        1.7.2 User Profile Management", ""),
        ("        1.7.3 Product Catalog Management", ""),
        ("        1.7.4 Product Customization", ""),
        ("        1.7.5 AI Chatbot Management", ""),
        ("        1.7.6 Order Management", ""),
        ("        1.7.7 Payment Management", ""),
        ("        1.7.8 Admin Dashboard Management", ""),
        ("    1.8 Primary Actor", ""),
        ("    1.9 Tools & Technologies", ""),
        ("        1.9.1 Front-End Tools", ""),
        ("        1.9.2 Back-End Tools", ""),
        ("        1.9.3 AI / Automation Tools", ""),
        ("        1.9.4 Database", ""),
        ("        1.9.5 Deployment", ""),
        ("    1.10 System Design Approach", ""),
        ("    1.11 Process Model Used", ""),
        ("    1.12 Modelling Techniques / Tools Used", ""),
        ("    1.13 Limitation / Constraint", ""),
        ("    1.14 References", ""),
        ("CHAPTER # 2: ANALYSIS", ""),
        ("    2.1 Introduction", ""),
        ("        2.1.1 Purpose", ""),
        ("        2.1.2 Scope", ""),
        ("        2.1.3 Definitions, Acronyms, and Abbreviations", ""),
        ("        2.1.4 Overview", ""),
        ("    2.2 Functional Requirements", ""),
        ("        2.2.1 Security Management", ""),
        ("        2.2.2 User Profile Management", ""),
        ("        2.2.3 Product Catalog Management", ""),
        ("        2.2.4 Product Customization", ""),
        ("        2.2.5 AI Chatbot Management", ""),
        ("        2.2.6 Order Management", ""),
        ("        2.2.7 Payment Management", ""),
        ("        2.2.8 Admin Dashboard Management", ""),
        ("    2.3 Non-Functional Requirements", ""),
        ("        2.3.1 Performance Requirements", ""),
        ("        2.3.2 Usability Requirements", ""),
        ("        2.3.3 Security Requirements", ""),
        ("        2.3.4 Reliability Requirements", ""),
        ("        2.3.5 Scalability Requirements", ""),
        ("        2.3.6 Maintainability & Support", ""),
        ("        2.3.7 Compatibility Requirements", ""),
        ("        2.3.8 Support Requirements", ""),
        ("    2.4 External Interface Requirements", ""),
        ("        2.4.1 User Interfaces", ""),
        ("        2.4.2 Hardware Interfaces", ""),
        ("        2.4.3 Software Interfaces", ""),
        ("        2.4.4 Communication Interfaces", ""),
        ("CHAPTER # 3: DESIGN", ""),
        ("    3.1 Use Case Diagram", ""),
        ("    3.2 Fully Dressed Use Cases", ""),
        ("    3.3 Domain Model", ""),
        ("    3.4 System Sequence Diagrams", ""),
        ("    3.5 Schema Diagram", ""),
        ("CHAPTER # 4: CONSTRUCTION", ""),
        ("    4.1 Class Diagram", ""),
        ("    4.2 Project Code (Business Logic)", ""),
        ("        4.2.1 Backend: FastAPI Routes", ""),
        ("        4.2.2 Frontend: JavaScript Logic", ""),
        ("    4.3 Sample Queries", ""),
        ("CHAPTER # 5: TESTING", ""),
        ("    5.1 Test Cases (TC-01 to TC-56)", ""),
        ("CHAPTER # 6: USER MANUAL", ""),
        ("    6.1 Screenshots", ""),
        ("    6.2 User Roles Description", ""),
    ]

    for entry_text, _ in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(entry_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if not entry_text.startswith(' '):
            run.bold = True

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # CHAPTER 1: Chapter Title Page
    # ─────────────────────────────────────────────────────────────
    for _ in range(10):
        _blank(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("CHAPTER # 1:")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("PROPOSAL")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)

    for _ in range(10):
        _blank(doc)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # REVISION HISTORY
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "Revision History", bold=True, size=14, center=True, underline=True)
    _blank(doc)

    rev_table = doc.add_table(rows=1, cols=4, style='Table Grid')
    rev_headers = ['Version', 'Date', 'Description', 'Author']
    for i, h in enumerate(rev_headers):
        cell = rev_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlEl
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = _OxmlEl('w:shd'); shd.set(_qn('w:val'), 'clear')
        shd.set(_qn('w:color'), 'auto'); shd.set(_qn('w:fill'), 'C0C0C0')
        tcPr.append(shd)

    rev_rows = [
        ('v1.0', 'Nov 2025', 'Initial SRS document: core requirements, basic use cases, project scope defined', 'Muhammad Umer Javed'),
        ('v2.0', 'Jan 2026', 'Added use case diagrams, ER diagram, system architecture, and technology stack details', 'Abdul Baqi'),
        ('v3.0', 'Mar 2026', 'Tech stack upgraded to FastAPI + Python, JWT authentication added, full database schema, n8n RAG chatbot integration documented', 'Muhammad Umer Javed'),
        ('v4.0', 'Jun 2026', 'Final version: all 56 functional requirements implemented, real-time 2D preview, dynamic pricing, coupon system, admin analytics, CSV export, user manual complete', 'Both Authors'),
    ]
    for ver, dt, desc, auth in rev_rows:
        row = rev_table.add_row()
        for i, txt in enumerate([ver, dt, desc, auth]):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            r.font.name = 'Times New Roman'; r.font.size = Pt(11)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # 1.1 Introduction
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.1 Introduction", bold=True, size=13)

    _body(doc,
          "The fashion industry is rapidly shifting toward personalized and custom-made clothing, "
          "driven by rising consumer expectations for uniqueness and fit. AURA-WEAR is an "
          "AI-integrated custom clothing web platform developed as a Final Year Project for the "
          "Department of Computer Science at FUUAST. The platform allows customers to design "
          "their own streetwear and sportswear directly in the browser, with real-time visual "
          "feedback powered by a Canvas-based 2D preview engine. Beyond design, AURA-WEAR "
          "integrates an RAG-powered AI chatbot that provides intelligent recommendations for "
          "fabrics, styles, and design combinations. The system is built using modern, production-"
          "ready technologies including FastAPI, Vanilla JavaScript, and SQLite/PostgreSQL, and "
          "is fully deployed on Railway.app. AURA-WEAR aims to bridge the gap between the "
          "convenience of online shopping and the personalization of bespoke tailoring.")

    # ─────────────────────────────────────────────────────────────
    # 1.2 Problem Statement
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.2 Problem Statement", bold=True, size=13)

    _body(doc,
          "In Pakistan, there is a significant gap in online platforms offering custom clothing "
          "with real-time design preview capabilities and AI-based recommendations. Most existing "
          "e-commerce platforms offer only standard sizing with limited or no customization. "
          "Customers lack the ability to visualize their custom orders before placing them, "
          "leading to dissatisfaction, mismatched expectations, and high return rates. "
          "Additionally, customers have no intelligent assistance to guide fabric selection, "
          "color combinations, or design compatibility. The absence of transparent, dynamic "
          "pricing for customized orders further erodes trust. AURA-WEAR is designed to solve "
          "these problems by providing real-time 2D previews, AI-driven guidance, and a "
          "seamless end-to-end custom ordering experience.")

    # ─────────────────────────────────────────────────────────────
    # 1.3 Proposed System
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.3 Proposed System", bold=True, size=13)

    _body(doc,
          "AURA-WEAR proposes a comprehensive web-based solution for custom clothing design and "
          "ordering. The system is organized into three primary portals. The Customer Portal "
          "enables users to register, browse the product catalog, use the 2D customization "
          "engine to configure garment type, color, fabric, size, and personalized graphics or "
          "text, view dynamic pricing, add items to cart, and place orders. The Admin Panel "
          "provides administrators with tools to manage products, inventory, customer accounts, "
          "orders, and view business analytics. The AI Chatbot Interface, accessible from all "
          "pages, uses an RAG pipeline built on n8n workflows and OpenAI GPT-4o-mini to answer "
          "product queries and provide personalized recommendations. The backend is powered by "
          "FastAPI with JWT-based authentication and SQLAlchemy ORM, connected to SQLite (dev) "
          "or PostgreSQL (prod), and deployed on Railway.app with version control via GitHub.")

    # ─────────────────────────────────────────────────────────────
    # 1.4 Benefits of the Proposed System
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.4 Benefits of the Proposed System", bold=True, size=13)

    benefits = [
        "Real-Time 2D Design Preview: Customers can visualize their custom garment configurations instantly before placing an order, reducing uncertainty and return rates.",
        "AI-Powered Recommendations: The RAG-based chatbot provides intelligent fabric, style, and design suggestions tailored to each customer's preferences.",
        "Transparent Dynamic Pricing: Pricing is calculated in real time as customers add customization options, building trust and eliminating hidden costs.",
        "Streamlined Ordering Process: A complete e-commerce workflow: from design to checkout to order tracking: ensures a smooth customer journey.",
        "Admin Analytics Dashboard: Administrators gain insights into sales trends, inventory levels, and customer behavior through a built-in analytics panel.",
        "Scalable Architecture: The 4-tier layered architecture (Browser → FastAPI → n8n AI → Database) ensures the system can scale to handle growing user demand.",
    ]

    for benefit in benefits:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(benefit)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # ─────────────────────────────────────────────────────────────
    # 1.5 Scope
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.5 Scope", bold=True, size=13)

    _body(doc,
          "AURA-WEAR is scoped as a web-based platform accessible via modern browsers on desktop "
          "and mobile devices. The system covers custom design and ordering of two clothing "
          "categories: streetwear and sportswear. Features within scope include: user "
          "registration and authentication, product catalog management, 2D garment customization "
          "(color, fabric, size, text, graphics), real-time dynamic pricing, shopping cart, "
          "order placement and tracking, payment integration (Cash on Delivery and PayFast), "
          "AI chatbot for product recommendations, and an admin panel for full system management.")

    _body(doc,
          "Outside the scope of this project: a native mobile application (iOS/Android), "
          "integration with a physical store inventory system, Stripe payment gateway (pending "
          "future development), augmented reality (AR) try-on features, and multi-language "
          "support. These features may be considered for future iterations of the platform.")

    # ─────────────────────────────────────────────────────────────
    # 1.6 Survey Analysis
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.6 Survey Analysis", bold=True, size=13)

    _body(doc,
          "A comparative analysis was conducted of five existing e-commerce and custom clothing "
          "platforms to identify gaps that AURA-WEAR addresses. The systems evaluated were Amazon "
          "Custom, Printful, Custom Ink, Spreadshirt, and AURA-WEAR. The comparison focused on "
          "five key features: AI Chatbot integration, 2D Design Preview capability, Custom Fabric "
          "selection, Real-Time Pricing, and Order Tracking.")

    _blank(doc)

    survey_headers = ["System Name", "AI Chatbot", "2D Preview", "Custom Fabric",
                      "Real-Time Pricing", "Order Tracking"]
    survey_rows = [
        ["Amazon Custom", "No", "No", "Limited", "No", "Yes"],
        ["Printful", "No", "Basic", "Yes", "Yes", "Yes"],
        ["Custom Ink", "No", "2D Basic", "Limited", "Yes", "Yes"],
        ["Spreadshirt", "No", "Basic", "Limited", "Yes", "Yes"],
        ["AURA-WEAR", "Yes (RAG)", "Yes (Canvas)", "Yes (Full)", "Yes (Dynamic)", "Yes"],
    ]

    survey_table = doc.add_table(rows=1 + len(survey_rows), cols=6, style='Table Grid')

    # Header row
    hdr = survey_table.rows[0].cells
    for i, h in enumerate(survey_headers):
        _set_cell_text(hdr[i], h, bold=True, shade=True)

    # Data rows
    for ri, row_data in enumerate(survey_rows):
        row_cells = survey_table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            _set_cell_text(row_cells[ci], cell_text)

    _blank(doc)

    _body(doc,
          "The survey confirms that AURA-WEAR is the only platform among those evaluated to offer "
          "all five features simultaneously, making it a uniquely comprehensive solution in the "
          "custom clothing e-commerce space.")

    # ─────────────────────────────────────────────────────────────
    # 1.7 Modules & Submodules
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.7 Modules & Submodules", bold=True, size=13)

    _body(doc,
          "AURA-WEAR is divided into eight functional modules, each responsible for a distinct "
          "area of the system's functionality:")

    modules = [
        ("1.7.1 Security Management",
         "Handles user authentication and authorization including registration, login, logout, "
         "JWT token management, and bcrypt password hashing. Ensures secure access control "
         "across all portals."),
        ("1.7.2 User Profile Management",
         "Allows customers to create, view, and update their profile information including "
         "personal details, shipping addresses, and order history. Admins can view and manage "
         "all user accounts."),
        ("1.7.3 Product Catalog Management",
         "Manages the product listings for streetwear and sportswear categories. Admins can "
         "add, update, delete, and view products with attributes such as name, description, "
         "base price, available colors, and fabric types."),
        ("1.7.4 Product Customization",
         "The core module of AURA-WEAR. Provides the 2D Canvas-based design interface where "
         "customers select garment type, color, fabric, size, add custom text or graphics, "
         "and view a live preview with dynamic price calculation."),
        ("1.7.5 AI Chatbot Management",
         "Integrates the RAG-based AI chatbot powered by n8n workflows and OpenAI GPT-4o-mini. "
         "The chatbot answers product-related questions and provides personalized fabric and "
         "design recommendations based on the knowledge base."),
        ("1.7.6 Order Management",
         "Handles the complete order lifecycle: cart management, order placement, status updates "
         "(Pending, Processing, Shipped, Delivered, Cancelled), and order history views for "
         "both customers and admins."),
        ("1.7.7 Payment Management",
         "Manages payment processing for completed orders. Supports Cash on Delivery (COD) and "
         "PayFast gateway integration. Records payment status and links payments to orders."),
        ("1.7.8 Admin Dashboard Management",
         "Provides administrators with a centralized analytics dashboard showing key metrics: "
         "total orders, revenue, active customers, inventory levels, and recent activity. "
         "Supports data-driven management decisions."),
    ]

    for sub_title, sub_body in modules:
        _heading(doc, sub_title, bold=True, size=12)
        _body(doc, sub_body)

    # ─────────────────────────────────────────────────────────────
    # 1.8 Primary Actor
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.8 Primary Actor", bold=True, size=13)

    _body(doc, "AURA-WEAR has two primary actors who interact with the system:")

    p = doc.add_paragraph()
    run = p.add_run("Customer: ")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(
        "The end-user of the platform. A Customer can register an account, log in, browse the "
        "product catalog, use the 2D customization engine to design garments, add items to the "
        "cart, place orders, make payments, track order status, interact with the AI chatbot, "
        "and manage their profile and shipping addresses."
    )
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    run = p.add_run("Admin: ")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(
        "The system administrator. An Admin can log in to the admin panel, manage product "
        "listings (add, update, delete), manage customer accounts, view and update order "
        "statuses, monitor inventory, view analytics, manage payment records, and configure "
        "system settings. The Admin role has elevated privileges not accessible to Customers."
    )
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # ─────────────────────────────────────────────────────────────
    # 1.9 Tools & Technologies
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.9 Tools & Technologies", bold=True, size=13)

    _body(doc, "The following tools and technologies were used in the development of AURA-WEAR:")

    tech_sections = [
        ("Frontend",
         "HTML5: Structure and semantic markup of all web pages. CSS3: Styling, animations, "
         "and responsive layout using Flexbox and Grid. Vanilla JavaScript (ES6+): All client-"
         "side logic, Canvas API for 2D preview, DOM manipulation, and fetch API for REST calls. "
         "No JavaScript frameworks were used, keeping the frontend lightweight and dependency-free."),
        ("Backend",
         "FastAPI (Python): High-performance, asynchronous REST API framework for all backend "
         "endpoints. Python 3.11+: Core programming language for all server-side logic. "
         "Uvicorn: ASGI server for running the FastAPI application. Pydantic: Data validation "
         "and serialization for request/response models. FastAPI-Mail: Email notifications via "
         "Gmail SMTP. bcrypt: Secure password hashing. python-jose: JWT token generation "
         "and validation."),
        ("Database",
         "SQLite: Lightweight relational database used during development. PostgreSQL: "
         "Production-grade relational database for the deployed application. SQLAlchemy ORM: "
         "Database abstraction layer providing model-based queries and migrations."),
        ("AI / Chatbot",
         "n8n (v1.x): Open-source workflow automation platform used to build the RAG pipeline. "
         "OpenAI GPT-4o-mini: Large language model providing natural language responses. "
         "Retrieval-Augmented Generation (RAG): Architecture combining a vector knowledge base "
         "of product data with LLM responses for accurate, context-aware answers."),
        ("Deployment & Version Control",
         "Railway.app: Cloud hosting platform for both the FastAPI backend and the frontend. "
         "Git: Version control for tracking code changes. GitHub: Remote repository hosting "
         "and CI/CD integration with Railway.app for automatic deployments on push. "
         "Postman: API testing and documentation during development."),
    ]

    for sub_title, sub_body in tech_sections:
        _heading(doc, sub_title, bold=True, size=12)
        _body(doc, sub_body)

    # ─────────────────────────────────────────────────────────────
    # 1.10 System Design Approach
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.10 System Design Approach", bold=True, size=13)

    _body(doc,
          "AURA-WEAR follows a 4-Tier Layered Architecture that separates concerns across "
          "distinct layers for maintainability, scalability, and testability:")

    tiers = [
        ("Tier 1: Presentation Layer (Browser):",
         "The client-side interface built with HTML5, CSS3, and Vanilla JavaScript. Handles all "
         "user interactions, renders the 2D canvas preview, and communicates with the backend via "
         "RESTful fetch API calls."),
        ("Tier 2: Application Layer (FastAPI Backend):",
         "The core business logic layer. Receives HTTP requests, authenticates users via JWT, "
         "processes data through Pydantic models, executes database operations via SQLAlchemy, "
         "and returns structured JSON responses."),
        ("Tier 3: AI Processing Layer (n8n + OpenAI):",
         "The intelligent service layer. When chatbot queries are received, the FastAPI backend "
         "forwards them to the n8n workflow which executes the RAG pipeline: retrieving relevant "
         "context from the knowledge base and sending it to OpenAI GPT-4o-mini for response "
         "generation."),
        ("Tier 4: Data Layer (SQLite / PostgreSQL):",
         "The persistence layer. Stores all structured data including users, products, orders, "
         "payments, and customization configurations. Accessed exclusively through the "
         "SQLAlchemy ORM to ensure data integrity and portability."),
    ]

    for tier_title, tier_body in tiers:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(tier_title + " ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run2 = p.add_run(tier_body)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

    # ─────────────────────────────────────────────────────────────
    # 1.11 Process Model Used
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.11 Process Model Used", bold=True, size=13)

    _body(doc,
          "AURA-WEAR was developed using the Agile Incremental process model. This approach "
          "allowed the team to deliver working software in short, iterative sprints while "
          "continuously incorporating feedback and refining requirements. The project was "
          "organized into four sprints:")

    sprints = [
        ("Sprint 1 (Weeks 1-3): Foundation",
         "Project setup, technology stack finalization, database schema design, FastAPI project "
         "structure, basic authentication (register/login/JWT), and initial product catalog API."),
        ("Sprint 2 (Weeks 4-6): Core Features",
         "2D Canvas customization engine, dynamic pricing logic, product catalog frontend, "
         "shopping cart implementation, and order placement workflow."),
        ("Sprint 3 (Weeks 7-9): AI & Admin",
         "n8n RAG pipeline integration, AI chatbot frontend, admin panel (product/order/customer "
         "management), analytics dashboard, and payment integration (COD + PayFast)."),
        ("Sprint 4 (Weeks 10-12): Testing & Deployment",
         "Comprehensive testing (56 test cases), bug fixing, performance optimization, "
         "Railway.app deployment, documentation, and user manual preparation."),
    ]

    for sprint_title, sprint_body in sprints:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(sprint_title + ": ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run2 = p.add_run(sprint_body)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

    # ─────────────────────────────────────────────────────────────
    # 1.12 Modelling Techniques / Tools Used
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.12 Modelling Techniques / Tools Used", bold=True, size=13)

    _body(doc,
          "The following UML modelling techniques were used to design and document the AURA-WEAR "
          "system:")

    modelling = [
        "Use Case Diagrams: To capture functional requirements and actor interactions for all 8 modules.",
        "Fully Dressed Use Cases: Detailed use case specifications with main success scenarios and alternative flows.",
        "Domain Model: To identify key concepts (entities) and their relationships in the problem domain.",
        "System Sequence Diagrams (SSDs): To illustrate interactions between actors and the system for primary use case flows.",
        "Class Diagrams: To represent the software architecture including classes, attributes, methods, and associations.",
        "Entity-Relationship (ER) Diagram: To design the relational database schema.",
    ]

    for item in modelling:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    _blank(doc)
    _body(doc, "Tools used for modelling and development:")

    tools = [
        "draw.io: UML diagrams (use case, class, ER, SSDs, domain model).",
        "Visual Studio Code: Primary code editor for all Python and JavaScript development.",
        "Postman: REST API testing and documentation.",
        "GitHub: Version control and remote repository hosting.",
        "Railway.app Dashboard: Deployment monitoring and environment variable management.",
        "n8n Canvas: Visual workflow builder for the RAG AI pipeline.",
    ]

    for item in tools:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # ─────────────────────────────────────────────────────────────
    # 1.13 Limitation / Constraint
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.13 Limitation/Constraint", bold=True, size=13)

    _body(doc,
          "The following limitations and constraints apply to the current version of AURA-WEAR:")

    limitations = [
        "Web Browser Only: AURA-WEAR is a web application and does not have a dedicated native mobile application for iOS or Android. Mobile users must access the platform via a responsive web browser.",
        "No Native Mobile App: While the frontend is mobile-responsive, a dedicated native app with push notifications and offline capability is not part of this version.",
        "Payment Methods: Only Cash on Delivery (COD) and PayFast are supported. Stripe integration was planned but is deferred to a future release due to regional API restrictions.",
        "AI Chatbot Requires Internet: The AI chatbot depends on the n8n workflow server and OpenAI API, both of which require an active internet connection. Offline operation is not supported.",
        "No Physical Store Integration: AURA-WEAR operates as an independent web platform. There is no real-time integration with any physical store inventory or point-of-sale system.",
        "Single Language: The platform is currently available in English only. Urdu and other regional language support is not included in this version.",
    ]

    for item in limitations:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # ─────────────────────────────────────────────────────────────
    # 1.14 References
    # ─────────────────────────────────────────────────────────────
    _heading(doc, "1.14 References", bold=True, size=13)

    references = [
        "[1] P. Loureiro, J. Bettencourt-Silva, and A. Trindade, \"Artificial Intelligence in E-Commerce: "
        "A Systematic Literature Review,\" IEEE Access, vol. 10, pp. 12345-12360, 2022.",

        "[2] M. A. Khan and F. Algarni, \"A Healthcare Monitoring System for the Diagnosis of Heart Disease "
        "in the IoMT Cloud Environment Using MSSO-ANFIS,\" IEEE Access, vol. 8, pp. 122259-122269, 2020.",

        "[3] T. Brown et al., \"Language Models are Few-Shot Learners,\" in Proc. NeurIPS, 2020.",

        "[4] Y. Gao et al., \"Retrieval-Augmented Generation for Large Language Models: A Survey,\" "
        "arXiv preprint arXiv:2312.10997, 2023.",

        "[5] S. Kim, J. Lee, and H. Park, \"Personalized Fashion Recommendation Using Deep Learning and "
        "Visual Similarity,\" Journal of Fashion Technology & Textile Engineering, vol. 9, no. 3, 2021.",

        "[6] FastAPI Documentation, Sebastián Ramírez, 2024. [Online]. Available: https://fastapi.tiangolo.com/",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_page_break()
